#!/usr/bin/env python3
"""
Force LibreOffice Embedding - Direct approach
"""

import os
import sys
import subprocess
import time
import tempfile
import shutil
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, QLabel, QFileDialog
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QWindow
import win32gui
import win32process
import win32con

class ForceLibreOfficeEmbedder:
    def __init__(self):
        self.process = None
        self.window_handle = None
        self.temp_file = None
        
    def find_libreoffice(self):
        """Find LibreOffice executable"""
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        
        for path in paths:
            if os.path.exists(path):
                return path
        return None
    
    def create_test_document(self):
        """Create a test document if it doesn't exist"""
        if not os.path.exists("test_template.docx"):
            from docx import Document
            doc = Document()
            doc.add_heading('Test Document for Embedding', 0)
            doc.add_paragraph('This is a test document to verify LibreOffice embedding.')
            doc.add_paragraph('If you can see this embedded in the application, it works!')
            doc.save("test_template.docx")
            print("✅ Created test document")
    
    def launch_and_find_window(self, file_path):
        """Launch LibreOffice and find its window"""
        executable = self.find_libreoffice()
        if not executable:
            raise Exception("LibreOffice not found")
        
        # Create temporary copy
        temp_dir = tempfile.mkdtemp()
        self.temp_file = os.path.join(temp_dir, os.path.basename(file_path))
        shutil.copy2(file_path, self.temp_file)
        
        print(f"🚀 Launching LibreOffice with: {self.temp_file}")
        
        # Launch LibreOffice with multiple attempts
        launch_attempts = [
            [executable, self.temp_file],
            [executable, "--show", self.temp_file],
            [executable, "--writer", self.temp_file],
            [executable, "--nologo", "--norestore", self.temp_file]
        ]
        
        for i, cmd in enumerate(launch_attempts):
            try:
                print(f"🔄 Attempt {i+1}: {' '.join(cmd)}")
                
                # Kill any existing process
                if self.process:
                    try:
                        self.process.terminate()
                        self.process.wait(timeout=3)
                    except:
                        self.process.kill()
                
                # Launch new process
                self.process = subprocess.Popen(cmd)
                print(f"✅ Process started with PID: {self.process.pid}")
                
                # Wait and look for window
                for wait_time in range(15):
                    time.sleep(1)
                    print(f"⏳ Waiting for window... ({wait_time+1}/15)")
                    
                    window = self.find_libreoffice_window()
                    if window:
                        print(f"✅ Found window: {window}")
                        self.window_handle = window
                        return True
                
                print(f"❌ No window found with attempt {i+1}")
                
            except Exception as e:
                print(f"❌ Attempt {i+1} failed: {e}")
                continue
        
        return False
    
    def find_libreoffice_window(self):
        """Find LibreOffice window"""
        def enum_windows_callback(hwnd, windows):
            if win32gui.IsWindowVisible(hwnd):
                try:
                    _, pid = win32process.GetWindowThreadProcessId(hwnd)
                    
                    # Check if this is our process
                    if self.process and pid == self.process.pid:
                        window_text = win32gui.GetWindowText(hwnd)
                        class_name = win32gui.GetClassName(hwnd)
                        
                        print(f"Found our process window: '{window_text}' (class: {class_name})")
                        windows.append(hwnd)
                        return True
                    
                    # Also check for LibreOffice windows
                    window_text = win32gui.GetWindowText(hwnd)
                    class_name = win32gui.GetClassName(hwnd)
                    
                    libreoffice_keywords = ['libreoffice', 'writer', 'calc', 'impress', 'soffice', 'document']
                    if any(keyword in window_text.lower() for keyword in libreoffice_keywords) or \
                       any(keyword in class_name.lower() for keyword in libreoffice_keywords):
                        print(f"Found LibreOffice window: '{window_text}' (PID: {pid})")
                        windows.append(hwnd)
                        
                except Exception as e:
                    pass
            return True
        
        windows = []
        win32gui.EnumWindows(enum_windows_callback, windows)
        
        if windows:
            return windows[0]
        return None
    
    def embed_window(self, parent_widget):
        """Embed the LibreOffice window into the parent widget"""
        if not self.window_handle:
            raise Exception("No window handle available")
        
        print(f"🔗 Embedding window handle: {self.window_handle}")
        
        # Get window info
        window_text = win32gui.GetWindowText(self.window_handle)
        print(f"📋 Window title: '{window_text}'")
        
        # Set window style for embedding
        try:
            style = win32gui.GetWindowLong(self.window_handle, win32con.GWL_EXSTYLE)
            win32gui.SetWindowLong(self.window_handle, win32con.GWL_EXSTYLE, style | win32con.WS_EX_MDICHILD)
            print("✅ Window style set for embedding")
        except Exception as e:
            print(f"⚠️ Could not set window style: {e}")
        
        # Create QWindow from handle
        window = QWindow.fromWinId(self.window_handle)
        
        # Create container widget
        container = QWidget.createWindowContainer(window, parent_widget)
        
        print("✅ Window embedded successfully!")
        return container
    
    def cleanup(self):
        """Clean up resources"""
        if self.process:
            try:
                self.process.terminate()
                self.process.wait(timeout=5)
            except:
                self.process.kill()
            self.process = None
        
        if self.temp_file and os.path.exists(self.temp_file):
            try:
                os.remove(self.temp_file)
                os.rmdir(os.path.dirname(self.temp_file))
            except:
                pass

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.embedder = ForceLibreOfficeEmbedder()
        self.embedded_widget = None
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Force LibreOffice Embedding")
        self.setGeometry(100, 100, 1200, 800)
        
        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Create layout
        layout = QVBoxLayout(central_widget)
        
        # Title
        title = QLabel("Force LibreOffice Embedding Test")
        title.setStyleSheet("font-size: 18px; font-weight: bold; padding: 10px;")
        layout.addWidget(title)
        
        # Status label
        self.status_label = QLabel("Ready to test embedding")
        self.status_label.setStyleSheet("color: #666666; padding: 5px;")
        layout.addWidget(self.status_label)
        
        # Buttons
        button_layout = QVBoxLayout()
        
        # Test with sample document
        test_btn = QPushButton("Test with Sample Document")
        test_btn.clicked.connect(self.test_with_sample)
        button_layout.addWidget(test_btn)
        
        # Upload document
        upload_btn = QPushButton("Upload Document")
        upload_btn.clicked.connect(self.upload_document)
        button_layout.addWidget(upload_btn)
        
        layout.addLayout(button_layout)
        
        # Preview area
        self.preview_area = QWidget()
        self.preview_layout = QVBoxLayout(self.preview_area)
        layout.addWidget(self.preview_area)
        
        # Create test document
        self.embedder.create_test_document()
    
    def test_with_sample(self):
        """Test with sample document"""
        self.status_label.setText("Testing with sample document...")
        self.load_document("test_template.docx")
    
    def upload_document(self):
        """Upload a document"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Document", "", "Word Documents (*.docx *.doc)"
        )
        
        if file_path:
            self.status_label.setText(f"Loading: {os.path.basename(file_path)}")
            self.load_document(file_path)
    
    def load_document(self, file_path):
        """Load and embed document"""
        # Clean up previous embedding
        if self.embedded_widget:
            self.embedded_widget.setParent(None)
            self.embedded_widget = None
        
        if self.embedder:
            self.embedder.cleanup()
        
        try:
            # Launch LibreOffice and find window
            if self.embedder.launch_and_find_window(file_path):
                # Embed the window
                self.embedded_widget = self.embedder.embed_window(self.preview_area)
                self.preview_layout.addWidget(self.embedded_widget)
                
                self.status_label.setText(f"✅ Document embedded: {os.path.basename(file_path)}")
            else:
                self.status_label.setText("❌ Could not find LibreOffice window")
                
        except Exception as e:
            self.status_label.setText(f"❌ Error: {str(e)}")
            print(f"Error: {e}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    window = MainWindow()
    window.show()
    
    print("🚀 Force LibreOffice Embedding Test Application")
    print("📋 Click 'Test with Sample Document' to begin")
    
    sys.exit(app.exec())
