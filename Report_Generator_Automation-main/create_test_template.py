#!/usr/bin/env python3
"""
Create a test DOCX template with headers, footers, and watermarks
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor

def create_test_template():
    """Create a test DOCX template with headers, footers, and watermarks"""
    
    # Create a new document
    doc = Document()
    
    # Set up page margins
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "DELTAPHI LABS - SECURITY ASSESSMENT REPORT"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.name = 'Calibri'
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add page number to footer
    footer_para.add_run(" ")
    page_num_run = footer_para.add_run()
    page_num_run._element.append(OxmlElement('w:fldChar'))
    page_num_run._element.append(OxmlElement('w:instrText'))
    page_num_run._element.find('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}).text = 'PAGE'
    page_num_run._element.append(OxmlElement('w:fldChar'))
    
    # Add watermark (as background text)
    # Note: This is a simplified watermark implementation
    watermark_para = doc.add_paragraph()
    watermark_para.text = "CONFIDENTIAL"
    watermark_run = watermark_para.runs[0]
    watermark_run.font.name = 'Arial'
    watermark_run.font.size = Pt(72)
    watermark_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title page content
    doc.add_page_break()
    
    # Title page
    title_para = doc.add_paragraph()
    title_para.text = "SECURITY ASSESSMENT REPORT"
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add some spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add company info
    company_para = doc.add_paragraph()
    company_para.text = "DELTAPHI LABS"
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_para.runs[0]
    company_run.font.name = 'Calibri'
    company_run.font.size = Pt(16)
    company_run.font.bold = True
    
    # Add date placeholder
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.text = "Date: [To be filled]"
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(12)
    
    # Add template page content
    doc.add_page_break()
    
    # Template page with sample content
    template_para = doc.add_paragraph()
    template_para.text = "TEMPLATE PAGE - This page contains the template formatting that will be preserved"
    template_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    template_run = template_para.runs[0]
    template_run.font.name = 'Calibri'
    template_run.font.size = Pt(14)
    template_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph("This template includes:")
    doc.add_paragraph("• Headers with company name")
    doc.add_paragraph("• Footers with page numbers")
    doc.add_paragraph("• Watermark (CONFIDENTIAL)")
    doc.add_paragraph("• Proper page margins and formatting")
    
    # Save the template
    template_path = "sample_document_template.docx"
    doc.save(template_path)
    print(f"✅ Test template created: {template_path}")
    print("This template includes headers, footers, and watermarks that will be preserved when generating reports.")
    print("IMPORTANT: Content will now start from page 1, not page 3!")
    
    return template_path

if __name__ == "__main__":
    create_test_template()
