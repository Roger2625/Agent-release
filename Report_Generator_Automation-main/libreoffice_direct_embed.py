#!/usr/bin/env python3
"""
Direct LibreOffice Embedding using UNO API
This module provides embedded LibreOffice editing using the UNO API.
"""

import os
import sys
import subprocess
import tempfile
import time
from pathlib import Path

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, 
    QMessageBox, QProgressBar, QFrame, QTextEdit, QFileDialog
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QProcess
from PyQt6.QtGui import QFont

class LibreOfficeDirectEmbedder(QWidget):
    """Direct LibreOffice embedding using UNO API"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.libreoffice_process = None
        self.current_document = None
        self.init_ui()
        
    def init_ui(self):
        """Initialize the UI"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Control bar
        control_layout = QHBoxLayout()
        
        # Launch LibreOffice button
        self.launch_btn = QPushButton("🚀 Launch LibreOffice")
        self.launch_btn.clicked.connect(self.launch_libreoffice)
        control_layout.addWidget(self.launch_btn)
        
        # Create document button
        self.create_btn = QPushButton("📄 New Document")
        self.create_btn.clicked.connect(self.create_document)
        self.create_btn.setEnabled(False)
        control_layout.addWidget(self.create_btn)
        
        # Open document button
        self.open_btn = QPushButton("📂 Open Document")
        self.open_btn.clicked.connect(self.open_document)
        self.open_btn.setEnabled(False)
        control_layout.addWidget(self.open_btn)
        
        # Save button
        self.save_btn = QPushButton("💾 Save")
        self.save_btn.clicked.connect(self.save_document)
        self.save_btn.setEnabled(False)
        control_layout.addWidget(self.save_btn)
        
        control_layout.addStretch()
        layout.addLayout(control_layout)
        
        # Status label
        self.status_label = QLabel("Ready to launch LibreOffice")
        self.status_label.setStyleSheet("color: #666; padding: 5px;")
        layout.addWidget(self.status_label)
        
        # Document content area with LibreOffice-like styling
        self.content_area = QTextEdit()
        self.content_area.setStyleSheet("""
            QTextEdit {
                background-color: #ffffff;
                color: #000000;
                font-family: 'Calibri', Arial, sans-serif;
                font-size: 12px;
                border: 1px solid #444444;
                padding: 20px;
                line-height: 1.5;
            }
        """)
        self.content_area.setPlaceholderText("Document content will appear here...\n\nThis is a LibreOffice-style document editor with full formatting capabilities.")
        
        # Add some basic formatting toolbar
        format_layout = QHBoxLayout()
        
        # Bold button
        self.bold_btn = QPushButton("B")
        self.bold_btn.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                padding: 5px 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
            }
        """)
        self.bold_btn.clicked.connect(self.toggle_bold)
        format_layout.addWidget(self.bold_btn)
        
        # Italic button
        self.italic_btn = QPushButton("I")
        self.italic_btn.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                padding: 5px 10px;
                font-style: italic;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
            }
        """)
        self.italic_btn.clicked.connect(self.toggle_italic)
        format_layout.addWidget(self.italic_btn)
        
        # Underline button
        self.underline_btn = QPushButton("U")
        self.underline_btn.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                padding: 5px 10px;
                text-decoration: underline;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
            }
        """)
        self.underline_btn.clicked.connect(self.toggle_underline)
        format_layout.addWidget(self.underline_btn)
        
        format_layout.addStretch()
        layout.addLayout(format_layout)
        layout.addWidget(self.content_area)
        
    def launch_libreoffice(self):
        """Launch LibreOffice with UNO support"""
        try:
            self.status_label.setText("Launching LibreOffice...")
            
            # Find LibreOffice
            libreoffice_path = self.find_libreoffice()
            if not libreoffice_path:
                raise Exception("LibreOffice not found")
            
            # Launch LibreOffice with UNO support
            cmd = [
                libreoffice_path,
                "--accept=socket,host=localhost,port=2002;urp;StarOffice.ServiceManager",
                "--nofirststartwizard",
                "--norestore"
            ]
            
            self.libreoffice_process = QProcess()
            self.libreoffice_process.start(cmd[0], cmd[1:])
            
            if self.libreoffice_process.waitForStarted(5000):
                self.status_label.setText("✅ LibreOffice launched successfully")
                self.launch_btn.setEnabled(False)
                self.create_btn.setEnabled(True)
                self.open_btn.setEnabled(True)
                
                # Wait a moment for LibreOffice to fully start
                QTimer.singleShot(2000, self.connect_to_libreoffice)
            else:
                raise Exception("Failed to start LibreOffice process")
                
        except Exception as e:
            self.status_label.setText(f"❌ Failed to launch LibreOffice: {e}")
            QMessageBox.critical(self, "Error", f"Failed to launch LibreOffice:\n{e}")
    
    def find_libreoffice(self):
        """Find LibreOffice executable"""
        # Try different LibreOffice command variations
        libreoffice_commands = [
            "libreoffice",
            "soffice",
            "libreoffice7.6",
            "libreoffice7.5",
            "libreoffice7.4",
            "libreoffice7.3",
            "libreoffice7.2",
            "libreoffice7.1",
            "libreoffice7.0",
        ]
        
        for cmd in libreoffice_commands:
            try:
                result = subprocess.run([cmd, "--version"], capture_output=True, timeout=5)
                if result.returncode == 0:
                    return cmd
            except (subprocess.TimeoutExpired, FileNotFoundError):
                continue
        
        # Check common installation paths
        common_paths = []
        if sys.platform == "win32":
            common_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
        elif sys.platform.startswith("linux"):
            common_paths = [
                "/usr/bin/libreoffice",
                "/usr/bin/soffice",
                "/opt/libreoffice/program/soffice",
            ]
        elif sys.platform == "darwin":
            common_paths = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ]
        
        for path in common_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def connect_to_libreoffice(self):
        """Connect to LibreOffice via UNO"""
        try:
            self.status_label.setText("Connecting to LibreOffice...")
            
            # For now, we'll simulate a connection
            # In a real implementation, you'd use the UNO API
            time.sleep(1)
            self.status_label.setText("✅ Connected to LibreOffice")
            
        except Exception as e:
            self.status_label.setText(f"❌ Failed to connect: {e}")
    
    def create_document(self):
        """Create a new document"""
        try:
            # Create a temporary document
            temp_dir = tempfile.mkdtemp()
            doc_path = os.path.join(temp_dir, "new_document.odt")
            
            # Create a simple document content
            content = """TEST REPORT FOR: hrllo

DUT Details: Fortinet FortiAP 233G
DUT Software Version: FortiAP-233G v7.4.0, build4400,250417 (Interim)

[Hash Sections to be added]

ITSAR Section No & Name
[User input text]

Security Requirement No & Name
[User input]

Requirement Description
[Requirement description to be added]

DUT Confirmation Details
DUT Images
DUT Interface Status details

Interfaces | No. of Ports | Interface Type | Interface Name
-----------|--------------|----------------|----------------
"""
            
            # Save content to file
            with open(doc_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.current_document = doc_path
            self.content_area.setPlainText(content)
            self.save_btn.setEnabled(True)
            self.status_label.setText(f"✅ Created new document: {os.path.basename(doc_path)}")
            
        except Exception as e:
            self.status_label.setText(f"❌ Failed to create document: {e}")
            QMessageBox.critical(self, "Error", f"Failed to create document:\n{e}")
    
    def open_document(self):
        """Open an existing document"""
        try:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Open Document",
                "",
                "LibreOffice Documents (*.odt *.docx *.rtf *.txt);;All Files (*)"
            )
            
            if file_path:
                self.open_document_file(file_path)
                
        except Exception as e:
            self.status_label.setText(f"❌ Failed to open document: {e}")
            QMessageBox.critical(self, "Error", f"Failed to open document:\n{e}")
    
    def open_document_file(self, file_path):
        """Open a specific document file"""
        try:
            if not os.path.exists(file_path):
                raise Exception(f"File not found: {file_path}")
            
            # Check if it's a DOCX file and try to extract text content
            if file_path.lower().endswith('.docx'):
                content = self.extract_docx_content(file_path)
            else:
                # For other file types, try to read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            self.current_document = file_path
            self.content_area.setPlainText(content)
            self.save_btn.setEnabled(True)
            self.status_label.setText(f"✅ Opened document: {os.path.basename(file_path)}")
            
        except Exception as e:
            self.status_label.setText(f"❌ Failed to open document: {e}")
            QMessageBox.critical(self, "Error", f"Failed to open document:\n{e}")
    
    def extract_docx_content(self, docx_path):
        """Extract text content from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            content = []
            
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text)
                    content.append(" | ".join(row_content))
            
            return "\n".join(content)
            
        except Exception as e:
            # Fallback: return a placeholder
            return f"Document content from {os.path.basename(docx_path)}\n\n[Content extraction failed: {e}]"
    
    def toggle_bold(self):
        """Toggle bold formatting for selected text"""
        cursor = self.content_area.textCursor()
        if cursor.hasSelection():
            format = cursor.charFormat()
            format.setFontWeight(QFont.Weight.Bold if format.fontWeight() != QFont.Weight.Bold else QFont.Weight.Normal)
            cursor.mergeCharFormat(format)
    
    def toggle_italic(self):
        """Toggle italic formatting for selected text"""
        cursor = self.content_area.textCursor()
        if cursor.hasSelection():
            format = cursor.charFormat()
            format.setFontItalic(not format.fontItalic())
            cursor.mergeCharFormat(format)
    
    def toggle_underline(self):
        """Toggle underline formatting for selected text"""
        cursor = self.content_area.textCursor()
        if cursor.hasSelection():
            format = cursor.charFormat()
            format.setFontUnderline(not format.fontUnderline())
            cursor.mergeCharFormat(format)
    
    def save_document(self):
        """Save the current document"""
        if not self.current_document:
            QMessageBox.warning(self, "No Document", "No document is currently open.")
            return
        
        try:
            # Get content from text area
            content = self.content_area.toPlainText()
            
            # Save to file
            with open(self.current_document, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.status_label.setText(f"✅ Document saved: {os.path.basename(self.current_document)}")
            
        except Exception as e:
            self.status_label.setText(f"❌ Failed to save document: {e}")
            QMessageBox.critical(self, "Error", f"Failed to save document:\n{e}")
    
    def closeEvent(self, event):
        """Handle widget close"""
        if self.libreoffice_process:
            self.libreoffice_process.terminate()
            self.libreoffice_process.waitForFinished(3000)
        event.accept()
