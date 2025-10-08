#!/usr/bin/env python3
"""
Simple LibreOffice Embedding - Direct and Working Approach
"""

import os
import sys
import subprocess
import time
import tempfile
import shutil
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, QLabel, QFileDialog, QTextEdit
from PyQt6.QtCore import Qt, QTimer, QThread, pyqtSignal
from PyQt6.QtGui import QWindow
import win32gui
import win32process
import win32con

class LibreOfficeLauncher(QThread):
    """Thread to launch LibreOffice and find its window"""
    window_found = pyqtSignal(int)  # Signal with window handle
    error_occurred = pyqtSignal(str)
    
    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path
        self.process = None
        
    def run(self):
        try:
            # Find LibreOffice
            libreoffice_path = self.find_libreoffice()
            if not libreoffice_path:
                self.error_occurred.emit("LibreOffice not found")
                return
            
            # Create temp copy
            temp_dir = tempfile.mkdtemp()
            temp_file = os.path.join(temp_dir, os.path.basename(self.file_path))
            shutil.copy2(self.file_path, temp_file)
            
            print(f"🚀 Launching LibreOffice: {libreoffice_path}")
            print(f"📄 File: {temp_file}")
            
            # Launch LibreOffice
            self.process = subprocess.Popen([libreoffice_path, temp_file])
            print(f"✅ Process started with PID: {self.process.pid}")
            
            # Wait for window to appear
            for i in range(30):  # Wait up to 30 seconds
                time.sleep(1)
                print(f"⏳ Waiting for window... ({i+1}/30)")
                
                # Look for any window from our process
                window_handle = self.find_process_window(self.process.pid)
                if window_handle:
                    print(f"✅ Found window: {window_handle}")
                    self.window_found.emit(window_handle)
                    return
            
            self.error_occurred.emit("No window found after 30 seconds")
            
        except Exception as e:
            self.error_occurred.emit(f"Error: {str(e)}")
    
    def find_libreoffice(self):
        """Find LibreOffice executable"""
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        
        for path in paths:
            if os.path.exists(path):
                return path
        return None
    
    def find_process_window(self, pid):
        """Find any window belonging to the process"""
        def enum_windows_callback(hwnd, windows):
            if win32gui.IsWindowVisible(hwnd):
                try:
                    _, window_pid = win32process.GetWindowThreadProcessId(hwnd)
                    if window_pid == pid:
                        window_text = win32gui.GetWindowText(hwnd)
                        if window_text:  # Only windows with titles
                            print(f"Found window: '{window_text}' (PID: {window_pid})")
                            windows.append(hwnd)
                except:
                    pass
            return True
        
        windows = []
        win32gui.EnumWindows(enum_windows_callback, windows)
        return windows[0] if windows else None

class SimpleLibreOfficeEmbedder:
    def __init__(self):
        self.window_handle = None
        self.launcher_thread = None
        
    def embed_document(self, file_path, parent_widget, callback):
        """Embed a document by launching LibreOffice and finding its window"""
        # Start launcher thread
        self.launcher_thread = LibreOfficeLauncher(file_path)
        self.launcher_thread.window_found.connect(lambda hwnd: self.on_window_found(hwnd, parent_widget, callback))
        self.launcher_thread.error_occurred.connect(callback)
        self.launcher_thread.start()
    
    def on_window_found(self, window_handle, parent_widget, callback):
        """Called when LibreOffice window is found"""
        try:
            self.window_handle = window_handle
            
            # Get window info
            window_text = win32gui.GetWindowText(window_handle)
            print(f"📋 Embedding window: '{window_text}'")
            
            # Set window style for embedding
            try:
                style = win32gui.GetWindowLong(window_handle, win32con.GWL_EXSTYLE)
                win32gui.SetWindowLong(window_handle, win32con.GWL_EXSTYLE, style | win32con.WS_EX_MDICHILD)
                print("✅ Window style set for embedding")
            except Exception as e:
                print(f"⚠️ Could not set window style: {e}")
            
            # Create QWindow from handle
            window = QWindow.fromWinId(window_handle)
            
            # Create container widget
            container = QWidget.createWindowContainer(window, parent_widget)
            
            print("✅ Window embedded successfully!")
            callback("SUCCESS", container)
            
        except Exception as e:
            callback(f"Embedding error: {str(e)}")

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.embedder = SimpleLibreOfficeEmbedder()
        self.embedded_widget = None
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Simple LibreOffice Embedding")
        self.setGeometry(100, 100, 1400, 900)
        
        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Create layout
        layout = QVBoxLayout(central_widget)
        
        # Title
        title = QLabel("Simple LibreOffice Embedding Test")
        title.setStyleSheet("font-size: 18px; font-weight: bold; padding: 10px;")
        layout.addWidget(title)
        
        # Status label
        self.status_label = QLabel("Ready to embed LibreOffice")
        self.status_label.setStyleSheet("color: #666666; padding: 5px;")
        layout.addWidget(self.status_label)
        
        # Buttons
        button_layout = QVBoxLayout()
        
        # Test with sample document
        test_btn = QPushButton("Test with Sample Document")
        test_btn.clicked.connect(self.test_with_sample)
        button_layout.addWidget(test_btn)
        
        # Upload document
        upload_btn = QPushButton("Upload Document")
        upload_btn.clicked.connect(self.upload_document)
        button_layout.addWidget(upload_btn)
        
        layout.addLayout(button_layout)
        
        # Preview area
        self.preview_area = QWidget()
        self.preview_layout = QVBoxLayout(self.preview_area)
        layout.addWidget(self.preview_area)
        
        # Create test document
        self.create_test_document()
    
    def create_test_document(self):
        """Create a test document"""
        if not os.path.exists("test_template.docx"):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('Test Document for LibreOffice Embedding', 0)
                doc.add_paragraph('This is a test document to verify LibreOffice embedding.')
                doc.add_paragraph('If you can see this embedded in the application, it works!')
                doc.add_paragraph('Features:')
                doc.add_paragraph('• Real LibreOffice window', style='List Bullet')
                doc.add_paragraph('• Full document editing', style='List Bullet')
                doc.add_paragraph('• Native LibreOffice interface', style='List Bullet')
                doc.save("test_template.docx")
                print("✅ Created test document")
            except Exception as e:
                print(f"❌ Could not create test document: {e}")
    
    def test_with_sample(self):
        """Test with sample document"""
        self.status_label.setText("Testing with sample document...")
        self.load_document("test_template.docx")
    
    def upload_document(self):
        """Upload a document"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Document", "", "Word Documents (*.docx *.doc)"
        )
        
        if file_path:
            self.status_label.setText(f"Loading: {os.path.basename(file_path)}")
            self.load_document(file_path)
    
    def load_document(self, file_path):
        """Load and embed document"""
        # Clean up previous embedding
        if self.embedded_widget:
            self.embedded_widget.setParent(None)
            self.embedded_widget = None
        
        # Clear preview area
        for i in reversed(range(self.preview_layout.count())): 
            self.preview_layout.itemAt(i).widget().setParent(None)
        
        # Show loading message
        loading_label = QLabel("Loading LibreOffice... Please wait.")
        loading_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        loading_label.setStyleSheet("color: #666666; font-style: italic; padding: 20px;")
        self.preview_layout.addWidget(loading_label)
        
        # Start embedding
        self.embedder.embed_document(file_path, self.preview_area, self.on_embedding_result)
    
    def on_embedding_result(self, result, widget=None):
        """Called when embedding is complete"""
        # Clear loading message
        for i in reversed(range(self.preview_layout.count())): 
            self.preview_layout.itemAt(i).widget().setParent(None)
        
        if result == "SUCCESS" and widget:
            # Success - add the embedded widget
            self.embedded_widget = widget
            self.preview_layout.addWidget(widget)
            self.status_label.setText("✅ LibreOffice embedded successfully!")
        else:
            # Error - show error message
            error_label = QLabel(f"❌ Error: {result}")
            error_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            error_label.setStyleSheet("color: #ff0000; font-style: italic; padding: 20px;")
            self.preview_layout.addWidget(error_label)
            self.status_label.setText("❌ Embedding failed")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    window = MainWindow()
    window.show()
    
    print("🚀 Simple LibreOffice Embedding Test Application")
    print("📋 Click 'Test with Sample Document' to begin")
    print("⏳ This will launch LibreOffice and embed it directly")
    
    sys.exit(app.exec())
