#!/usr/bin/env python3
"""
Enhanced Security Assessment Report Generator
Generates comprehensive security assessment reports with multiple sections
"""

import os
import sys
from datetime import datetime
from docx import Document
import logging
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor


# --- Logging Setup ---
def setup_document_logging():
    """Setup logging to save debug logs to ~/.wingzai/logs/ with timestamps"""
    try:
        # Create logs directory
        logs_dir = os.path.expanduser("~/.wingzai/logs")
        os.makedirs(logs_dir, exist_ok=True)
        
        # Create log filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename = f"enhanced_document_generator_{timestamp}.log"
        log_path = os.path.join(logs_dir, log_filename)
        
        # Clear any existing handlers to avoid conflicts
        for handler in logging.root.handlers[:]:
            logging.root.removeHandler(handler)
        
        # Configure logging
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_path, encoding='utf-8'),
                logging.StreamHandler(sys.stdout)  # Also log to console
            ],
            force=True  # Force reconfiguration
        )
        
        # Create logger for enhanced document generator
        logger = logging.getLogger('enhanced_document_generator')
        logger.setLevel(logging.DEBUG)
        logger.info(f"Enhanced Document Generator logging initialized. Log file: {log_path}")
        
        return logger
    except Exception as e:
        print(f"Failed to setup document logging: {e}")
        return None

# Initialize logging
document_logger = setup_document_logging()

def log_debug(message):
    """Log debug message to both file and console"""
    if document_logger:
        document_logger.debug(message)
    else:
        print(f"DEBUG: {message}")

def log_info(message):
    """Log info message to both file and console"""
    if document_logger:
        document_logger.info(message)
    else:
        print(f"INFO: {message}")

def log_warning(message):
    """Log warning message to both file and console"""
    if document_logger:
        document_logger.warning(message)
    else:
        print(f"WARNING: {message}")

def log_error(message):
    """Log error message to both file and console"""
    if document_logger:
        document_logger.error(message)
    else:
        print(f"ERROR: {message}")


class DocumentHelper:
    """Helper class for document formatting and styling"""
    
    @staticmethod
    def set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=False, color=None):
        """Sets font style for a run"""
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # Set color if provided
        if color:
            from docx.shared import RGBColor
            # Convert hex color to RGB
            if color.startswith('#'):
                color = color[1:]
            if len(color) == 6:
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
    
    @staticmethod
    def safe_add_paragraph_with_style(doc, text="", style_name=None):
        """Safely add a paragraph with a style, falling back to Normal if style doesn't exist"""
        try:
            if style_name and style_name in [style.name for style in doc.styles]:
                return doc.add_paragraph(text, style=style_name)
            else:
                # Fallback to Normal style and manually format as bullet point
                paragraph = doc.add_paragraph(text)
                if style_name == 'List Bullet':
                    # Manually create bullet point formatting
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    # Add bullet character
                    if text:
                        paragraph.text = f"• {text}"
                return paragraph
        except Exception:
            # Ultimate fallback - just add normal paragraph
            return doc.add_paragraph(text)
    
    @staticmethod
    def add_smart_subheading_paragraph(doc, text=""):
        """Add a paragraph with smart formatting based on text content using tags:
        - <bullets>*text* → • text (bullet and text inside * are bold)
        - <n>*text* → 1. text (number and text inside * are bold)
        - <->*text* → - text (dash and text inside * are bold)
        - <number>*text* → number. text (number and text inside * are bold, e.g., <3>*hello* → 3. hello)
        - <alphabet>*text* → alphabet. text (letter and text inside * are bold, e.g., <B>*hello* → B. hello)
        - *<bullets>text* → • text (formatting tags inside bold markers)
        - Plain text with *text* → text (only text inside * is bold)
        """
        if not text:
            return doc.add_paragraph()
        
        text_stripped = text.strip()
        paragraph = doc.add_paragraph()
        
        # Check if the entire text is wrapped in bold markers
        if text_stripped.startswith('*') and text_stripped.endswith('*') and text_stripped.count('*') >= 2:
            # Extract content between the outer asterisks
            inner_content = text_stripped[1:-1].strip()
            
            # Process the inner content for formatting tags
            if inner_content.startswith('<bullets>'):
                content = inner_content[9:].strip()  # Remove '<bullets>'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                # When content is inside *<bullets>...* format, make the entire content bold
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "• ", f"*{content}*")
                
            elif inner_content.startswith('<n>'):
                content = inner_content[3:].strip()  # Remove '<n>'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "1. ", content)
                
            elif inner_content.startswith('<->'):
                content = inner_content[3:].strip()  # Remove '<->'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "- ", content)
                
            elif inner_content.startswith('<') and '>' in inner_content:
                tag_end = inner_content.find('>')
                if tag_end > 1:
                    tag_content = inner_content[1:tag_end]  # Get content between < and >
                    # Check if it's a number (including multi-digit numbers)
                    if tag_content.isdigit():
                        number = tag_content
                        content = inner_content[tag_end + 1:].strip()  # Get text after >
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                        DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, f"{number}. ", content)
                    # Check if it's a single alphabet character (A-Z, a-z)
                    elif len(tag_content) == 1 and tag_content.isalpha():
                        alphabet = tag_content.upper()  # Convert to uppercase for consistency
                        content = inner_content[tag_end + 1:].strip()  # Get text after >
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                        DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, f"{alphabet}. ", content)
                    else:
                        # Not a number or single alphabet, treat as plain text
                        DocumentHelper.add_formatted_text_with_bold(paragraph, inner_content)
                else:
                    # Invalid tag format, treat as plain text
                    DocumentHelper.add_formatted_text_with_bold(paragraph, inner_content)
                    
            elif inner_content.startswith('-'):
                content = inner_content[1:].strip()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "• ", content)
                
            else:
                # No formatting tags, just apply bold formatting to the original text
                DocumentHelper.add_formatted_text_with_bold(paragraph, text_stripped)
        
        # Check for tag-based formatting (not wrapped in bold)
        elif text_stripped.startswith('<bullets>'):
            # Remove the tag and add as bullet point with normal formatting
            content = text_stripped[9:].strip()  # Remove '<bullets>'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "• ", content)
            
        elif text_stripped.startswith('<n>'):
            # Remove the tag and add as numbered list with normal formatting
            content = text_stripped[3:].strip()  # Remove '<n>'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "1. ", content)
            
        elif text_stripped.startswith('<->'):
            # Remove the tag and add as dash list with normal formatting
            content = text_stripped[3:].strip()  # Remove '<->'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "- ", content)
            
        elif text_stripped.startswith('<note>'):
            # Remove the tag and add as note with red "Note:" prefix
            content = text_stripped[6:].strip()  # Remove '<note>'
            # Add "Note:" prefix in red
            note_prefix_run = paragraph.add_run("Note: ")
            DocumentHelper.set_font_style(note_prefix_run, font_name='Calibri', font_size=12, bold=True, color='#FF0000')  # Red color
            # Add content with markdown formatting support
            DocumentHelper.add_formatted_text_with_bold(paragraph, content)
            
        elif text_stripped.startswith('<') and '>' in text_stripped:
            # Check for custom numbered/alphabet tags like <3>*hello* or <B>*hello*
            tag_end = text_stripped.find('>')
            if tag_end > 1:
                tag_content = text_stripped[1:tag_end]  # Get content between < and >
                # Check if it's a number
                if tag_content.isdigit():
                    number = tag_content
                    content = text_stripped[tag_end + 1:].strip()  # Get text after >
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, f"{number}. ", content)
                # Check if it's a single alphabet character (A-Z, a-z)
                elif len(tag_content) == 1 and tag_content.isalpha():
                    alphabet = tag_content.upper()  # Convert to uppercase for consistency
                    content = text_stripped[tag_end + 1:].strip()  # Get text after >
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, f"{alphabet}. ", content)
                else:
                    # Not a number or single alphabet, treat as plain text
                    DocumentHelper.add_formatted_text_with_bold(paragraph, text)
            else:
                # Invalid tag format, treat as plain text
                DocumentHelper.add_formatted_text_with_bold(paragraph, text)
            
        else:
            # Plain text - use star-based bold formatting
            DocumentHelper.add_formatted_text_with_bold(paragraph, text)
        
        return paragraph
    
    @staticmethod
    def add_formatted_text_with_bold(paragraph, text):
        """Add text to paragraph with selective bold formatting for text inside * markers"""
        # Split text by * markers to identify bold sections
        parts = text.split('*')
        
        for i, part in enumerate(parts):
            if part:  # Skip empty parts
                run = paragraph.add_run(part)
                # Odd indices (1, 3, 5...) are inside * markers, so make them bold
                is_bold = (i % 2 == 1)
                DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=is_bold)
    
    @staticmethod
    def add_formatted_text_with_bold_and_prefix(paragraph, prefix, content):
        """Add text with bold prefix and selective bold formatting for content inside * markers"""
        # Add prefix in bold
        prefix_run = paragraph.add_run(prefix)
        DocumentHelper.set_font_style(prefix_run, font_name='Calibri', font_size=12, bold=True)
        
        # Add content with selective bold formatting
        DocumentHelper.add_formatted_text_with_bold(paragraph, content)
    
    @staticmethod
    def add_formatted_text_with_normal_prefix(paragraph, prefix, content):
        """Add text with normal prefix and selective bold formatting for content inside * markers"""
        # Add prefix in normal (not bold)
        prefix_run = paragraph.add_run(prefix)
        DocumentHelper.set_font_style(prefix_run, font_name='Calibri', font_size=12, bold=False)
        
        # Add content with selective bold formatting
        DocumentHelper.add_formatted_text_with_bold(paragraph, content)
    
    @staticmethod
    def add_smart_subheading_paragraph_with_spacing(doc, text="", line_spacing=1.0):
        """Add a paragraph with smart formatting and custom line spacing"""
        if not text:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = line_spacing
            return paragraph
        
        text_stripped = text.strip()
        paragraph = doc.add_paragraph()
        
        # Set line spacing
        paragraph.paragraph_format.line_spacing = line_spacing
        
        # Check if the entire text is wrapped in bold markers
        if text_stripped.startswith('*') and text_stripped.endswith('*') and text_stripped.count('*') >= 2:
            # Extract content between the outer asterisks
            inner_content = text_stripped[1:-1].strip()
            
            # Process the inner content for formatting tags
            if inner_content.startswith('<bullets>'):
                content = inner_content[9:].strip()  # Remove '<bullets>'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                # When content is inside *<bullets>...* format, make the entire content bold
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "• ", f"*{content}*")
                
            elif inner_content.startswith('<n>'):
                content = inner_content[3:].strip()  # Remove '<n>'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "1. ", content)
                
            elif inner_content.startswith('<->'):
                content = inner_content[3:].strip()  # Remove '<->'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "- ", content)
                
            elif inner_content.startswith('<') and '>' in inner_content:
                tag_end = inner_content.find('>')
                if tag_end > 1:
                    tag_content = inner_content[1:tag_end]  # Get content between < and >
                    # Check if it's a number (including multi-digit numbers)
                    if tag_content.isdigit():
                        number = tag_content
                        content = inner_content[tag_end + 1:].strip()  # Get text after >
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                        DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, f"{number}. ", content)
                    # Check if it's a single alphabet character (A-Z, a-z)
                    elif len(tag_content) == 1 and tag_content.isalpha():
                        alphabet = tag_content.upper()  # Convert to uppercase for consistency
                        content = inner_content[tag_end + 1:].strip()  # Get text after >
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                        DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, f"{alphabet}. ", content)
                    else:
                        # Not a number or single alphabet, treat as plain text
                        DocumentHelper.add_formatted_text_with_bold(paragraph, inner_content)
                else:
                    # Invalid tag format, treat as plain text
                    DocumentHelper.add_formatted_text_with_bold(paragraph, inner_content)
                    
            elif inner_content.startswith('-'):
                content = inner_content[1:].strip()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                DocumentHelper.add_formatted_text_with_bold_and_prefix(paragraph, "• ", content)
                
            else:
                # No formatting tags, just apply bold formatting to the original text
                DocumentHelper.add_formatted_text_with_bold(paragraph, text_stripped)
        
        # Check for tag-based formatting (not wrapped in bold)
        elif text_stripped.startswith('<bullets>'):
            # Remove the tag and add as bullet point with normal formatting
            content = text_stripped[9:].strip()  # Remove '<bullets>'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "• ", content)
            
        elif text_stripped.startswith('<n>'):
            # Remove the tag and add as numbered list with normal formatting
            content = text_stripped[3:].strip()  # Remove '<n>'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "1. ", content)
            
        elif text_stripped.startswith('<->'):
            # Remove the tag and add as dash list with normal formatting
            content = text_stripped[3:].strip()  # Remove '<->'
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, "- ", content)
            
        elif text_stripped.startswith('<note>'):
            # Remove the tag and add as note with red "Note:" prefix
            content = text_stripped[6:].strip()  # Remove '<note>'
            # Add "Note:" prefix in red
            note_prefix_run = paragraph.add_run("Note: ")
            DocumentHelper.set_font_style(note_prefix_run, font_name='Calibri', font_size=12, bold=True, color='#FF0000')  # Red color
            # Add content with markdown formatting support
            content_run = paragraph.add_run(content)
            DocumentHelper.set_font_style(content_run, font_name='Calibri', font_size=12, bold=True)
            
        elif text_stripped.startswith('<') and '>' in text_stripped:
            # Check for custom numbered/alphabet tags like <3>hello or <B>hello
            tag_end = text_stripped.find('>')
            if tag_end > 1:
                tag_content = text_stripped[1:tag_end]  # Get content between < and >
                # Check if it's a number (including multi-digit numbers)
                if tag_content.isdigit():
                    number = tag_content
                    content = text_stripped[tag_end + 1:].strip()  # Get text after >
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, f"{number}. ", content)
                # Check if it's a single alphabet character (A-Z, a-z)
                elif len(tag_content) == 1 and tag_content.isalpha():
                    alphabet = tag_content.upper()  # Convert to uppercase for consistency
                    content = text_stripped[tag_end + 1:].strip()  # Get text after >
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    DocumentHelper.add_formatted_text_with_normal_prefix(paragraph, f"{alphabet}. ", content)
                else:
                    # Not a number or single alphabet, treat as plain text
                    DocumentHelper.add_formatted_text_with_bold(paragraph, text)
            else:
                # Invalid tag format, treat as plain text
                DocumentHelper.add_formatted_text_with_bold(paragraph, text)
            
        else:
            # Plain text - use star-based bold formatting
            DocumentHelper.add_formatted_text_with_bold(paragraph, text)
        
        return paragraph

class EnhancedDocumentGenerator:
    """Enhanced document generator for security assessment reports"""
    
    @staticmethod
    def format_text_for_html(text):
        """Format text with proper HTML formatting for bullets, numbers, and alphabets"""
        if not text:
            return ""
        
        # Split text into lines
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append('<br>')
                continue
            
            # Helper function to apply bold formatting to content
            def apply_bold_formatting(content):
                import re
                return re.sub(r'\*([^*]+)\*', r'<strong>\1</strong>', content)
            
            # Check if the entire line is wrapped in bold markers
            if line.startswith('*') and line.endswith('*') and line.count('*') >= 2:
                # Extract content between the outer asterisks
                inner_content = line[1:-1].strip()
                
                # Process the inner content for formatting tags
                if inner_content.startswith('<bullets>'):
                    content = inner_content[9:].strip()  # Remove '<bullets>' prefix
                    formatted_content = apply_bold_formatting(content)
                    formatted_lines.append(f'<div style="margin-left: 20px;">• {formatted_content}</div>')
                
                elif inner_content.startswith('<n>'):
                    content = inner_content[3:].strip()  # Remove '<n>' prefix
                    formatted_content = apply_bold_formatting(content)
                    formatted_lines.append(f'<div style="margin-left: 20px;">1. {formatted_content}</div>')
                
                elif inner_content.startswith('<->'):
                    content = inner_content[3:].strip()  # Remove '<->' prefix
                    formatted_content = apply_bold_formatting(content)
                    formatted_lines.append(f'<div style="margin-left: 20px;">- {formatted_content}</div>')
                
                elif inner_content.startswith('<') and '>' in inner_content:
                    tag_end = inner_content.find('>')
                    if tag_end > 1:
                        tag_content = inner_content[1:tag_end]
                        content = inner_content[tag_end + 1:].strip()
                        
                        # Check if it's a number
                        if tag_content.isdigit():
                            formatted_content = apply_bold_formatting(content)
                            formatted_lines.append(f'<div style="margin-left: 20px;">{tag_content}. {formatted_content}</div>')
                        # Check if it's a single alphabet character
                        elif len(tag_content) == 1 and tag_content.isalpha():
                            alphabet = tag_content.upper()
                            formatted_content = apply_bold_formatting(content)
                            formatted_lines.append(f'<div style="margin-left: 20px;">{alphabet}. {formatted_content}</div>')
                        else:
                            # Not a recognized tag, treat as plain text with bold formatting
                            formatted_content = apply_bold_formatting(inner_content)
                            formatted_lines.append(f'<div>{formatted_content}</div>')
                    else:
                        formatted_content = apply_bold_formatting(inner_content)
                        formatted_lines.append(f'<div>{formatted_content}</div>')
                
                elif inner_content.startswith('-'):
                    content = inner_content[1:].strip()
                    formatted_content = apply_bold_formatting(content)
                    formatted_lines.append(f'<div style="margin-left: 20px;">• {formatted_content}</div>')
                
                else:
                    # No formatting tags, just apply bold formatting
                    formatted_content = apply_bold_formatting(inner_content)
                    formatted_lines.append(f'<div>{formatted_content}</div>')
            
            # Handle bullet points (not wrapped in bold)
            elif line.startswith('<bullets>'):
                content = line[9:].strip()  # Remove '<bullets>' prefix
                formatted_content = apply_bold_formatting(content)
                formatted_lines.append(f'<div style="margin-left: 20px;">• {formatted_content}</div>')
            
            # Handle numbered items with <n> tag (not wrapped in bold)
            elif line.startswith('<n>'):
                content = line[3:].strip()  # Remove '<n>' prefix
                formatted_content = apply_bold_formatting(content)
                formatted_lines.append(f'<div style="margin-left: 20px;">1. {formatted_content}</div>')
            
            # Handle dash items with <-> tag (not wrapped in bold)
            elif line.startswith('<->'):
                content = line[3:].strip()  # Remove '<->' prefix
                formatted_content = apply_bold_formatting(content)
                formatted_lines.append(f'<div style="margin-left: 20px;">- {formatted_content}</div>')
            
            # Handle numbered items like <1>, <2>, etc. (not wrapped in bold)
            elif line.startswith('<') and '>' in line:
                tag_end = line.find('>')
                if tag_end > 1:
                    tag_content = line[1:tag_end]
                    content = line[tag_end + 1:].strip()
                    
                    # Check if it's a number (including multi-digit numbers)
                    if tag_content.isdigit():
                        formatted_content = apply_bold_formatting(content)
                        formatted_lines.append(f'<div style="margin-left: 20px;">{tag_content}. {formatted_content}</div>')
                    # Check if it's a single alphabet character
                    elif len(tag_content) == 1 and tag_content.isalpha():
                        alphabet = tag_content.upper()
                        formatted_content = apply_bold_formatting(content)
                        formatted_lines.append(f'<div style="margin-left: 20px;">{alphabet}. {formatted_content}</div>')
                    else:
                        # Not a recognized tag, treat as plain text with bold formatting
                        formatted_content = apply_bold_formatting(line)
                        formatted_lines.append(f'<div>{formatted_content}</div>')
                else:
                    formatted_content = apply_bold_formatting(line)
                    formatted_lines.append(f'<div>{formatted_content}</div>')
            
            # Handle lines starting with dashes (not wrapped in bold)
            elif line.startswith('-'):
                content = line[1:].strip()
                formatted_content = apply_bold_formatting(content)
                formatted_lines.append(f'<div style="margin-left: 20px;">• {formatted_content}</div>')
            
            # Handle lines starting with numbers and dots
            elif line and line[0].isdigit() and '.' in line:
                formatted_content = apply_bold_formatting(line)
                formatted_lines.append(f'<div style="margin-left: 20px;">{formatted_content}</div>')
            
            # Handle lines starting with alphabets and dots
            elif line and len(line) > 1 and line[0].isalpha() and line[1] == '.':
                formatted_content = apply_bold_formatting(line)
                formatted_lines.append(f'<div style="margin-left: 20px;">{formatted_content}</div>')
            
            # Handle bold text with asterisks (plain text)
            elif '*' in line:
                formatted_content = apply_bold_formatting(line)
                formatted_lines.append(f'<div>{formatted_content}</div>')
            
            # Plain text
            else:
                formatted_lines.append(f'<div>{line}</div>')
        
        return ''.join(formatted_lines)
    
    def __init__(self):
        self.doc = None
        self.section_counter = 0  # Counter for dynamic section numbering
    
    def _extract_image_path(self, image_info):
        """Helper function to extract file path from image info dictionary or string"""
        if isinstance(image_info, dict):
            return image_info.get('absolute_path') or image_info.get('path', '')
        else:
            return str(image_info)
    
    def _get_image_filename(self, image_info, fallback_path=''):
        """Helper function to get filename from image info dictionary or path"""
        if isinstance(image_info, dict):
            return image_info.get('filename', os.path.basename(fallback_path))
        else:
            return os.path.basename(str(image_info))
    
    def _add_centered_image(self, image_path, width_cm=15):
        """Helper function to add a centered image with specified width in cm"""
        if not image_path or not os.path.exists(image_path):
            return False
            
        try:
            # Convert cm to inches (1 cm = 0.393701 inches)
            width_inches = width_cm * 0.393701
            
            # Add the image
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            
            return True
        except Exception as e:
            print(f"Error adding centered image {image_path}: {e}")
            return False
        
    def generate_docx(self, data, output_path):
        """Generate complete DOCX document following the specified structure"""
        self.doc = Document()
        
        # Reset section counter for dynamic numbering
        self.reset_section_counter()
        
        # Set up document styles
        self.setup_document_styles()
        
        # Generate document following the exact structure
        self.add_header(data)
        self.add_section_1_itsar(data)
        self.add_section_2_security_requirement(data)
        self.add_section_3_requirement_description(data)
        self.add_section_4_dut_confirmation(data)
        self.add_section_5_dut_configuration(data)
        self.add_section_6_preconditions(data)
        self.add_section_7_test_objective(data)
        self.add_section_8_test_plan(data)
        self.add_section_9_expected_result(data)
        self.add_section_10_evidence_format(data)
        self.add_section_11_test_execution(data)
        self.add_section_12_test_case_result(data)
        self.add_footer()
        
        # Save document with error handling
        try:
            self.doc.save(output_path)
        except PermissionError as e:
            print(f"Permission denied when saving document: {e}")
            print("The file might be open in another application (like Microsoft Word)")
            print("Please close the file and try again")
            raise e
        except Exception as e:
            print(f"Error saving DOCX: {e}")
            raise e
        
    def setup_document_styles(self):
        """Set up document-wide styles"""
        if not self.doc:
            return
            
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        
        # Set up sections
        if len(self.doc.sections) > 0:
            section = self.doc.sections[0]
            section.page_height = Inches(11.69)  # A4 height
            section.page_width = Inches(8.27)    # A4 width
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
    
    def add_styled_heading(self, text, is_main_header=True):
        """Add a heading with custom styling based on whether it's a main header or subheading"""
        if not self.doc:
            return None
            
        # Create paragraph for the heading
        heading_para = self.doc.add_paragraph(text)
        
        # Apply styling to the run
        for run in heading_para.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
            run.font.bold = True
            
            # Set font size based on header type
            if is_main_header:
                # Main headers: keep current font size (typically 16-18pt)
                run.font.size = Pt(14)
            else:
                # Subheadings: keep current font size (typically 14pt)
                run.font.size = Pt(12)
            
            # Set East Asian fonts
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Add spacing after heading
        heading_para.paragraph_format.space_after = Pt(6)
        
        return heading_para
    
    def get_next_section_number(self):
        """Get the next section number for dynamic numbering"""
        self.section_counter += 1
        return self.section_counter
    
    def reset_section_counter(self):
        """Reset the section counter (called at start of document generation)"""
        self.section_counter = 0
        self.subsection_counter = 0  # Counter for dynamic sub-numbering
    
    def add_numbered_section_heading(self, heading_text, is_main_header=True):
        """Add a section heading with automatic numbering for top-level sections"""
        if not self.doc:
            return None
        
        # Get next section number
        section_num = self.get_next_section_number()
        
        # Reset subsection counter when starting a new section
        self.subsection_counter = 0
        
        # Remove any existing number from the heading text
        # This handles cases where the UI still has hardcoded numbers
        import re
        heading_text = re.sub(r'^\d+\.\s*', '', heading_text)
        
        # Create the numbered heading
        numbered_heading = f"{section_num}. {heading_text}"
        
        # Use the existing styled heading method
        return self.add_styled_heading(numbered_heading, is_main_header)
    
    def add_numbered_subheading(self, heading_text, is_main_header=False):
        """Add a subheading with automatic sub-numbering based on current section"""
        if not self.doc:
            return None
        
        # Get current section number and increment subsection counter
        if self.section_counter > 0:
            section_num = self.section_counter
        else:
            # If no section has been created yet, start with section 1
            section_num = 1
        
        self.subsection_counter += 1
        
        # Remove any existing number from the heading text
        # This handles cases where the UI still has hardcoded numbers
        import re
        heading_text = re.sub(r'^\d+\.\d+\.\s*', '', heading_text)
        
        # Create the sub-numbered heading
        sub_numbered_heading = f"{section_num}.{self.subsection_counter}. {heading_text}"
        
        # Use the existing styled heading method
        return self.add_styled_heading(sub_numbered_heading, is_main_header)
            
    def add_header(self, data):
        """Add header with test report title"""
        if not self.doc:
            return
            
        # Header
            
        # Main title
        test_report = data.get('test_report', 'TEST REPORT FOR: [Test Case Title]')
        # Create title as a paragraph instead of heading to avoid default heading styles
        title = self.doc.add_paragraph(test_report)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply custom formatting to the title
        for run in title.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black color
            # Set the font for East Asian characters as well
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        title.paragraph_format.space_after = Pt(6)  # Reduced spacing
        
        # Add DUT Details after title
        self.add_dut_details(data)
        
    def add_dut_details(self, data):
        """Add DUT Details after the title"""
        if not self.doc:
            return
            
        # Add DUT fields
        dut_fields = data.get('dut_fields', [])
        if dut_fields:
            for field in dut_fields:
                label = field.get('label', '')
                value = field.get('value', '')
                if label and value:
                    # Create paragraph with label and value on same line
                    para = self.doc.add_paragraph()
                    label_run = para.add_run(f"{label}: ")
                    label_run.font.bold = True
                    label_run.font.name = 'Calibri'
                    label_run.font.size = Pt(12)
                    label_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    
                    value_run = para.add_run(value)
                    value_run.font.name = 'Calibri'
                    value_run.font.size = Pt(12)
                    value_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    
                    para.paragraph_format.space_after = Pt(6)  # Space after the entire field

                # Add DUT field images if available
                dut_images = field.get('images', [])
                if dut_images:
                    for image_info in dut_images:
                        image_path = self._extract_image_path(image_info)
                        
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                filename = self._get_image_filename(image_info, image_path)
                                self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                
                # Add DUT field scripts if available
                dut_scripts = field.get('scripts', [])
                if dut_scripts:
                    for script_data in dut_scripts:
                        description = script_data.get('description', '')
                        filename = script_data.get('filename', '')
                        
                        if description:
                            para = self.doc.add_paragraph()
                            run = para.add_run(description)
                            DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        if filename:
                            display_filename = script_data.get('original_filename', filename)
                            self.doc.add_paragraph("")  # Add spacing
        else:
            # Add placeholder if no DUT fields
            para = self.doc.add_paragraph("[DUT Details to be added]")
            para.runs[0].font.name = 'Calibri'
            para.runs[0].font.size = Pt(12)
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Add DUT placeholders below DUT details
        self.add_dut_placeholders(data)
        
        # Add hash sections
        self.add_hash_sections(data)
        
        # Add ITSAR information
        self.add_itsar_information(data)
        
        # Add spacing
        # self.doc.add_paragraph()
    
    def add_section_placeholders(self, data, section_key):
        """Add placeholders for any section"""
        if not self.doc:
            return
            
        # Get placeholders from scenario_placeholders
        scenario_placeholders = data.get('scenario_placeholders', {})
        section_placeholders = scenario_placeholders.get(section_key, [])
        
        if section_placeholders:
            # Add each placeholder (centered) without heading
            for placeholder_info in section_placeholders:
                if isinstance(placeholder_info, dict):
                    placeholder_text = placeholder_info.get('placeholder', '')
                    if placeholder_text:
                        # Create paragraph for placeholder (centered)
                        para = self.doc.add_paragraph()
                        run = para.add_run(placeholder_text)
                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=11, bold=False)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.paragraph_format.space_after = Pt(3)
                elif isinstance(placeholder_info, str):
                    # Handle string placeholders (centered)
                    para = self.doc.add_paragraph()
                    run = para.add_run(placeholder_info)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=11, bold=False)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_after = Pt(3)
        
    def add_dut_placeholders(self, data):
        """Add DUT placeholders below DUT details"""
        self.add_section_placeholders(data, 'DUT Details')
    
    def add_hash_sections(self, data):
        """Add hash sections after DUT details and placeholders"""
        if not self.doc:
            return
            
        # Multiple Hash Sections
        hash_sections = data.get('hash_sections', [])
        
        for section in hash_sections:
            heading = section.get('heading', '')
            direct_hash_value = section.get('direct_hash_value', '')
            hash_fields = section.get('hash_fields', [])
            
            # Display hash section if there's a heading or any content
            if heading or direct_hash_value or hash_fields:
                # Add hash section heading if provided
                if heading:
                    hash_heading = self.doc.add_paragraph()
                    # Add colon to heading if it doesn't already end with one
                    formatted_heading = heading.strip()
                    if not formatted_heading.endswith(':'):
                        formatted_heading += ':'
                    hash_heading_run = hash_heading.add_run(formatted_heading)
                    DocumentHelper.set_font_style(hash_heading_run, font_name='Calibri', font_size=12, bold=True)
                    hash_heading.paragraph_format.space_after = Pt(6)  # Consistent spacing
                
                # Add direct hash value if provided
                if direct_hash_value:
                    direct_hash_para = self.doc.add_paragraph()
                    direct_hash_run = direct_hash_para.add_run(direct_hash_value)
                    DocumentHelper.set_font_style(direct_hash_run, font_name='Calibri', font_size=12, bold=False)
                    direct_hash_para.paragraph_format.space_after = Pt(1)
                
                # Add direct hash images if available
                direct_hash_images = section.get('direct_hash_images', [])
                if direct_hash_images:
                    for image_info in direct_hash_images:
                        image_path = self._extract_image_path(image_info)
                        
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                print(f"Error adding direct hash image {image_path}")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                
                # Add direct hash scripts if available
                direct_hash_scripts = section.get('direct_hash_scripts', [])
                if direct_hash_scripts:
                    for script_info in direct_hash_scripts:
                        if isinstance(script_info, dict):
                            # New format with full script info
                            script_description = script_info.get('description', '')
                            script_filename = script_info.get('original_filename', script_info.get('filename', ''))
                            
                            if script_description:
                                script_desc_para = self.doc.add_paragraph()
                                script_desc_run = script_desc_para.add_run(script_description)
                                DocumentHelper.set_font_style(script_desc_run, font_name='Calibri', font_size=12, bold=False)
                                script_desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                script_desc_para.paragraph_format.space_after = Pt(1)
                        else:
                            # Old format with just path - skip displaying script filename
                            pass
                
                # Add individual hash fields
                for field in hash_fields:
                    field_heading = field.get('label', '')
                    field_value = field.get('value', '')
                    if field_heading:  # Show field if it has a heading, even if value is empty
                        # Add hash subheading
                        hash_subheading = self.doc.add_paragraph()
                        # Add colon to field heading if it doesn't already end with one
                        formatted_field_heading = field_heading.strip()
                        if not formatted_field_heading.endswith(':'):
                            formatted_field_heading += ':'
                        hash_subheading_run = hash_subheading.add_run(formatted_field_heading)
                        DocumentHelper.set_font_style(hash_subheading_run, font_name='Calibri', font_size=12, bold=True)
                        hash_subheading.paragraph_format.space_after = Pt(6)  # Consistent spacing
                    
                        # Add hash value (even if empty)
                        hash_value_para = self.doc.add_paragraph()
                        if field_value:
                            hash_value_run = hash_value_para.add_run(field_value)
                        else:
                            hash_value_run = hash_value_para.add_run("[Hash value to be added]")
                        DocumentHelper.set_font_style(hash_value_run, font_name='Calibri', font_size=12, bold=False)
                        hash_value_para.paragraph_format.space_after = Pt(1)
                        
                        # Add hash field images if available
                        field_images = field.get('images', [])
                        if field_images:
                            for image_info in field_images:
                                image_path = self._extract_image_path(image_info)
                                
                                if image_path and os.path.exists(image_path):
                                    if not self._add_centered_image(image_path):
                                        print(f"Error adding hash field image {image_path}")
                                else:
                                    filename = self._get_image_filename(image_info, image_path)
                                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                        
                        # Add hash field scripts if available
                        field_scripts = field.get('scripts', [])
                        if field_scripts:
                            for script_info in field_scripts:
                                if isinstance(script_info, dict):
                                    # New format with full script info
                                    script_description = script_info.get('description', '')
                                    script_filename = script_info.get('original_filename', script_info.get('filename', ''))
                                    
                                    if script_description:
                                        script_desc_para = self.doc.add_paragraph()
                                        script_desc_run = script_desc_para.add_run(script_description)
                                        DocumentHelper.set_font_style(script_desc_run, font_name='Calibri', font_size=12, bold=False)
                                        script_desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        script_desc_para.paragraph_format.space_after = Pt(1)
                                else:
                                    # Old format with just path - skip displaying script filename
                                    pass
                        
                        hash_value_para.paragraph_format.space_after = Pt(2)
        
        # If no hash sections, add placeholder
        if not hash_sections:
            para = self.doc.add_paragraph("[Hash Sections to be added]")
            para.runs[0].font.name = 'Calibri'
            para.runs[0].font.size = Pt(12)
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Add Hash Sections placeholders
        self.add_section_placeholders(data, 'Hash Sections')
        
        # Add spacing
        self.doc.add_paragraph()
    
    def add_itsar_information(self, data):
        """Add ITSAR Information after hash sections"""
        if not self.doc:
            return
            
        # ITSAR Information - Fixed formatting
        itsar_fields = data.get('itsar_fields', [])
        if itsar_fields:
            for field in itsar_fields:
                key = field.get('key', '')
                value = field.get('value', '')
                if key:  # Show field if it has a key, even if value is empty
                    # Check if value contains newlines (multi-line input)
                    if value and '\n' in value:
                        # Multi-line input - first show the heading, then bullet points
                        # Add heading first
                        heading_para = self.doc.add_paragraph()
                        heading_run = heading_para.add_run(f"{key}:")
                        heading_run.font.bold = True
                        heading_run.font.name = 'Calibri'
                        heading_run.font.size = Pt(12)
                        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                        heading_para.paragraph_format.space_after = Pt(6)  # Space after heading
                        
                        # Then add each line as numbered item (like DUT details format)
                        lines = value.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                # Add each line as a regular paragraph (not bullet point)
                                para = self.doc.add_paragraph()
                                line_run = para.add_run(line)
                                line_run.font.name = 'Calibri'
                                line_run.font.size = Pt(12)
                                line_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                                para.paragraph_format.space_after = Pt(6)  # Same spacing as DUT details
                    else:
                        # Single-line input or empty value - output in same single line
                        # Create paragraph with key and value on same line
                        para = self.doc.add_paragraph()
                        key_run = para.add_run(f"{key}: ")
                        key_run.font.bold = True
                        key_run.font.name = 'Calibri'
                        key_run.font.size = Pt(12)
                        
                        # Use value if available, otherwise show placeholder or empty
                        display_value = value if value else ""
                        value_run = para.add_run(display_value)
                        value_run.font.name = 'Calibri'
                        value_run.font.size = Pt(12)
                        
                        # Set East Asian fonts
                        key_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                        
                        para.paragraph_format.space_after = Pt(6)  # Same spacing as DUT details
                
                # Add ITSAR field images if available
                itsar_images = field.get('images', [])
                if itsar_images:
                    for image_info in itsar_images:
                        image_path = self._extract_image_path(image_info)
                        
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                filename = self._get_image_filename(image_info, image_path)
                                self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                
                # Add ITSAR field scripts if available
                itsar_scripts = field.get('scripts', [])
                if itsar_scripts:
                    for script_data in itsar_scripts:
                        description = script_data.get('description', '')
                        filename = script_data.get('filename', '')
                        
                        if description:
                            para = self.doc.add_paragraph()
                            run = para.add_run(description)
                            DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        if filename:
                            display_filename = script_data.get('original_filename', filename)
                            self.doc.add_paragraph("")  # Add spacing
        
        # Add ITSAR Information placeholders
        self.add_section_placeholders(data, 'ITSAR Information')
        
        # Add spacing
        self.doc.add_paragraph()
        
    def add_section_1_itsar(self, data):
        """Section 1: ITSAR Section No & Name"""
        if not self.doc:
            return
            
        # Add page break before Section 1
        self.doc.add_page_break()
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        print(f"🔍 DEBUG Section 1: sections_1_7 = {sections_1_7}")
        section_1_data = sections_1_7.get('section_1', {})
        print(f"🔍 DEBUG Section 1: section_1_data = {section_1_data}")
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_1_data.get('heading', 'ITSAR Section No & Name')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Get ITSAR section from sections 1-7 data or fallback to old data
        itsar_section = section_1_data.get('content', data.get('itsar_section', ''))
        if itsar_section:
            DocumentHelper.add_smart_subheading_paragraph(self.doc, itsar_section)
        else:
            self.doc.add_paragraph("[User input text]")
        
        # Add placeholder content if available
        placeholder_content_management = section_1_data.get('placeholder_content_management', [])
        print(f"🔍 DEBUG Section 1: placeholder_content_management = {placeholder_content_management}")
        self.add_placeholder_content(placeholder_content_management)
            
        # Add section 1 images if available
        section_1_images = section_1_data.get('images', [])
        if section_1_images:
            for image_info in section_1_images:
                image_path = self._extract_image_path(image_info)
                
                if image_path and os.path.exists(image_path):
                    if not self._add_centered_image(image_path):
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add section 1 scripts if available
        section_1_scripts = section_1_data.get('scripts', [])
        if section_1_scripts:
            for script_data in section_1_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    # Add description as centered placeholder text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(description)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    self.doc.add_paragraph("")  # Add spacing
            
            # Add Section 1 placeholders
            self.add_section_placeholders(data, 'ITSAR Section No & Name')
        
        # Add section 1 notes if available
        section_1_notes = section_1_data.get('notes', [])
        print(f"🔍 DEBUG Section 1: section_1_notes = {section_1_notes}")
        if section_1_notes:
            print(f"🔍 DEBUG Section 1: Calling add_section_notes with {len(section_1_notes)} notes")
            self.add_section_notes(section_1_notes)
            
    def add_section_2_security_requirement(self, data):
        """Section 2: Security Requirement"""
        if not self.doc:
            return
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        section_2_data = sections_1_7.get('section_2', {})
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_2_data.get('heading', 'Security Requirement No & Name')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Get security requirement from sections 1-7 data or fallback to old data
        security_req = section_2_data.get('content', data.get('security_req', ''))
        if security_req:
            self.doc.add_paragraph(security_req)
        else:
            self.doc.add_paragraph("[User input]")
        
        # Add placeholder content if available
        placeholder_content_management = section_2_data.get('placeholder_content_management', [])
        self.add_placeholder_content(placeholder_content_management)
            
        # Add section 2 images if available
        section_2_images = section_2_data.get('images', [])
        if section_2_images:
            for image_info in section_2_images:
                image_path = self._extract_image_path(image_info)
                
                if image_path and os.path.exists(image_path):
                    if not self._add_centered_image(image_path):
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add section 2 scripts if available
        section_2_scripts = section_2_data.get('scripts', [])
        if section_2_scripts:
            for script_data in section_2_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    # Add description as centered placeholder text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(description)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    self.doc.add_paragraph("")  # Add spacing
            
            # Add Section 2 placeholders
            self.add_section_placeholders(data, 'Security Requirement No & Name')
        
        # Add section 2 notes if available
        section_2_notes = section_2_data.get('notes', [])
        if section_2_notes:
            self.add_section_notes(section_2_notes)
            
    def add_section_3_requirement_description(self, data):
        """Section 3: Requirement Description"""
        if not self.doc:
            return
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        section_3_data = sections_1_7.get('section_3', {})
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_3_data.get('heading', 'Requirement Description')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Get requirement description from sections 1-7 data or fallback to old data
        req_description = section_3_data.get('content', data.get('req_description', ''))
        if req_description:
            self.doc.add_paragraph(req_description)
        else:
            self.doc.add_paragraph("[Requirement description to be added]")
        
        # Add placeholder content if available
        placeholder_content_management = section_3_data.get('placeholder_content_management', [])
        self.add_placeholder_content(placeholder_content_management)
            
        # Add section 3 images if available
        section_3_images = section_3_data.get('images', [])
        if section_3_images:
            for image_info in section_3_images:
                image_path = self._extract_image_path(image_info)
                
                if image_path and os.path.exists(image_path):
                    if not self._add_centered_image(image_path):
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add section 3 scripts if available
        section_3_scripts = section_3_data.get('scripts', [])
        if section_3_scripts:
            for script_data in section_3_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    # Add description as centered placeholder text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(description)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    self.doc.add_paragraph("")  # Add spacing
            
            # Add Section 3 placeholders
            self.add_section_placeholders(data, 'Requirement Description')
        
        # Add section 3 notes if available
        section_3_notes = section_3_data.get('notes', [])
        if section_3_notes:
            self.add_section_notes(section_3_notes)
            
    def add_section_4_dut_confirmation(self, data):
        """Section 4: DUT Confirmation Details with DUT Images"""
        if not self.doc:
            return
            
        # Get sections 1-7 data for proper mapping
        sections_1_7 = data.get('sections_1_7', {})
        print(f"🔍 DEBUG Section 4: sections_1_7 = {sections_1_7}")
        section_4_data = sections_1_7.get('section_4', {})
        print(f"🔍 DEBUG Section 4: section_4_data = {section_4_data}")
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_4_data.get('heading', 'DUT Confirmation Details')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Section 4 content is now handled in the pairs processing below
        
        # DUT Images subsection
        # self.doc.add_heading("DUT Images", 2)
        
        # Add screenshots if available
        screenshots = data.get('screenshots', [])
        if screenshots:
            for i, screenshot_path in enumerate(screenshots[:4]):  # Limit to 4 images
                if os.path.exists(screenshot_path):
                    if not self._add_centered_image(screenshot_path):
                        # Remove "Figure:" text - show only the image
                        self.doc.add_paragraph(f"[Image could not be loaded: {os.path.basename(screenshot_path)}]")
                else:
                    self.doc.add_paragraph(f"[Image placeholder {i+1}]")
        # DUT Version info section removed as requested
        
        # Connection details
        machine_ip = data.get('machine_ip', '')
        target_ip = data.get('target_ip', '')
        if machine_ip and target_ip:
            connection_text = f"DUT is connected to a testing machine (IP address {machine_ip}) via Ethernet LAN1 Connection (IP address {target_ip})."
            DocumentHelper.add_smart_subheading_paragraph(self.doc, connection_text)
            
        # Add section 4 text-image pairs if available (using smart subheading formatting)
        section_4_pairs = section_4_data.get('pairs', [])
        if section_4_pairs:
            for pair in section_4_pairs:
                text_content = pair.get('text', '')
                images = pair.get('images', [])
                scripts = pair.get('scripts', [])
                
                if text_content:
                    # Split text content into lines and process each line separately
                    lines = text_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:  # Only process non-empty lines
                            paragraph = DocumentHelper.add_smart_subheading_paragraph(self.doc, line)
                
                # Add pair-specific placeholders right after the text content
                pair_placeholders = pair.get('placeholders', [])
                for placeholder_info in pair_placeholders:
                    placeholder_text = placeholder_info.get('placeholder', '').strip()
                    if placeholder_text:
                        # Add placeholder text with indentation (no bullet, just indented)
                        placeholder_paragraph = self.doc.add_paragraph()
                        placeholder_paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent to align with bullet text
                        placeholder_run = placeholder_paragraph.add_run(placeholder_text)
                        DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, bold=False, italic=True)
                
                # Add scripts if available (like Section 6)
                upload_scripts_list = pair.get('upload_scripts_list', [])
                if upload_scripts_list:
                    for script_data in upload_scripts_list:
                        filename = script_data.get('filename', '')
                        placeholder_text = script_data.get('placeholder', script_data.get('description', ''))
                        
                        print(f"🔍 DEBUG Section 4 Document: Processing script with placeholder: '{placeholder_text}'")
                        
                        if filename:
                            # Use original filename if available, otherwise use the stored filename
                            display_filename = script_data.get('original_filename', filename)
                            
                            # Add script placeholder text if available
                            if placeholder_text:
                                placeholder_paragraph = self.doc.add_paragraph()
                                placeholder_paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent to align with bullet text
                                placeholder_run = placeholder_paragraph.add_run(placeholder_text)
                                DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                print(f"🔍 DEBUG Section 4 Document: Added placeholder text: '{placeholder_text}'")
                            
                            # self.doc.add_paragraph(f"Script file: {display_filename}")
                            self.doc.add_paragraph("")  # Add spacing
                
                images_list = pair.get('images_list', [])
                for image_info in images_list:
                    image_path = self._extract_image_path(image_info)
                    
                    if image_path and os.path.exists(image_path):
                        if not self._add_centered_image(image_path):
                            # Remove "Figure:" text - show only the image
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                    else:
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add interfaces description if available
        interfaces_desc = section_4_data.get('interfaces_desc', '')
        if interfaces_desc:
            # Process the interfaces description with smart formatting
            DocumentHelper.add_smart_subheading_paragraph(self.doc, interfaces_desc)
        
        # Placeholder content is now displayed inline with pairs above
        
        # Add section 4 notes if available
        section_4_notes = section_4_data.get('notes', [])
        if section_4_notes:
            self.add_section_notes(section_4_notes)
        
        # Interface status - MOVED TO END OF SECTION 4
        
        # Use Section 4 interface data if available, otherwise fall back to global interfaces
        section_4_interfaces = section_4_data.get('interfaces', [])
        print(f"🔍 DEBUG Section 4: section_4_interfaces = {section_4_interfaces}")
        global_interfaces = data.get('interfaces', [])
        print(f"🔍 DEBUG Section 4: global_interfaces = {global_interfaces}")
        if section_4_interfaces:
            print(f"🔍 DEBUG Section 4: Using section_4_interfaces")
            self.add_interface_table_from_section_4(section_4_interfaces)
        elif global_interfaces:
            print(f"🔍 DEBUG Section 4: Using global_interfaces")
            # Only show interface table if user has provided interfaces
            self.add_interface_table(global_interfaces)
        else:
            print(f"🔍 DEBUG Section 4: No interfaces found")
            
    def add_interface_table(self, interfaces):
        """Add interface status table"""
        if not self.doc:
            return
            
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row with light blue background
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Interfaces'
        header_cells[1].text = 'No. of Ports'
        header_cells[2].text = 'Interface Type'
        header_cells[3].text = 'Interface Name'
        
        # Apply light blue background color to header cells
        for cell in header_cells:
            # Set background color to light blue (RGB: 173, 216, 230)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'ADD8E6')  # Light blue hex color
            tcPr.append(shd)
            
            # Make text bold and center-aligned
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Data rows
        if interfaces:
            for interface in interfaces:
                if len(interface) >= 4:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(interface[:4]):
                        row_cells[i].text = str(value)
                    
    def add_interface_table_from_section_4(self, interfaces):
        """Add interface status table from Section 4 data (new JSON format)"""
        if not self.doc:
            return
            
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row with light blue background
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Interfaces'
        header_cells[1].text = 'No. of Ports'
        header_cells[2].text = 'Interface Type'
        header_cells[3].text = 'Interface Name'
        
        # Apply light blue background color to header cells
        for cell in header_cells:
            # Set background color to light blue (RGB: 173, 216, 230)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'ADD8E6')  # Light blue hex color
            tcPr.append(shd)
            
            # Make text bold and center-aligned
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Data rows - handle new JSON format
        if interfaces:
            for interface in interfaces:
                if isinstance(interface, dict):
                    # New JSON format: {"interface": "...", "ports": "...", "type": "...", "name": "..."}
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(interface.get('interface', ''))
                    row_cells[1].text = str(interface.get('ports', ''))
                    row_cells[2].text = str(interface.get('type', ''))
                    row_cells[3].text = str(interface.get('name', ''))
                elif isinstance(interface, list) and len(interface) >= 4:
                    # Old format: ["interface", "ports", "type", "name"]
                    row_cells = table.add_row().cells
                    for i, value in enumerate(interface[:4]):
                        row_cells[i].text = str(value)
        else:
            # Default interface data
            default_interfaces = [
                {"interface": "USB", "ports": "1 Port", "type": "Physical", "name": "USB Port"},
                {"interface": "Ethernet Interface", "ports": "2 Ports", "type": "Physical", "name": "eth0/PoE (LAN1) 2.5G, eth1 (LAN2) 1G"},
                {"interface": "Console", "ports": "1 Port", "type": "Physical", "name": "Console"},
                {"interface": "Wlan (WIFI)", "ports": "3 Ports", "type": "Wireless", "name": "wifi0, wifi1, wifi2"},
                {"interface": "Bridge Interface", "ports": "1 Port", "type": "Bridge Interface", "name": "br0"},
                {"interface": "Loopback", "ports": "1 Port", "type": "Logical", "name": "lo"}
            ]
            for interface in default_interfaces:
                row_cells = table.add_row().cells
                row_cells[0].text = str(interface.get('interface', ''))
                row_cells[1].text = str(interface.get('ports', ''))
                row_cells[2].text = str(interface.get('type', ''))
                row_cells[3].text = str(interface.get('name', ''))
                    
    def add_section_5_dut_configuration(self, data):
        """Section 5: DUT Configuration (Text + Image)"""
        if not self.doc:
            return
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        print(f"🔍 DEBUG Section 5: sections_1_7 = {sections_1_7}")
        section_5_data = sections_1_7.get('section_5', {})
        print(f"🔍 DEBUG Section 5: section_5_data = {section_5_data}")
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_5_data.get('heading', 'DUT Configuration')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Note: DUT fields data should not be added to Section 5
        # DUT Details from DUT configuration should only appear in their designated section
        # Section 5 should only contain content from sections_1_7 data
        
        # Section 5 content is now handled in the pairs processing below
            
        # Add section 5 text-image pairs if available using resolved objects
        section_5_pairs = section_5_data.get('pairs', [])
        if section_5_pairs:
            for pair in section_5_pairs:
                text_content = pair.get('text', '')
                resolved_images = pair.get('resolved_images', [])
                resolved_scripts = pair.get('resolved_scripts', [])

                if text_content:
                    # Split text content into lines and process each line separately
                    lines = text_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:  # Only process non-empty lines
                            paragraph = DocumentHelper.add_smart_subheading_paragraph(self.doc, line)

                # Add pair-specific placeholders right after the text content
                pair_placeholders = pair.get('placeholders', [])
                for placeholder_info in pair_placeholders:
                    placeholder_text = placeholder_info.get('placeholder', '').strip()
                    if placeholder_text:
                        # Add placeholder text with indentation (no bullet, just indented)
                        placeholder_paragraph = self.doc.add_paragraph()
                        placeholder_paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent to align with bullet text
                        placeholder_run = placeholder_paragraph.add_run(placeholder_text)
                        DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, bold=False, italic=True)

                # Add scripts if available (like Section 6)
                upload_scripts_list = pair.get('upload_scripts_list', [])
                if upload_scripts_list:
                    for script_info in upload_scripts_list:
                        if isinstance(script_info, dict):
                            filename = script_info.get('filename', '')
                            placeholder_text = script_info.get('placeholder', script_info.get('description', ''))
                            
                            print(f"🔍 DEBUG Section 5 Document: Processing script with placeholder: '{placeholder_text}'")

                            if filename:
                                # Use original filename if available, otherwise use the stored filename
                                display_filename = script_info.get('original_filename', filename)
                                
                                # Add script placeholder text if available
                                if placeholder_text:
                                    placeholder_paragraph = self.doc.add_paragraph()
                                    placeholder_paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent to align with bullet text
                                    placeholder_run = placeholder_paragraph.add_run(placeholder_text)
                                    DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                    print(f"🔍 DEBUG Section 5 Document: Added placeholder text: '{placeholder_text}'")
                                
                                # self.doc.add_paragraph(f"Script file: {display_filename}")
                                self.doc.add_paragraph("")  # Add spacing

                # Add images using images_list
                images_list = pair.get('images_list', [])
                for image_info in images_list:
                    image_path = self._extract_image_path(image_info)
                    
                    if image_path and os.path.exists(image_path):
                        if not self._add_centered_image(image_path):
                            # Remove "Figure:" text - show only the image
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                    else:
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        else:
            # Configuration image placeholder
            self.doc.add_paragraph("[Configuration image to be inserted]")
        
        # Placeholder content is now displayed inline with pairs above
        
        # Add section 5 notes if available
        section_5_notes = section_5_data.get('notes', [])
        if section_5_notes:
            self.add_section_notes(section_5_notes)
            
    def add_section_6_preconditions(self, data):
        """Section 6: Preconditions"""
        if not self.doc:
            return
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        section_6_data = sections_1_7.get('section_6', {})
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_6_data.get('heading', 'Preconditions')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Preconditions from sections 1-7 data (prioritize user input from UI)
        preconditions = section_6_data.get('content', '')
        if preconditions:
            # Split by newlines to handle multiple items
            lines = preconditions.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Use smart formatting for each line
                    paragraph = DocumentHelper.add_smart_subheading_paragraph(self.doc, line)
        
        # Add placeholder content if available
        placeholder_content_management = section_6_data.get('placeholder_content_management', [])
        self.add_placeholder_content(placeholder_content_management)
            
        # Add section 6 images if available
        section_6_images = section_6_data.get('images', [])
        if section_6_images:
            for image_info in section_6_images:
                image_path = self._extract_image_path(image_info)
                
                if image_path and os.path.exists(image_path):
                    if not self._add_centered_image(image_path):
                        # Remove "Figure:" text - show only the image
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add section 6 scripts if available
        section_6_scripts = section_6_data.get('scripts', [])
        if section_6_scripts:
            for script_data in section_6_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    # Add description as centered placeholder text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(description)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    # self.doc.add_paragraph(f"Script file: {display_filename}")
                    self.doc.add_paragraph("")  # Add spacing
            
            # Add Section 6 placeholders
            self.add_section_placeholders(data, 'Preconditions')
        
        # Add section 6 notes if available
        section_6_notes = section_6_data.get('notes', [])
        if section_6_notes:
            self.add_section_notes(section_6_notes)
            
    def add_section_7_test_objective(self, data):
        """Section 7: Test Objective"""
        if not self.doc:
            return
            
        # Get sections 1-7 data
        sections_1_7 = data.get('sections_1_7', {})
        section_7_data = sections_1_7.get('section_7', {})
        
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        section_heading = section_7_data.get('heading', 'Test Objective')
        self.add_numbered_section_heading(section_heading, is_main_header=True)
        
        # Test objective from sections 1-7 data or fallback to old data
        test_objective = section_7_data.get('content', data.get('test_objective', ''))
        if test_objective:
            self.doc.add_paragraph(test_objective)
        
        # Add placeholder content if available
        placeholder_content_management = section_7_data.get('placeholder_content_management', [])
        self.add_placeholder_content(placeholder_content_management)
            
        # Add section 7 images if available
        section_7_images = section_7_data.get('images', [])
        if section_7_images:
            for image_info in section_7_images:
                image_path = self._extract_image_path(image_info)
                
                if image_path and os.path.exists(image_path):
                    if not self._add_centered_image(image_path):
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add section 7 scripts if available
        section_7_scripts = section_7_data.get('scripts', [])
        if section_7_scripts:
            for script_data in section_7_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    # Add description as centered placeholder text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(description)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    self.doc.add_paragraph("")  # Add spacing


            
        # Add section 7 text-image pairs if available
        section_7_pairs = section_7_data.get('pairs', [])
        if section_7_pairs:
            for pair in section_7_pairs:
                text_content = pair.get('text', '')
                images = pair.get('images', [])
                
                if text_content:
                    # Split text content into lines and process each line separately
                    lines = text_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:  # Only process non-empty lines
                            DocumentHelper.add_smart_subheading_paragraph(self.doc, line)
                
                for image_info in images:
                    image_path = self._extract_image_path(image_info)
                    
                    if image_path and os.path.exists(image_path):
                        if not self._add_centered_image(image_path):
                            # Remove "Figure:" text - show only the image
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                    else:
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image placeholder: {filename}]")
            
            # Add Section 7 placeholders
            self.add_section_placeholders(data, 'Test Objective')
        
        # Add section 7 notes if available
        section_7_notes = section_7_data.get('notes', [])
        if section_7_notes:
            self.add_section_notes(section_7_notes)
            
    def add_section_8_test_plan(self, data):
        """Section 8: Test Plan with subsections"""
        if not self.doc:
            return
            
        self.add_numbered_section_heading("Test Plan", is_main_header=True)
        
        # 8.0 Test Plan Overview (Internal Reference Only)
        test_plan_overview = data.get('test_plan_overview', '')
        if test_plan_overview:
            # self.doc.add_heading("8.0. Test Plan Overview", 2)
            self.doc.add_paragraph(test_plan_overview)
        
        # Section 1 - Custom editable section
        section1_heading = data.get('section1_heading', '')
        section1_content = data.get('section1_content', '')
        if section1_heading or section1_content:
            if section1_heading:
                self.add_numbered_subheading(section1_heading, is_main_header=False)
            if section1_content:
                self.doc.add_paragraph(section1_content)
        
        # 8.1 Number of Test Scenarios
        test_scenarios_heading = data.get('test_scenarios_heading', 'Number of Test Scenarios')
        self.add_numbered_subheading(test_scenarios_heading, is_main_header=False)
        scenarios = data.get('test_scenarios', [])
        if scenarios:
            for scenario in scenarios:
                scenario_key = scenario.get('key', '')
                scenario_description = scenario.get('description', '')
                if scenario_key and scenario_description:
                    # Add scenario key on first line with bold formatting
                    scenario_header = f"{scenario_key}:"
                    scenario_para = self.doc.add_paragraph()
                    scenario_run = scenario_para.add_run(scenario_header)
                    scenario_run.font.bold = True
                    # Process description with smart formatting
                    if scenario_description:
                        # Split text by lines to handle multiple paragraphs
                        lines = scenario_description.split('\n')
                        
                        for line in lines:
                            if not line.strip():
                                # Empty line - add paragraph break
                                self.doc.add_paragraph()
                                continue
                                
                            # Use smart formatting for each line
                            DocumentHelper.add_smart_subheading_paragraph(self.doc, line)
                
                # Handle script placeholders for test scenarios
                scenario_scripts = scenario.get('scripts', [])
                print(f"🔍 DEBUG 8.1: scenario_scripts = {scenario_scripts}")
                if scenario_scripts:
                    for script_info in scenario_scripts:
                        print(f"🔍 DEBUG 8.1: script_info = {script_info}")
                        if isinstance(script_info, dict):
                            script_description = script_info.get('description', '')
                            print(f"🔍 DEBUG 8.1: script_description = '{script_description}'")
                            if script_description:
                                print(f"🔍 DEBUG 8.1: Adding placeholder '{script_description}' to document")
                                # Add description as centered placeholder text
                                paragraph = self.doc.add_paragraph()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run(script_description)
                                DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                paragraph.paragraph_format.space_after = Pt(6)
                
                # Also check for placeholders in scenario_placeholders data structure
                scenario_placeholders = scenario.get('placeholders', [])
                print(f"🔍 DEBUG 8.1: scenario_placeholders = {scenario_placeholders}")
                if scenario_placeholders:
                    for placeholder_text in scenario_placeholders:
                        print(f"🔍 DEBUG 8.1: Adding placeholder '{placeholder_text}' to document")
                        # Add placeholder as centered text
                        paragraph = self.doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(placeholder_text)
                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                        paragraph.paragraph_format.space_after = Pt(6)
                
                # Also check for images in test scenarios
                scenario_images = scenario.get('images', [])
                print(f"🔍 DEBUG 8.1: scenario_images = {scenario_images}")
                if scenario_images:
                    for image_info in scenario_images:
                        print(f"🔍 DEBUG 8.1: image_info = {image_info}")
                        image_path = self._extract_image_path(image_info)
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                filename = self._get_image_filename(image_info, image_path)
                                self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        else:
            self.doc.add_paragraph("Number of test scenarios: [To be determined]")
            
        # 8.2 Test Bed Diagram
        test_bed_diagram_heading = data.get('test_bed_diagram_heading', 'Test Bed Diagram')
        self.add_numbered_subheading(test_bed_diagram_heading, is_main_header=False)
        
        # Add test bed diagram notes if available
        test_bed_notes = data.get('test_bed_diagram_notes', '')
        if not test_bed_notes and 'test_bed_diagram' in data:
            test_bed_diagram = data['test_bed_diagram']
            if isinstance(test_bed_diagram, dict):
                test_bed_notes = test_bed_diagram.get('notes', '')
        
        if test_bed_notes:
            print(f"🔍 DEBUG 8.2: Processing test bed notes: {test_bed_notes}")
            # Use DocumentHelper to format the notes with smart formatting
            DocumentHelper.add_smart_subheading_paragraph(self.doc, test_bed_notes)
        
        test_bed_images = data.get('test_bed_images', [])
        test_bed_scripts = data.get('test_bed_scripts', [])
        
        # Handle test bed images using resolved images
        resolved_test_bed_images = data.get('resolved_test_bed_images', test_bed_images)
        print(f"🔍 DEBUG 8.2: resolved_test_bed_images = {resolved_test_bed_images}")
        
        # Also check for images in test_bed_diagram structure
        if not resolved_test_bed_images and 'test_bed_diagram' in data:
            test_bed_diagram = data['test_bed_diagram']
            if isinstance(test_bed_diagram, dict) and 'images' in test_bed_diagram:
                resolved_test_bed_images = test_bed_diagram['images']
                print(f"🔍 DEBUG 8.2: Found {len(resolved_test_bed_images)} images in test_bed_diagram")
        
        if resolved_test_bed_images:
            for i, image_info in enumerate(resolved_test_bed_images):
                print(f"🔍 DEBUG 8.2: Processing image {i}: {image_info}")
                image_path = self._extract_image_path(image_info)
                if image_path and os.path.exists(image_path):
                    if self._add_centered_image(image_path, width_cm=15):  # Keep 15cm width as requested
                        # Image was successfully added, add the caption
                        caption_para = self.doc.add_paragraph(f"Figure 8.2.{i+1}: Test Bed Diagram")
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Apply font styling to the caption - bold, italic, Calibri 10pt
                        for run in caption_para.runs:
                            DocumentHelper.set_font_style(run, font_name='Calibri', font_size=10, bold=True, italic=True)
                    else:
                        # Image failed to load, show error message
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                elif isinstance(image_info, str) and image_info.startswith('_') and image_info.endswith('_placeholder'):
                    # This is a placeholder, add it as text
                    self.doc.add_paragraph(f"[{image_info}]")
                else:
                    filename = self._get_image_filename(image_info, image_path)
                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Handle test bed scripts using resolved scripts or fallback to regular scripts
        resolved_test_bed_scripts = data.get('resolved_test_bed_scripts', test_bed_scripts)
        
        # Check for scripts in test_bed_diagram structure first (this is where the data actually is)
        if not resolved_test_bed_scripts and 'test_bed_diagram' in data:
            test_bed_diagram = data['test_bed_diagram']
            if isinstance(test_bed_diagram, dict) and 'scripts' in test_bed_diagram:
                resolved_test_bed_scripts = test_bed_diagram['scripts']
                print(f"🔍 DEBUG 8.2: Found {len(resolved_test_bed_scripts)} scripts in test_bed_diagram")
        
        # Check for scripts in the main data structure as well
        if not resolved_test_bed_scripts:
            # Look for test bed scripts in the main data
            all_scripts = data.get('scripts', [])
            if all_scripts:
                # Filter for test bed diagram scripts
                resolved_test_bed_scripts = [script for script in all_scripts if 
                    isinstance(script, dict) and 
                    script.get('subheading') == '8.2. Test Bed Diagram']
        
        print(f"🔍 DEBUG 8.2: resolved_test_bed_scripts = {resolved_test_bed_scripts}")
        print(f"🔍 DEBUG 8.2: test_bed_scripts = {test_bed_scripts}")
        print(f"🔍 DEBUG 8.2: test_bed_diagram in data = {'test_bed_diagram' in data}")
        print(f"🔍 DEBUG 8.2: data keys = {list(data.keys())}")
        if 'test_bed_diagram' in data:
            print(f"🔍 DEBUG 8.2: test_bed_diagram = {data['test_bed_diagram']}")
        print(f"🔍 DEBUG 8.2: all scripts in data = {data.get('scripts', [])}")
        
        # Check if test_bed_diagram is in configuration
        if 'configuration' in data and 'test_bed_diagram' in data['configuration']:
            print(f"🔍 DEBUG 8.2: Found test_bed_diagram in configuration: {data['configuration']['test_bed_diagram']}")
            resolved_test_bed_scripts = data['configuration']['test_bed_diagram'].get('scripts', [])
        
        if resolved_test_bed_scripts:
            print(f"🔍 DEBUG 8.2: Processing {len(resolved_test_bed_scripts)} scripts")
            for i, script_info in enumerate(resolved_test_bed_scripts):
                print(f"🔍 DEBUG 8.2: Script {i}: {script_info}")
                # Always check for description first - this is the placeholder
                if isinstance(script_info, dict):
                    script_description = script_info.get('description', '')
                    print(f"🔍 DEBUG 8.2: Script {i} description = '{script_description}'")
                    if script_description:
                        print(f"🔍 DEBUG 8.2: Adding placeholder '{script_description}' to document")
                        # Add description as centered placeholder text
                        paragraph = self.doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(script_description)
                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                        paragraph.paragraph_format.space_after = Pt(6)
                        continue  # Skip the rest of the processing for this script
                
                # If no description, process as regular script (if it has absolute_path and exists)
                if isinstance(script_info, dict) and 'absolute_path' in script_info and os.path.exists(script_info['absolute_path']):
                    # This is a resolved script file
                    try:
                        script_path = script_info['absolute_path']
                        script_filename = script_info.get('filename', '')
                        script_content = script_info.get('content', '')

                        # If we have content directly, use it, otherwise read from file
                        if script_content:
                            final_content = script_content
                        else:
                            with open(script_path, 'r', encoding='utf-8') as f:
                                final_content = f.read()

                        self.doc.add_paragraph(f"Script: {script_filename}")
                        self.doc.add_paragraph(final_content)
                    except Exception as e:
                        self.doc.add_paragraph(f"[Script could not be loaded: {script_info.get('filename', 'unknown')}]")
                else:
                    # Fallback for scripts without absolute_path or that don't exist
                    script_filename = script_info.get('filename', '') if isinstance(script_info, dict) else str(script_info)
                    self.doc.add_paragraph(f"[Script placeholder: {script_filename}]")
        
        if not test_bed_images and not test_bed_scripts and not resolved_test_bed_scripts:
            # No test bed content available - section will be empty
            pass
        
        # Add Test Bed Diagram placeholders
        self.add_section_placeholders(data, '8.2. Test Bed Diagram')
        
        # 8.3 Tools Required
        tools_required_heading = data.get('tools_required_heading', 'Tools Required')
        self.add_numbered_subheading(tools_required_heading, is_main_header=False)
        tools = data.get('tools_required', [])
        print(f"🔍 DEBUG 8.3: tools_required = {tools}")
        for i, tool in enumerate(tools):
            print(f"🔍 DEBUG 8.3: Tool {i}: {tool}")
            # Extract tool name from dictionary or use as string
            tool_name = tool.get('name', tool) if isinstance(tool, dict) else tool
            print(f"🔍 DEBUG 8.3: Tool {i} name = '{tool_name}'")
            
            # If tool name is empty, check for script descriptions as placeholders
            if not tool_name and isinstance(tool, dict):
                scripts = tool.get('scripts', [])
                print(f"🔍 DEBUG 8.3: Tool {i} has {len(scripts)} scripts")
                if scripts:
                    # Use the first script's description as placeholder
                    first_script = scripts[0]
                    if isinstance(first_script, dict):
                        tool_name = first_script.get('description', '')
                        print(f"🔍 DEBUG 8.3: Tool {i} using script description as name = '{tool_name}'")
            
            if tool_name:
                # Check if this is a placeholder (script description) - if tool name came from script description
                if isinstance(tool, dict) and tool.get('scripts'):
                    # This is a placeholder from script, display as centered text
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(tool_name)
                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                    paragraph.paragraph_format.space_after = Pt(6)
                else:
                    # This is a regular tool name, use smart formatting
                    paragraph = DocumentHelper.add_smart_subheading_paragraph(self.doc, tool_name)
            
            # Handle images for tools
            if isinstance(tool, dict):
                tool_images = tool.get('images', [])
                print(f"🔍 DEBUG 8.3: Tool {i} has {len(tool_images)} images")
                if tool_images:
                    for image_info in tool_images:
                        print(f"🔍 DEBUG 8.3: Tool {i} image: {image_info}")
                        image_path = self._extract_image_path(image_info)
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                filename = self._get_image_filename(image_info, image_path)
                                self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
        
        # Add Tools Required placeholders
        self.add_section_placeholders(data, '8.3. Tools Required')
            
        # 8.4 Test Execution Steps
        test_execution_steps_heading = data.get('test_execution_steps_heading', 'Test Execution Steps')
        self.add_numbered_subheading(test_execution_steps_heading, is_main_header=False)
        execution_steps = data.get('execution_steps', [])
        manual_steps = data.get('manual_execution_steps', [])
        print(f"🔍 DEBUG 8.4: execution_steps = {execution_steps}")
        print(f"🔍 DEBUG 8.4: manual_steps = {manual_steps}")
        
        # Add notes right after the heading if they exist in any execution step
        if execution_steps:
            # Collect all notes from all execution steps
            all_notes = []
            for scenario_steps in execution_steps:
                note_section = scenario_steps.get('note_section', [])
                for note in note_section:
                    if note.strip():
                        # Remove <note> tag if present
                        clean_note = note.strip()
                        if clean_note.startswith('<note>'):
                            clean_note = clean_note[6:].strip()  # Remove '<note>'
                        all_notes.append(clean_note)
            
            # Add all notes right after the heading
            for note in all_notes:
                if note:
                    # Add note with special formatting - red "Note:" and normal content
                    note_para = self.doc.add_paragraph()
                    
                    # Add "Note:" prefix in red, bold, italic, 12pt Calibri
                    note_prefix_run = note_para.add_run("Note: ")
                    DocumentHelper.set_font_style(note_prefix_run, font_name='Calibri', font_size=12, bold=True, italic=True, color='#FF0000')
                    
                    # Add note content in normal color, bold, italic, 12pt Calibri
                    note_content_run = note_para.add_run(note)
                    DocumentHelper.set_font_style(note_content_run, font_name='Calibri', font_size=12, bold=True, italic=True)
                    
                    # Add some spacing after note
                    note_para.paragraph_format.space_after = Pt(6)
        
        if execution_steps:
            print(f"🔍 DEBUG 8.4: Processing {len(execution_steps)} execution step scenarios")
            for i, scenario_steps in enumerate(execution_steps):
                print(f"🔍 DEBUG 8.4: Scenario {i}: {scenario_steps}")
                steps = scenario_steps.get('steps', [])
                custom_inputs = scenario_steps.get('Custom_input', [])
                note_section = scenario_steps.get('note_section', [])
                print(f"🔍 DEBUG 8.4: Scenario {i} has {len(steps)} steps, {len(custom_inputs)} custom inputs, and {len(note_section)} notes")
                print(f"🔍 DEBUG 8.4: Custom_inputs array: {custom_inputs}")
                print(f"🔍 DEBUG 8.4: note_section array: {note_section}")
                for idx, custom_input in enumerate(custom_inputs):
                    if custom_input:
                        print(f"🔍 DEBUG 8.4: Custom input {idx}: '{custom_input}'")
                for idx, note in enumerate(note_section):
                    if note:
                        print(f"🔍 DEBUG 8.4: Note {idx}: '{note}'")
                
                # Custom inputs will be added after each corresponding step
                # Notes will be added after the scenario description
                
                # Add scenario key and description if they exist
                scenario_key = scenario_steps.get('scenario_key', '')
                scenario_description = scenario_steps.get('scenario_description', '')
                if scenario_key or scenario_description:
                    # Add scenario key with bold formatting
                    if scenario_key:
                        scenario_header = f"{scenario_key}:"
                        scenario_para = self.doc.add_paragraph()
                        scenario_run = scenario_para.add_run(scenario_header)
                        scenario_run.font.bold = True
                        scenario_run.font.name = 'Calibri'
                        scenario_run.font.size = Pt(12)
                        scenario_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    
                    # Add scenario description with smart formatting
                    if scenario_description:
                        # Split text by lines to handle multiple paragraphs
                        lines = scenario_description.split('\n')
                        
                        for line in lines:
                            if not line.strip():
                                # Empty line - add paragraph break
                                self.doc.add_paragraph()
                                continue
                                
                            # Use smart formatting for each line
                            DocumentHelper.add_smart_subheading_paragraph(self.doc, line)
                    
                    # No extra spacing for tight layout
                
                # Notes are now processed at the beginning of the section, not here
                
                # Process each step and its corresponding custom input
                for j, step in enumerate(steps):
                    print(f"🔍 DEBUG 8.4: Step {j}: {step}")
                    # Check if step is a dict with script info, otherwise treat as string
                    if isinstance(step, dict):
                        step_text = step.get('text', step.get('step', ''))
                        step_scripts = step.get('scripts', [])
                        print(f"🔍 DEBUG 8.4: Step {j} text = '{step_text}', scripts = {step_scripts}")
                        
                        # Add the step text with tight line spacing
                        if step_text:
                            DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_text, line_spacing=0.8)
                        
                        # Add script placeholders for this step
                        if step_scripts:
                            print(f"🔍 DEBUG 8.4: Step {j} has {len(step_scripts)} scripts")
                            for k, script_info in enumerate(step_scripts):
                                print(f"🔍 DEBUG 8.4: Step {j} Script {k}: {script_info}")
                                if isinstance(script_info, dict):
                                    script_description = script_info.get('description', '')
                                    print(f"🔍 DEBUG 8.4: Step {j} Script {k} description = '{script_description}'")
                                    if script_description:
                                        print(f"🔍 DEBUG 8.4: Adding placeholder '{script_description}' to document")
                                        # Add description as centered placeholder text
                                        paragraph = self.doc.add_paragraph()
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = paragraph.add_run(script_description)
                                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                        paragraph.paragraph_format.space_after = Pt(6)
                        
                        # Handle images for this step
                        step_images = step.get('images', [])
                        print(f"🔍 DEBUG 8.4: Step {j} has {len(step_images)} images")
                        if step_images:
                            for k, image_info in enumerate(step_images):
                                print(f"🔍 DEBUG 8.4: Step {j} Image {k}: {image_info}")
                                image_path = self._extract_image_path(image_info)
                                if image_path and os.path.exists(image_path):
                                    if not self._add_centered_image(image_path):
                                        filename = self._get_image_filename(image_info, image_path)
                                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                                else:
                                    filename = self._get_image_filename(image_info, image_path)
                                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                        
                        # Add custom input for this step if it exists
                        if j < len(custom_inputs) and custom_inputs[j]:
                            step_custom_input = custom_inputs[j]
                            print(f"🔍 DEBUG 8.4: Adding custom input for step {j}: '{step_custom_input}'")
                            
                            # Use smart formatting for custom input with tight line spacing
                            DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
                            
                            # No extra spacing for tight layout
                    else:
                        # Handle string steps (backward compatibility)
                        print(f"🔍 DEBUG 8.4: Step {j} is string: '{step}'")
                        DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step, line_spacing=0.8)
                        
                        # Add custom input for this step if it exists
                        if j < len(custom_inputs) and custom_inputs[j]:
                            step_custom_input = custom_inputs[j]
                            print(f"🔍 DEBUG 8.4: Adding custom input for step {j}: '{step_custom_input}'")
                            
                            # Use smart formatting for custom input with tight line spacing
                            DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
                            
                            # No extra spacing for tight layout
                        
                        # No extra spacing for tight layout
        
        # Add manual execution steps if any
        if manual_steps:
            self.doc.add_paragraph("Manual Steps (Not linked to scenarios):")
            for j, step in enumerate(manual_steps):
                # Check if step is a dict with script info, otherwise treat as string
                if isinstance(step, dict):
                    step_text = step.get('text', step.get('step', ''))
                    step_scripts = step.get('scripts', [])
                    
                    # Add the step text with tight line spacing
                    if step_text:
                        DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_text, line_spacing=0.8)
                    
                    # Add script placeholders for this step
                    if step_scripts:
                        for script_info in step_scripts:
                            if isinstance(script_info, dict):
                                script_description = script_info.get('description', '')
                                if script_description:
                                    # Add description as centered placeholder text
                                    paragraph = self.doc.add_paragraph()
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run(script_description)
                                    DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                    paragraph.paragraph_format.space_after = Pt(6)
                else:
                    # Handle string steps (backward compatibility) with tight line spacing
                    DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step, line_spacing=0.8)
        
        if not execution_steps and not manual_steps:
            self.doc.add_paragraph("[Test execution steps to be defined]")
        
        # Add Test Execution Steps placeholders
        self.add_section_placeholders(data, '8.4. Test Execution Steps')
        
        # Also check for execution step scripts in the data structure
        execution_step_scripts = data.get('execution_step_scripts', [])
        print(f"🔍 DEBUG 8.4: execution_step_scripts = {execution_step_scripts}")
        
        # Check for execution step scripts in the main data structure as well
        if not execution_step_scripts:
            # Look for execution step scripts in the main data
            all_scripts = data.get('scripts', [])
            if all_scripts:
                # Filter for execution step scripts
                execution_step_scripts = [script for script in all_scripts if 
                    isinstance(script, dict) and 
                    script.get('subheading') == '8.4. Test Execution Steps']
                print(f"🔍 DEBUG 8.4: Found {len(execution_step_scripts)} execution step scripts in main data")
        
        # Check if execution steps are in configuration
        if not execution_step_scripts and 'configuration' in data and 'execution_steps' in data['configuration']:
            execution_steps_config = data['configuration']['execution_steps']
            if execution_steps_config:
                # Look for scripts in execution steps
                for step_data in execution_steps_config:
                    if isinstance(step_data, dict) and 'scripts' in step_data:
                        execution_step_scripts.extend(step_data['scripts'])
                print(f"🔍 DEBUG 8.4: Found {len(execution_step_scripts)} execution step scripts in configuration")
        
        if execution_step_scripts:
            print(f"🔍 DEBUG 8.4: Processing {len(execution_step_scripts)} fallback execution step scripts")
            for i, script_info in enumerate(execution_step_scripts):
                print(f"🔍 DEBUG 8.4: Fallback Script {i}: {script_info}")
                if isinstance(script_info, dict):
                    script_description = script_info.get('description', '')
                    print(f"🔍 DEBUG 8.4: Fallback Script {i} description = '{script_description}'")
                    if script_description:
                        print(f"🔍 DEBUG 8.4: Adding fallback placeholder '{script_description}' to document")
                        # Add description as centered placeholder text
                        paragraph = self.doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(script_description)
                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                        paragraph.paragraph_format.space_after = Pt(6)
        
        # Add section 8 notes if available
        section_8_notes = data.get('section_8_notes', [])
        if section_8_notes:
            self.add_section_notes(section_8_notes)
            
    def add_section_9_expected_result(self, data):
        """Section 9: Expected Result"""
        if not self.doc:
            return
            
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        expected_results_heading = data.get('expected_results_heading', 'Expected Result')
        self.add_numbered_section_heading(expected_results_heading, is_main_header=True)
        
        # Get expected results from new data format
        expected_results_list = data.get('expected_results', [])
        
        if expected_results_list:
            # Use the new expected results format
            for i, result_data in enumerate(expected_results_list):
                scenario_key = result_data.get('scenario_key', '')
                expected_result = result_data.get('expected_result', '')
                
                if expected_result:
                    # Use smart formatting for expected result with star-based bold support
                    DocumentHelper.add_smart_subheading_paragraph(self.doc, expected_result)
                else:
                    self.doc.add_paragraph("[Expected result to be defined]")
                
                # Handle images for expected results
                result_images = result_data.get('images', [])
                if result_images:
                    for image_info in result_images:
                        image_path = self._extract_image_path(image_info)
                        if image_path and os.path.exists(image_path):
                            if not self._add_centered_image(image_path):
                                filename = self._get_image_filename(image_info, image_path)
                                self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                        else:
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                
                # Handle script placeholders for expected results
                result_scripts = result_data.get('scripts', [])
                if result_scripts:
                    for script_info in result_scripts:
                        if isinstance(script_info, dict):
                            script_description = script_info.get('description', '')
                            if script_description:
                                # Add description as centered placeholder text
                                paragraph = self.doc.add_paragraph()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run(script_description)
                                DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                paragraph.paragraph_format.space_after = Pt(6)
        else:
            # Fallback to old format
            expected_results = data.get('expected_results', '')
            if expected_results:
                self.doc.add_paragraph(expected_results)
            # No hardcoded fallback - only show content if user has provided it
        
        # Add Expected Results placeholders
        self.add_section_placeholders(data, '9. Expected Results for Pass')
                
    def add_section_10_evidence_format(self, data):
        """Section 10: Expected Format of Evidence"""
        if not self.doc:
            return
            
        # Use dynamic heading from UI or fallback to default (without hardcoded number)
        expected_format_evidence_heading = data.get('expected_format_evidence_heading', 'Expected Format of Evidence')
        self.add_numbered_section_heading(expected_format_evidence_heading, is_main_header=True)
        
        # Get evidence format from new data format
        evidence_format_list = data.get('evidence_format', [])
        
        if evidence_format_list:
            # Use the new evidence format list - remove bullet points
            for evidence_item in evidence_format_list:
                if evidence_item:
                    # Extract evidence text from dictionary or use as string
                    evidence_text = evidence_item.get('evidence_text', evidence_item) if isinstance(evidence_item, dict) else evidence_item
                    if evidence_text:
                        DocumentHelper.add_smart_subheading_paragraph(self.doc, evidence_text)
                    
                    # Handle images for evidence format
                    if isinstance(evidence_item, dict):
                        evidence_images = evidence_item.get('images', [])
                        if evidence_images:
                            for image_info in evidence_images:
                                image_path = self._extract_image_path(image_info)
                                if image_path and os.path.exists(image_path):
                                    if not self._add_centered_image(image_path):
                                        filename = self._get_image_filename(image_info, image_path)
                                        self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                                else:
                                    filename = self._get_image_filename(image_info, image_path)
                                    self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                    
                    # Handle script placeholders for evidence format
                    if isinstance(evidence_item, dict):
                        evidence_scripts = evidence_item.get('scripts', [])
                        if evidence_scripts:
                            for script_info in evidence_scripts:
                                if isinstance(script_info, dict):
                                    script_description = script_info.get('description', '')
                                    if script_description:
                                        # Add description as centered placeholder text
                                        paragraph = self.doc.add_paragraph()
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = paragraph.add_run(script_description)
                                        DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                                        paragraph.paragraph_format.space_after = Pt(6)
        else:
            # Fallback to old format
            evidence_format = data.get('evidence_format', '')
            if evidence_format:
                self.doc.add_paragraph(evidence_format)
            # No hardcoded fallback - only show content if user has provided it
        
        # Add Evidence Format placeholders
        self.add_section_placeholders(data, '10. Expected Format of Evidence')
            
    def add_section_11_test_execution(self, data):
        """Section 11: Test Execution (11.1.1 to 11.1.N for each scenario)"""
        if not self.doc:
            return
            
        self.add_styled_heading("11. Test Execution", is_main_header=True)
        
        # Check for test execution cases from Section 11
        test_execution_cases = data.get('test_execution_cases', [])
        print(f"🔍 DEBUG 11: Found {len(test_execution_cases)} test execution cases")
        print(f"🔍 DEBUG 11: data keys = {list(data.keys())}")
        if test_execution_cases:
            print(f"🔍 DEBUG 11: test_execution_cases[0] = {test_execution_cases[0]}")
            if 'steps_data' in test_execution_cases[0] and test_execution_cases[0]['steps_data']:
                print(f"🔍 DEBUG 11: steps_data[0] = {test_execution_cases[0]['steps_data'][0]}")
                if 'placeholders' in test_execution_cases[0]['steps_data'][0]:
                    print(f"🔍 DEBUG 11: placeholders = {test_execution_cases[0]['steps_data'][0]['placeholders']}")
            else:
                print(f"🔍 DEBUG 11: steps_data is empty or missing")
        
        if test_execution_cases:
            # FIXED: Use ONLY test_execution_cases when available to prevent duplication
            # Clear any execution_blocks from data to prevent fallback processing
            print(f"🔍 DEBUG 11: test_execution_cases found, clearing execution_blocks to prevent duplication")
            if 'execution_blocks' in data:
                data['execution_blocks'] = []  # Clear execution_blocks to prevent fallback
            
            # Use the resolved test execution cases (with resolved filenames)
            resolved_test_execution_cases = data.get('test_execution_cases_resolved', test_execution_cases)
            print(f"🔍 DEBUG 11: Using {len(resolved_test_execution_cases)} resolved test execution cases")
            for case_data in resolved_test_execution_cases:
                print(f"🔍 DEBUG 11: Processing case {case_data.get('case_number', 'unknown')} with {len(case_data.get('steps_data', []))} steps_data")
                print(f"🔍 DEBUG 11: case_data keys = {list(case_data.keys())}")
                if 'steps_data' in case_data:
                    for i, step_data in enumerate(case_data['steps_data']):
                        print(f"🔍 DEBUG 11: step_data {i} keys = {list(step_data.keys())}")
                        if 'placeholders' in step_data:
                            print(f"🔍 DEBUG 11: step_data {i} placeholders = {step_data['placeholders']}")
                self.add_test_execution_case(case_data, data)
        else:
            # Fallback to old method - only when test_execution_cases is empty
            scenarios = data.get('test_scenarios', [])
            execution_blocks = data.get('execution_blocks', [])
            
            print(f"🔍 DEBUG 11: No test_execution_cases found, using fallback method")
            print(f"🔍 DEBUG 11: execution_blocks count: {len(execution_blocks)}, scenarios count: {len(scenarios)}")
            
            # Use execution blocks if available, otherwise generate from scenarios
            if execution_blocks:
                for i, block in enumerate(execution_blocks):
                    self.add_execution_block(i+1, block)
            elif scenarios:
                for i, scenario in enumerate(scenarios):
                    self.add_execution_block_from_scenario(i+1, scenario)
            else:
                # Placeholder execution block
                self.add_placeholder_execution_block(1)
                
    def add_test_execution_case(self, case_data, data=None):
        """Add a test execution case with proper step images and evidence mapping"""
        if not self.doc:
            return
            
        case_number = case_data.get('case_number', '')
        scenario_key = case_data.get('scenario_key', '')
        itsar_description = case_data.get('itsar_description', '')
        test_case_name = case_data.get('test_case_name', '')
        test_case_description = case_data.get('test_case_description', '')
        test_observations = case_data.get('test_observations', '')
        evidence_provided = case_data.get('evidence_provided', '')
        execution_steps = case_data.get('execution_steps', [])
        
        # Add case heading
        self.add_styled_heading(f"11.1.{case_number} {scenario_key}", is_main_header=False)
        
        # Add ITSAR description
        if itsar_description:
            DocumentHelper.add_smart_subheading_paragraph(self.doc, f" {itsar_description}")
        
        # Add test case name
        if test_case_name:
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("a. Test Case Name:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            DocumentHelper.add_smart_subheading_paragraph(self.doc, test_case_name)
        
        # Add test case description
        if test_case_description:
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("b. Test Case Description:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            DocumentHelper.add_smart_subheading_paragraph(self.doc, test_case_description)
        
        # Add execution steps - MOVED TO APPEAR IMMEDIATELY AFTER TEST CASE DESCRIPTION
        # FIXED: Only process execution steps if they haven't been processed already
        steps_data = case_data.get('steps_data', [])
        has_steps_data = len(steps_data) > 0
        
        if execution_steps and not has_steps_data:
            # Only process execution steps if there's no steps_data (to prevent duplication)
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("c. Execution Steps:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            
            # Get custom inputs from original Section 8.4 data
            custom_inputs = []
            if data and 'execution_steps' in data:
                for scenario_steps in data['execution_steps']:
                    if scenario_steps.get('scenario_key') == scenario_key:
                        custom_inputs = scenario_steps.get('Custom_input', [])
                        print(f"🔍 DEBUG 11: Found {len(custom_inputs)} custom inputs for scenario {scenario_key}")
                        break
            
            # Process execution steps without steps_data (simple case)
            for i, step in enumerate(execution_steps):
                step_number = i + 1
                print(f"🔍 DEBUG 11: Processing execution step {step_number}: '{step[:50]}...'")
                
                DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step, line_spacing=0.8)
                
                # Add custom input for this step if it exists (same logic as Section 8.4)
                if i < len(custom_inputs) and custom_inputs[i]:
                    step_custom_input = custom_inputs[i]
                    print(f"🔍 DEBUG 11: Adding custom input for step {step_number}: '{step_custom_input}'")
                    DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
        
        elif has_steps_data:
            # FIXED: Process execution steps from steps_data (when uploads are present)
            print(f"🔍 DEBUG 11: Processing execution steps from steps_data for case {case_number}")
            
            # Add execution steps heading
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("c. Execution Steps:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            
            # Get custom inputs from original Section 8.4 data
            custom_inputs = []
            if data and 'execution_steps' in data:
                for scenario_steps in data['execution_steps']:
                    if scenario_steps.get('scenario_key') == scenario_key:
                        custom_inputs = scenario_steps.get('Custom_input', [])
                        print(f"🔍 DEBUG 11: Found {len(custom_inputs)} custom inputs for scenario {scenario_key}")
                        break
            
            # FIXED: Process steps_data directly to avoid duplicate processing
            # This prevents the same step data from being processed multiple times
            # which was causing values to be placed 4 times instead of once
            print(f"🔍 DEBUG 11: Processing {len(steps_data)} steps_data entries for case {case_number}")
            
            # Process execution steps from steps_data
            for step_info in steps_data:
                step_number = step_info.get('step_index', 0)
                step_text = step_info.get('step_text', '')
                print(f"🔍 DEBUG 11: Processing execution step {step_number}: '{step_text[:50]}...'")
                
                # Add the step text from steps_data
                DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_text, line_spacing=0.8)
                
                # Add custom input for this step if it exists (same logic as Section 8.4)
                if step_number <= len(custom_inputs) and custom_inputs[step_number - 1]:
                    step_custom_input = custom_inputs[step_number - 1]
                    print(f"🔍 DEBUG 11: Adding custom input for step {step_number}: '{step_custom_input}'")
                    DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
                
                # Process the additional data (images, scripts, placeholders) from steps_data
                print(f"🔍 DEBUG 11: Found step_info for step {step_number} with {len(step_info.get('resolved_scripts', []))} scripts and {len(step_info.get('placeholders', []))} placeholders")
                
                # Add images for this step (images are now filenames, need to resolve from global arrays)
                resolved_images = step_info.get('resolved_images', [])
                print(f"🔍 DEBUG 11: Processing {len(resolved_images)} resolved_images for step {step_number}")
                for image_info in resolved_images:
                    print(f"🔍 DEBUG 11: image_info = {image_info}")
                    image_path = self._extract_image_path(image_info)
                    print(f"🔍 DEBUG 11: extracted image_path = {image_path}")
                    if image_path and os.path.exists(image_path):
                        if self._add_centered_image(image_path):
                            # Add figure caption with user-provided description
                            description = image_info.get('description', '')
                            figure_number = image_info.get('figure_number', f"11.1.{case_number}")
                            
                            if description:
                                caption_para = self.doc.add_paragraph(f"Figure {figure_number}: {description}")
                            else:
                                caption_para = self.doc.add_paragraph(f"Figure {figure_number}: [No description provided]")
                            
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Apply font styling to the caption - Calibri, size 10, bold, italic
                            for run in caption_para.runs:
                                DocumentHelper.set_font_style(run, font_name='Calibri', font_size=10, bold=True, italic=True)
                        else:
                            # Image failed to load, show error message
                            filename = self._get_image_filename(image_info, image_path)
                            self.doc.add_paragraph(f"[Image could not be loaded: {filename}]")
                    else:
                        filename = self._get_image_filename(image_info, image_path)
                        self.doc.add_paragraph(f"[Image placeholder: {filename}]")
                
                # Add scripts for this step (scripts are now filenames, need to resolve from global arrays)
                resolved_scripts = step_info.get('resolved_scripts', [])
                for script_info in resolved_scripts:
                    if isinstance(script_info, dict):
                        description = script_info.get('description', '')
                        filename = script_info.get('original_filename', script_info.get('filename', ''))

                        if description:
                            # Display script description as centered, italic placeholder
                            placeholder_para = self.doc.add_paragraph()
                            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            placeholder_run = placeholder_para.add_run(description)
                            DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, italic=True)
                            print(f"🔍 DEBUG 11: Added script placeholder '{description}' for step {step_number}")

                        # if filename:
                        #     self.doc.add_paragraph(f"Script file: {filename}")
                
                # Add placeholders for this step
                step_placeholders = step_info.get('placeholders', [])
                print(f"🔍 DEBUG 11: Processing {len(step_placeholders)} step_placeholders for step {step_number}")
                for placeholder_info in step_placeholders:
                    print(f"🔍 DEBUG 11: placeholder_info = {placeholder_info}")
                    if isinstance(placeholder_info, dict):
                        placeholder_name = placeholder_info.get('name', '')
                        if placeholder_name:
                            # Display placeholder as centered, italic text
                            placeholder_para = self.doc.add_paragraph()
                            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            placeholder_run = placeholder_para.add_run(placeholder_name)
                            DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, italic=True)
                            print(f"🔍 DEBUG 11: Added step placeholder '{placeholder_name}' for step {step_number}")
                    elif isinstance(placeholder_info, str):
                        # Display string placeholder as centered, italic text
                        placeholder_para = self.doc.add_paragraph()
                        placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        placeholder_run = placeholder_para.add_run(placeholder_info)
                        DocumentHelper.set_font_style(placeholder_run, font_name='Calibri', font_size=12, italic=True)
                        print(f"🔍 DEBUG 11: Added string placeholder '{placeholder_info}' for step {step_number}")
        
        # Add test observations
        if test_observations:
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("d. Test Observations:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            self.doc.add_paragraph(test_observations)
        
        # Add evidence provided
        if evidence_provided:
            subheading_para = self.doc.add_paragraph()
            subheading_run = subheading_para.add_run("e. Evidence Provided:")
            DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
            self.doc.add_paragraph(evidence_provided)
        
        # Add spacing between cases
        self.doc.add_paragraph()
        
    def add_execution_block(self, number, block_data):
        """Add execution block with proper structure"""
        if not self.doc:
            return
            
        block_title = f"11.1.{number} Test Case Number:"
        self.add_styled_heading(block_title, is_main_header=False)
        
        # Test Case Name
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("a. Test Case Name:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph(block_data.get('test_case_name', '[Test Case Name]'))
        
        # Test Case Description
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("b. Test Case Description:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph(block_data.get('description', '[Test Case Description]'))
        
        # Execution Steps - MOVED TO APPEAR IMMEDIATELY AFTER TEST CASE DESCRIPTION
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("c. Execution Steps:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        steps = block_data.get('steps', [])
        custom_inputs = block_data.get('Custom_input', [])
        
        if isinstance(steps, list):
            for i, step in enumerate(steps):
                DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step, line_spacing=0.8)
                
                # Add custom input for this step if it exists
                if i < len(custom_inputs) and custom_inputs[i]:
                    step_custom_input = custom_inputs[i]
                    print(f"🔍 DEBUG 11: Adding custom input for step {i}: '{step_custom_input}'")
                    DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
        else:
            DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, str(steps), line_spacing=0.8)
            
            # Add custom input if it exists
            if custom_inputs and custom_inputs[0]:
                DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, custom_inputs[0], line_spacing=0.8)
            
        # Test Observations
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("d. Test Observations:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph(block_data.get('observations', '[Test Observations]'))
        
        # Evidence Provided
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("e. Evidence Provided:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        evidence = block_data.get('evidence', '[Evidence Provided]')
        self.doc.add_paragraph(evidence)
        
        # Add space between blocks
        self.doc.add_paragraph()
        
    def add_execution_block_from_scenario(self, number, scenario):
        """Add execution block generated from scenario"""
        if not self.doc:
            return
            
        block_title = f"11.1.{number} Test Case Number:"
        self.add_styled_heading(block_title, is_main_header=False)
        
        # Test Case Name
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("a. Test Case Name:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph(scenario.get('test_case_name', '[Test Case Name]'))
        
        # Test Case Description
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("b. Test Case Description:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph(scenario.get('description', '[Test Case Description]'))
        
        # Execution Steps - MOVED TO APPEAR IMMEDIATELY AFTER TEST CASE DESCRIPTION
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("c. Execution Steps:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        steps = scenario.get('steps', [])
        custom_inputs = scenario.get('Custom_input', [])
        
        for i, step in enumerate(steps):
            DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step, line_spacing=0.8)
            
            # Add custom input for this step if it exists
            if i < len(custom_inputs) and custom_inputs[i]:
                step_custom_input = custom_inputs[i]
                print(f"🔍 DEBUG 11: Adding custom input for step {i}: '{step_custom_input}'")
                DocumentHelper.add_smart_subheading_paragraph_with_spacing(self.doc, step_custom_input, line_spacing=0.8)
            
        # Test Observations
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("d. Test Observations:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Test Observations to be filled during testing]")
        
        # Evidence Provided
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("e. Evidence Provided:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Screenshots and evidence to be inserted]")
        
        # Add space between blocks
        self.doc.add_paragraph()
        
    def add_placeholder_content(self, placeholder_content_management):
        """Add placeholder content with center alignment"""
        print(f"🔍 DEBUG add_placeholder_content: called with {placeholder_content_management}")
        if not self.doc or not placeholder_content_management:
            print(f"🔍 DEBUG add_placeholder_content: returning early - doc={self.doc is not None}, content={placeholder_content_management}")
            return
            
        for placeholder_item in placeholder_content_management:
            placeholder_text = placeholder_item.get('placeholder', '')
            print(f"🔍 DEBUG add_placeholder_content: processing placeholder_text = '{placeholder_text}'")
            if placeholder_text:
                # Create centered paragraph for placeholder
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add placeholder text with styling
                run = paragraph.add_run(placeholder_text)
                DocumentHelper.set_font_style(run, font_name='Calibri', font_size=12, bold=False, italic=True)
                
                # Add some spacing after placeholder
                paragraph.paragraph_format.space_after = Pt(6)
                print(f"🔍 DEBUG add_placeholder_content: added placeholder '{placeholder_text}' to document")

    def add_section_notes(self, notes_list):
        """Add section notes with red 'Note:' prefix and bold text"""
        if not self.doc or not notes_list:
            print(f"🔍 DEBUG add_section_notes: doc={bool(self.doc)}, notes_list={notes_list}")
            return
            
        print(f"🔍 DEBUG add_section_notes: Processing {len(notes_list)} notes")
        for i, note_info in enumerate(notes_list):
            print(f"🔍 DEBUG add_section_notes: Note {i}: {note_info}")
            note_text = note_info.get('text', '')
            print(f"🔍 DEBUG add_section_notes: Note {i} text: '{note_text}'")
            if note_text:
                # Create paragraph for note
                paragraph = self.doc.add_paragraph()
                
                # Add "Note:" prefix in red
                note_prefix_run = paragraph.add_run("Note: ")
                DocumentHelper.set_font_style(note_prefix_run, font_name='Calibri', font_size=12, bold=True, color='#FF0000')  # Red color
                
                # Add note text with markdown formatting support
                DocumentHelper.add_formatted_text_with_bold(paragraph, note_text)
                
                # Add some spacing after note
                paragraph.paragraph_format.space_after = Pt(6)
                print(f"🔍 DEBUG add_section_notes: Added note '{note_text}' to document")

    def add_placeholder_execution_block(self, number):
        """Add placeholder execution block"""
        if not self.doc:
            return
            
        block_title = f"11.1.{number} Test Case Number:"
        self.add_styled_heading(block_title, is_main_header=False)
        
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("a. Test Case Name:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Test Case Name]")
        
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("b. Test Case Description:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Test Case Description]")
        
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("c. Execution Steps:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Execution Steps]")
        
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("d. Test Observations:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Test Observations]")
        
        subheading_para = self.doc.add_paragraph()
        subheading_run = subheading_para.add_run("e. Evidence Provided:")
        DocumentHelper.set_font_style(subheading_run, font_name='Calibri', font_size=12, bold=True)
        self.doc.add_paragraph("[Evidence Provided]")
        
        self.doc.add_paragraph()
        
    def add_section_12_test_case_result(self, data):
        """Section 12: Test Case Result (Table with Pass/Fail and Remarks)"""
        if not self.doc:
            return
            
        self.add_styled_heading("12. Test Case Result", is_main_header=True)
        
        # Get test case results from new data format
        test_case_results = data.get('test_case_results', [])
        
        if test_case_results:
            # Create results table with new format
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'S. No'
            header_cells[0].paragraphs[0].runs[0].bold = True
            header_cells[1].text = 'TEST CASE No.'
            header_cells[1].paragraphs[0].runs[0].bold = True
            header_cells[2].text = 'PASS FAIL'
            header_cells[2].paragraphs[0].runs[0].bold = True
            header_cells[3].text = 'Remarks'
            header_cells[3].paragraphs[0].runs[0].bold = True
            
            # Add data rows from new format
            for result_data in test_case_results:
                row_cells = table.add_row().cells
                row_cells[0].text = str(result_data.get('serial_no', ''))
                row_cells[1].text = result_data.get('test_case_no', '')
                row_cells[2].text = result_data.get('pass_fail', 'NA')
                row_cells[3].text = result_data.get('remarks', '')
        else:
            # Fallback to old format - dynamically generate based on scenarios
            scenarios = data.get('test_scenarios', [])
            execution_blocks = data.get('execution_blocks', [])
            test_execution_cases = data.get('test_execution_cases', [])
            
            # Determine the number of scenarios to generate results for
            num_scenarios = 0
            if test_execution_cases:
                num_scenarios = len(test_execution_cases)
            elif execution_blocks:
                num_scenarios = len(execution_blocks)
            elif scenarios:
                num_scenarios = len(scenarios)
            else:
                num_scenarios = 1  # Default to 1 if no scenarios found
            
            # Create results table
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'S. No'
            header_cells[1].text = 'TEST CASE No.'
            header_cells[2].text = 'PASS FAIL'
            header_cells[3].text = 'Remarks'
            
            # Add data rows dynamically based on number of scenarios
            for i in range(num_scenarios):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = f"1.9.4.{i+1}"
                
                # Display the field names in DOCX instead of values
                result_key = f"result_{i+1}"
                remarks_key = f"remarks_{i+1}"
                
                row_cells[2].text = result_key
                row_cells[3].text = remarks_key
            
    def add_footer(self):
        """Add footer with company information"""
        if not self.doc or len(self.doc.sections) == 0:
            return
            
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        # footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | DELTAPHI LABS | Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _encode_image_as_base64(self, image_path):
        """Helper function to encode image as base64 data URL"""
        try:
            import base64
            with open(image_path, 'rb') as img_file:
                img_data = base64.b64encode(img_file.read()).decode('utf-8')
                # Get file extension for MIME type
                file_ext = os.path.splitext(image_path)[1].lower()
                mime_type = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.bmp': 'image/bmp',
                    '.webp': 'image/webp'
                }.get(file_ext, 'image/png')
                
                return f"data:{mime_type};base64,{img_data}"
        except Exception as e:
            print(f"🔧 DEBUG: Failed to encode image as base64: {e}")
            return None

    def generate_preview(self, data):
        """Generate HTML preview that matches the exact document structure"""
        html_content = """
        <html>
        <head>
            <style>
                body { 
                    font-family: 'Calibri', Arial, sans-serif; 
                    font-size: 12px; 
                    line-height: 1.4; 
                    background-color: #ffffff; 
                    color: #000000; 
                    margin: 20px;
                    max-width: 800px;
                }
                h1 { 
                    color: #000000; 
                    font-size: 18px; 
                    font-weight: bold;
                    text-align: center;
                    padding-bottom: 10px;
                    margin-bottom: 20px;
                }
                h2 { 
                    color: #000000; 
                    font-size: 14px; 
                    font-weight: bold;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                h3 { 
                    color: #000000; 
                    font-size: 12px; 
                    font-weight: bold;
                    margin-top: 15px;
                    margin-bottom: 8px;
                }
                .field { 
                    margin: 8px 0; 
                    line-height: 1.6;
                }
                .label { 
                    font-weight: bold; 
                    color: #000000; 
                }
                .value { 
                    color: #000000; 
                }
                table { 
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 15px 0; 
                    border: 1px solid #000000;
                }
                th, td { 
                    border: 1px solid #000000; 
                    padding: 8px; 
                    text-align: left; 
                    font-size: 11px;
                }
                th { 
                    background-color: #f0f0f0; 
                    font-weight: bold;
                }
                .section { 
                    margin: 15px 0; 
                    padding: 10px; 
                    border-left: 3px solid #000000; 
                    background-color: #fafafa;
                }
                .dut-details {
                    margin: 15px 0;
                    padding: 10px;
                    background-color: #f9f9f9;
                    border: 1px solid #ddd;
                }
                .hash-section {
                    margin: 5px 0;
                    padding: 5px;
                    background-color: #f9f9f9;
                    border: 1px solid #ddd;
                }
                .screenshot-info {
                    margin: 10px 0;
                    padding: 8px;
                    background-color: #f9f9f9;
                    border-left: 4px solid #666666;
                }
                .placeholder {
                    color: #666666;
                    font-style: italic;
                }
                .placeholder-content {
                    color: #000000;
                    font-style: italic;
                    text-align: center;
                    margin: 10px 0;
                    padding: 5px;
                    background-color: #f0f0f0;
                    border: 1px dashed #ccc;
                }
                .hash-heading {
                    color: #000000;
                    font-size: 12px;
                    font-weight: bold;
                    font-family: 'Calibri', Arial, sans-serif;
                    margin-bottom: 2px;
                    margin-top: 1px;
                }
                /* Page Break Styling */
                .page-break-indicator {
                    width: 100%;
                    height: 40px;
                    background: linear-gradient(to right, #ffffff 0%, #cccccc 50%, #ffffff 100%);
                    border-top: 2px dashed #999999;
                    border-bottom: 2px dashed #999999;
                    margin: 30px 0;
                    position: relative;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }
                .page-break-indicator::before {
                    content: "📄 PAGE BREAK 📄";
                    background-color: #ffffff;
                    color: #666666;
                    font-size: 12px;
                    font-weight: bold;
                    padding: 8px 15px;
                    border: 2px solid #999999;
                    border-radius: 5px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                .major-page-break {
                    width: 100%;
                    height: 50px;
                    background: linear-gradient(to right, #ffffff 0%, #000000 50%, #ffffff 100%);
                    border-top: 3px solid #000000;
                    border-bottom: 3px solid #000000;
                    margin: 40px 0;
                    position: relative;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }
                .major-page-break::before {
                    content: "🔄 NEW PAGE 🔄";
                    background-color: #ffffff;
                    color: #000000;
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px 20px;
                    border: 3px solid #000000;
                    border-radius: 8px;
                    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
                }
                .section-break {
                    width: 100%;
                    height: 20px;
                    background: linear-gradient(to right, #ffffff 0%, #dddddd 50%, #ffffff 100%);
                    border-top: 1px solid #dddddd;
                    border-bottom: 1px solid #dddddd;
                    margin: 20px 0;
                    position: relative;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }
                .section-break::before {
                    content: "📋 SECTION BREAK 📋";
                    background-color: #ffffff;
                    color: #666666;
                    font-size: 10px;
                    font-weight: bold;
                    padding: 5px 10px;
                    border: 1px solid #dddddd;
                    border-radius: 3px;
                }
            </style>
        </head>
        <body>
        """
        

        
        # Main Title
        test_report = data.get('test_report', 'TEST REPORT FOR: [Test Case Title]')
        html_content += f'<h1>{test_report}</h1>'
        
        # DUT Details Section (like in the document)
        dut_fields = data.get('dut_fields', [])
        if dut_fields:
            html_content += '<div class="dut-details">'
            for field in dut_fields:
                label = field.get('label', '')
                value = field.get('value', '')
                if label and value:
                    # Add label and value on same line with proper spacing
                    html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="label">{label}:</span> <span class="value">{value}</span></div>'
            html_content += '</div>'
        else:
            html_content += '<div class="dut-details">'
            html_content += '<div class="field"><span class="placeholder">[DUT Details to be added]</span></div>'
            html_content += '</div>'
        
        # Multiple Hash Sections
        hash_sections = data.get('hash_sections', [])
        
        for section in hash_sections:
            heading = section.get('heading', '')
            direct_hash_value = section.get('direct_hash_value', '')
            hash_fields = section.get('hash_fields', [])
            
            # Display hash section if there's a heading or any content
            if heading or direct_hash_value or hash_fields:
                html_content += f'<div class="hash-section">'
                 # Add colon to heading if it doesn't already end with one
                formatted_heading = heading.strip()
                if not formatted_heading.endswith(':'):
                    formatted_heading += ':'
                html_content += f'<h2 class="hash-heading" style="margin-bottom: 12px;">{formatted_heading}</h2>'
                
                # Add direct hash value if provided
                if direct_hash_value:
                    html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="value">{direct_hash_value}</span></div>'
                
                
                # Add direct hash value if provided
                if direct_hash_value:
                    html_content += f'<div class="field"><span class="value">{direct_hash_value}</span></div>'
                
                # Add direct hash images if available
                direct_hash_images = section.get('direct_hash_images', [])
                if direct_hash_images:
                    html_content += '<h3>Direct Hash Images</h3>'
                    for image_path in direct_hash_images:
                        if os.path.exists(image_path):
                            filename = os.path.basename(image_path)
                            # Try to encode the image as base64 for display
                            data_url = self._encode_image_as_base64(image_path)
                            if data_url:
                                print(f"🔧 DEBUG: Generated base64 data URL for direct hash image: {filename}")
                                html_content += f'''
                                <div class="image-container" style="text-align: center; margin: 20px 0;">
                                    <div class="screenshot-info">📷 {filename}</div>
                                    <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                                </div>
                                '''
                            else:
                                # Fallback to file:// URL
                                file_url = f"file:///{image_path.replace(os.sep, '/')}"
                                print(f"🔧 DEBUG: Generated file URL for direct hash image: {file_url}")
                                html_content += f'''
                                <div class="image-container" style="text-align: center; margin: 20px 0;">
                                    <div class="screenshot-info">📷 {filename}</div>
                                    <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                                </div>
                                '''
                
                # Add direct hash scripts if available
                direct_hash_scripts = section.get('direct_hash_scripts', [])
                if direct_hash_scripts:
                    html_content += '<h3>Direct Hash Scripts</h3>'
                    for script_info in direct_hash_scripts:
                        if isinstance(script_info, dict):
                            # New format with full script info
                            script_description = script_info.get('description', '')
                            script_filename = script_info.get('original_filename', script_info.get('filename', ''))
                            
                            if script_description:
                                html_content += f'<div class="field"><span class="value">{script_description}</span></div>'
                            if script_filename:
                                html_content += f'<div class="script-info">📄 {script_filename}</div>'
                        else:
                            # Old format with just path
                            script_path = script_info
                            if os.path.exists(script_path):
                                script_filename = os.path.basename(script_path)
                                html_content += f'<div class="script-info">📄 {script_filename}</div>'
                
                # Add individual hash fields
                for field in hash_fields:
                    field_heading = field.get('label', '')
                    field_value = field.get('value', '')
                    if field_heading:  # Show field if it has a heading, even if value is empty
                        # Add colon to field heading if it doesn't already end with one
                        formatted_field_heading = field_heading.strip()
                        if not formatted_field_heading.endswith(':'):
                            formatted_field_heading += ':'
                        html_content += f'<h3 class="hash-heading" style="margin-bottom: 12px;">{formatted_field_heading}</h3>'
                        if field_value:
                            html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="value">{field_value}</span></div>'
                        else:
                            html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="placeholder">[Hash value to be added]</span></div>'
                        
                        # Add hash field images if available
                        field_images = field.get('images', [])
                        if field_images:
                            html_content += '<h4>Field Images</h4>'
                            for image_path in field_images:
                                if os.path.exists(image_path):
                                    filename = os.path.basename(image_path)
                                    # Try to encode the image as base64 for display
                                    data_url = self._encode_image_as_base64(image_path)
                                    if data_url:
                                        print(f"🔧 DEBUG: Generated base64 data URL for field image: {filename}")
                                        html_content += f'''
                                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                                            <div class="screenshot-info">📷 {filename}</div>
                                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                                        </div>
                                        '''
                                    else:
                                        # Fallback to file:// URL
                                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                                        print(f"🔧 DEBUG: Generated file URL for field image: {file_url}")
                                        html_content += f'''
                                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                                            <div class="screenshot-info">📷 {filename}</div>
                                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                                        </div>
                                        '''
                        
                        # Add hash field scripts if available
                        field_scripts = field.get('scripts', [])
                        if field_scripts:
                            html_content += '<h4>Field Scripts</h4>'
                            for script_info in field_scripts:
                                if isinstance(script_info, dict):
                                    # New format with full script info
                                    script_description = script_info.get('description', '')
                                    script_filename = script_info.get('original_filename', script_info.get('filename', ''))
                                    
                                    if script_description:
                                        html_content += f'<div class="field"><span class="value">{script_description}</span></div>'
                                    if script_filename:
                                        html_content += f'<div class="script-info">📄 {script_filename}</div>'
                                else:
                                    # Old format with just path
                                    script_path = script_info
                                    if os.path.exists(script_path):
                                        script_filename = os.path.basename(script_path)
                                        html_content += f'<div class="script-info">📄 {script_filename}</div>'
                html_content += '</div>'
        
        # If no hash sections, add placeholder
        if not hash_sections:
            html_content += '<div class="hash-section">'
            html_content += '<div class="field"><span class="placeholder">[Hash Sections to be added]</span></div>'
            html_content += '</div>'
        
        # ITSAR Information for HTML preview
        itsar_fields = data.get('itsar_fields', [])
        if itsar_fields:
            for field in itsar_fields:
                key = field.get('key', '')
                value = field.get('value', '')
                if key:
                    # Check if value contains newlines (multi-line input)
                    if value and '\n' in value:
                        # Multi-line input - first show the heading, then bullet points
                        # Add heading first
                        html_content += f'<div class="field" style="margin-bottom: 6px;"><span class="label">{key}:</span></div>'
                        # Then add each line as numbered item (like DUT details format)
                        lines = value.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                # Add each line as a regular field (not bullet point)
                                html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="value">{line}</span></div>'
                    else:
                        # Single-line input or empty value - output in same single line
                        display_value = value if value else ""
                        html_content += f'<div class="field" style="margin-bottom: 12px;"><span class="label">{key}:</span> <span class="value">{display_value}</span></div>'
        
        # Section 1: ITSAR Section No & Name (Page Break)
        sections_1_7 = data.get('sections_1_7', {})
        section_1_data = sections_1_7.get('section_1', {})
        section_1_heading = section_1_data.get('heading', '1. ITSAR Section No & Name')
        
        html_content += '<div class="major-page-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{section_1_heading}</h2>"
        itsar_section = section_1_data.get('content', data.get('itsar_section', ''))
        if itsar_section:
            html_content += f'<div class="field"><span class="value">{itsar_section}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">[User input text]</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_1_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        # Add section 1 images if available
        section_1_images = section_1_data.get('images', [])
        if section_1_images:
            for i, image_info in enumerate(section_1_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                
                if os.path.exists(image_path):
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(image_path)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for section 1 image: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 1.{i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for section 1 image: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 1.{i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Image placeholder 1.{i}]</span></div>'
        
        html_content += '</div>'
        
        # Section 2: Security Requirement
        section_2_data = sections_1_7.get('section_2', {})
        section_2_heading = section_2_data.get('heading', '2. Security Requirement')
        
        html_content += '<div class="section">'
        html_content += f"<h2>{section_2_heading}</h2>"
        security_req = section_2_data.get('content', data.get('security_req', ''))
        if security_req:
            html_content += f'<div class="field"><span class="value">{security_req}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">[User input]</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_2_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        # Add section 2 images if available
        section_2_images = section_2_data.get('images', [])
        if section_2_images:
            for i, image_info in enumerate(section_2_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                
                if os.path.exists(image_path):
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(image_path)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for section 2 image: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 2.{i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for section 2 image: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 2.{i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Image placeholder 2.{i}]</span></div>'
        
        html_content += '</div>'
        
        # Section 3: Requirement Description
        section_3_data = sections_1_7.get('section_3', {})
        section_3_heading = section_3_data.get('heading', '3. Requirement Description')
        
        html_content += '<div class="section">'
        html_content += f"<h2>{section_3_heading}</h2>"
        req_description = section_3_data.get('content', data.get('req_description', ''))
        if req_description:
            formatted_description = self.format_text_for_html(req_description)
            html_content += f'<div class="field"><span class="value">{formatted_description}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">[Requirement description to be added]</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_3_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        # Add section 3 images if available
        section_3_images = section_3_data.get('images', [])
        if section_3_images:
            for i, image_info in enumerate(section_3_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                
                if os.path.exists(image_path):
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(image_path)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for section 3 image: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 3.{i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for section 3 image: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 3.{i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Image placeholder 3.{i}]</span></div>'
        
        html_content += '</div>'
        
        # Section 4: DUT Confirmation Details (Page Break)
        section_4_data = sections_1_7.get('section_4', {})
        section_4_heading = section_4_data.get('heading', '4. DUT Confirmation Details')
        
        html_content += '<div class="major-page-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{section_4_heading}</h2>"
        html_content += "<h3>DUT Images</h3>"
        screenshots = data.get('screenshots', [])
        if screenshots:
            for i, screenshot in enumerate(screenshots[:4], 1):
                if os.path.exists(screenshot):
                    filename = os.path.basename(screenshot)
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(screenshot)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for screenshot: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure {i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{screenshot.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for screenshot: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure {i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="screenshot-info">Figure {i}: [Image not found]</div>'
        # DUT Version info section removed as requested
        
        # DUT Interface Status details
        # Check for interface data in section_4 first, then fall back to global interfaces
        section_4_interfaces = section_4_data.get('interfaces', [])
        global_interfaces = data.get('interfaces', [])
        interfaces = section_4_interfaces if section_4_interfaces else global_interfaces
        
        if interfaces:
            html_content += "<h3>DUT Interface Status details</h3>"
            html_content += '<table>'
            html_content += '<tr><th>Interfaces</th><th>No. of Ports</th><th>Interface Type</th><th>Interface Name</th></tr>'
            for interface in interfaces:
                if isinstance(interface, dict):
                    # New JSON format: {"interface": "...", "ports": "...", "type": "...", "name": "..."}
                    html_content += f'<tr><td>{interface.get("interface", "")}</td><td>{interface.get("ports", "")}</td><td>{interface.get("type", "")}</td><td>{interface.get("name", "")}</td></tr>'
                elif len(interface) >= 4:
                    # Old array format: [interface, ports, type, name]
                    html_content += f'<tr><td>{interface[0]}</td><td>{interface[1]}</td><td>{interface[2]}</td><td>{interface[3]}</td></tr>'
            html_content += '</table>'
        
        # Add section 4 text-image pairs if available
        section_4_pairs = section_4_data.get('pairs', [])
        if section_4_pairs:
            for pair in section_4_pairs:
                text_content = pair.get('text', '')
                images = pair.get('images', [])
                scripts = pair.get('scripts', [])
                
                if text_content:
                    formatted_text = self.format_text_for_html(text_content)
                    html_content += f'<div class="field"><span class="value">{formatted_text}</span></div>'
                
                # Add scripts if available (like Section 6)
                if scripts:
                    for script_data in scripts:
                        description = script_data.get('description', '')
                        filename = script_data.get('filename', '')
                        
                        if description:
                            html_content += f'<div class="field"><span class="value">{description}</span></div>'
                        
                        if filename:
                            # Use original filename if available, otherwise use the stored filename
                            display_filename = script_data.get('original_filename', filename)
                            html_content += f'<div class="field"><span class="label">Script file:</span> <span class="value">{display_filename}</span></div>'
                
                for image_path in images:
                    if os.path.exists(image_path):
                        try:
                            filename = os.path.basename(image_path)
                            # Convert file path to file:// URL for display
                            file_url = f"file:///{image_path.replace(os.sep, '/')}"
                            html_content += f'''
                            <div class="image-container" style="text-align: center; margin: 20px 0;">
                                <div class="screenshot-info">Figure: {filename}</div>
                                <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                            </div>
                            '''
                        except Exception as e:
                            html_content += f'<div class="field"><span class="placeholder">[Image could not be loaded: {os.path.basename(image_path)}]</span></div>'
                    else:
                        html_content += f'<div class="field"><span class="placeholder">[Image placeholder: {os.path.basename(image_path)}]</span></div>'
        
        # Connection details
        machine_ip = data.get('machine_ip', '')
        target_ip = data.get('target_ip', '')
        if machine_ip and target_ip:
            connection_text = f"DUT is connected to a testing machine (IP address {machine_ip}) via Ethernet LAN1 Connection (IP address {target_ip})."
            html_content += f'<div class="field"><span class="value">{connection_text}</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_4_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        # Add section 4 notes if available
        section_4_notes = section_4_data.get('notes', [])
        if section_4_notes:
            for note_info in section_4_notes:
                note_text = note_info.get('text', '')
                if note_text:
                    html_content += f'<div class="field"><span style="color: #FF0000; font-weight: bold;">Note:</span> <span style="font-weight: bold;">{note_text}</span></div>'
        
        html_content += '</div>'
        
        # Section 5: DUT Configuration
        section_5_data = sections_1_7.get('section_5', {})
        section_5_heading = section_5_data.get('heading', '5. DUT Configuration')
        
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{section_5_heading}</h2>"
        
        # Add section 5 text-image pairs if available
        section_5_pairs = section_5_data.get('pairs', [])
        if section_5_pairs:
            for pair in section_5_pairs:
                text_content = pair.get('text', '')
                images = pair.get('images', [])
                scripts = pair.get('scripts', [])
                
                if text_content:
                    formatted_text = self.format_text_for_html(text_content)
                    html_content += f'<div class="field"><span class="value">{formatted_text}</span></div>'
                
                # Add scripts if available (like Section 6)
                if scripts:
                    for script_data in scripts:
                        description = script_data.get('description', '')
                        filename = script_data.get('filename', '')
                        
                        if description:
                            html_content += f'<div class="field"><span class="value">{description}</span></div>'
                        
                        if filename:
                            # Use original filename if available, otherwise use the stored filename
                            display_filename = script_data.get('original_filename', filename)
                            html_content += f'<div class="field"><span class="label">Script file:</span> <span class="value">{display_filename}</span></div>'
                
                for image_path in images:
                    if os.path.exists(image_path):
                        try:
                            filename = os.path.basename(image_path)
                            # Convert file path to file:// URL for display
                            file_url = f"file:///{image_path.replace(os.sep, '/')}"
                            html_content += f'''
                            <div class="image-container" style="text-align: center; margin: 20px 0;">
                                <div class="screenshot-info">Figure: {filename}</div>
                                <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                            </div>
                            '''
                        except Exception as e:
                            html_content += f'<div class="field"><span class="placeholder">[Image could not be loaded: {os.path.basename(image_path)}]</span></div>'
                    else:
                        html_content += f'<div class="field"><span class="placeholder">[Image placeholder: {os.path.basename(image_path)}]</span></div>'
        else:
            # Fallback to old data structure
            dut_config = section_5_data.get('content', data.get('dut_configuration', ''))
            if dut_config:
                html_content += f'<div class="field"><span class="value">{dut_config}</span></div>'
            else:
                html_content += '<div class="field"><span class="placeholder">[Configuration details to be added]</span></div>'
            html_content += '<div class="field"><span class="placeholder">[Configuration image to be inserted]</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_5_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        # Add section 5 notes if available
        section_5_notes = section_5_data.get('notes', [])
        if section_5_notes:
            for note_info in section_5_notes:
                note_text = note_info.get('text', '')
                if note_text:
                    html_content += f'<div class="field"><span style="color: #FF0000; font-weight: bold;">Note:</span> <span style="font-weight: bold;">{note_text}</span></div>'
        
        html_content += '</div>'
        
        # Section 6: Preconditions
        section_6_data = sections_1_7.get('section_6', {})
        section_6_heading = section_6_data.get('heading', '6. Preconditions')
        
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{section_6_heading}</h2>"
        preconditions = section_6_data.get('content', data.get('preconditions', ''))
        if preconditions:
            formatted_preconditions = self.format_text_for_html(preconditions)
            html_content += f'<div class="field"><span class="value">{formatted_preconditions}</span></div>'

        # Add section 6 images if available
        section_6_images = section_6_data.get('images', [])
        if section_6_images:
            for i, image_info in enumerate(section_6_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                
                if os.path.exists(image_path):
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(image_path)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for section 6 image: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 6.{i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for section 6 image: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 6.{i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Image placeholder: {filename}]</span></div>'
        
        # Add section 6 scripts if available
        section_6_scripts = section_6_data.get('scripts', [])
        if section_6_scripts:
            for script_data in section_6_scripts:
                description = script_data.get('description', '')
                filename = script_data.get('filename', '')
                
                if description:
                    html_content += f'<div class="field"><span class="value">{description}</span></div>'
                
                if filename:
                    # Use original filename if available, otherwise use the stored filename
                    display_filename = script_data.get('original_filename', filename)
                    html_content += f'<div class="field"><span class="label">Script file:</span> <span class="value">{display_filename}</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_6_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        html_content += '</div>'
        
        # Section 7: Test Objective
        section_7_data = sections_1_7.get('section_7', {})
        section_7_heading = section_7_data.get('heading', '7. Test Objective')
        
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{section_7_heading}</h2>"
        test_objective = section_7_data.get('content', data.get('test_objective', ''))
        if test_objective:
            formatted_test_objective = self.format_text_for_html(test_objective)
            html_content += f'<div class="field"><span class="value">{formatted_test_objective}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">[Test objective content to be added]</span></div>'
        
        # Add section 7 images if available
        section_7_images = section_7_data.get('images', [])
        if section_7_images:
            for i, image_info in enumerate(section_7_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                
                if os.path.exists(image_path):
                    # Try to encode the image as base64 for display
                    data_url = self._encode_image_as_base64(image_path)
                    if data_url:
                        print(f"🔧 DEBUG: Generated base64 data URL for section 7 image: {filename}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 7.{i}: {filename}</div>
                            <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                    else:
                        # Fallback to file:// URL
                        file_url = f"file:///{image_path.replace(os.sep, '/')}"
                        print(f"🔧 DEBUG: Generated file URL for section 7 image: {file_url}")
                        html_content += f'''
                        <div class="image-container" style="text-align: center; margin: 20px 0;">
                            <div class="screenshot-info">Figure 7.{i}: {filename}</div>
                            <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                        </div>
                        '''
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Image placeholder: {filename}]</span></div>'
        
        # Add section 7 text-image pairs if available
        section_7_pairs = data.get('section_7_pairs', [])
        if section_7_pairs:
            for pair in section_7_pairs:
                text_content = pair.get('text', '')
                images = pair.get('images', [])
                
                if text_content:
                    html_content += f'<div class="field"><span class="value">{text_content}</span></div>'
                
                for image_path in images:
                    if os.path.exists(image_path):
                        try:
                            filename = os.path.basename(image_path)
                            # Convert file path to file:// URL for display
                            file_url = f"file:///{image_path.replace(os.sep, '/')}"
                            html_content += f'''
                            <div class="image-container" style="text-align: center; margin: 20px 0;">
                                <div class="screenshot-info">Figure: {filename}</div>
                                <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                            </div>
                            '''
                        except Exception as e:
                            html_content += f'<div class="field"><span class="placeholder">[Image could not be loaded: {os.path.basename(image_path)}]</span></div>'
                    else:
                        html_content += f'<div class="field"><span class="placeholder">[Image placeholder: {os.path.basename(image_path)}]</span></div>'
        
        # Add placeholder content if available
        placeholder_content_management = section_7_data.get('placeholder_content_management', [])
        html_content = self.add_placeholder_content_to_html(html_content, placeholder_content_management)
        
        html_content += '</div>'
        
        # Section 8: Test Plan (Page Break)
        html_content += '<div class="major-page-break"></div>'
        html_content += '<div class="section">'
        html_content += "<h2>8. Test Plan</h2>"
        html_content += "<h3>8.0. Test Plan Overview</h3>"
        test_plan_overview = data.get('test_plan_overview', '')
        if test_plan_overview:
            html_content += f'<div class="field"><span class="value">{test_plan_overview}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">[Test Plan Overview to be added]</span></div>'
        
        test_scenarios_heading = data.get('test_scenarios_heading', 'Number of Test Scenarios')
        # For HTML preview, we'll show the heading as is since it's just for preview
        html_content += f"<h3>{test_scenarios_heading}</h3>"
        scenarios = data.get('test_scenarios', [])
        if scenarios:
            html_content += f'<div class="field"><span class="value">Total number of test scenarios: {len(scenarios)}</span></div>'
            for scenario in scenarios:
                scenario_key = scenario.get('key', '')
                scenario_description = scenario.get('description', '')
                if scenario_key and scenario_description:
                    html_content += f'<div class="field"><span class="value">{scenario_key}: {scenario_description}</span></div>'
        else:
            html_content += '<div class="field"><span class="placeholder">Number of test scenarios: [To be determined]</span></div>'
        
        test_bed_diagram_heading = data.get('test_bed_diagram_heading', 'Test Bed Diagram')
        # For HTML preview, we'll show the heading as is since it's just for preview
        html_content += f"<h3>{test_bed_diagram_heading}</h3>"
        
        # Add test bed diagram notes if available
        test_bed_notes = data.get('test_bed_diagram_notes', '')
        if not test_bed_notes and 'test_bed_diagram' in data:
            test_bed_diagram = data['test_bed_diagram']
            if isinstance(test_bed_diagram, dict):
                test_bed_notes = test_bed_diagram.get('notes', '')
        
        if test_bed_notes:
            html_content += f'<div class="field"><span class="content">{test_bed_notes}</span></div>'
        # Get test bed images from multiple sources
        test_bed_images = data.get('test_bed_images', [])
        
        # Also check for images in test_bed_diagram structure
        if not test_bed_images and 'test_bed_diagram' in data:
            test_bed_diagram = data['test_bed_diagram']
            if isinstance(test_bed_diagram, dict) and 'images' in test_bed_diagram:
                test_bed_images = test_bed_diagram['images']
                print(f"🔧 DEBUG: Found {len(test_bed_images)} images in test_bed_diagram structure")
        
        if test_bed_images:
            for i, image_info in enumerate(test_bed_images, 1):
                # Handle both string paths and image info dictionaries
                if isinstance(image_info, dict):
                    image_path = image_info.get('path', '')
                    filename = image_info.get('filename', os.path.basename(image_path))
                    print(f"🔧 DEBUG: Test bed image {i}: path='{image_path}', filename='{filename}', exists={os.path.exists(image_path) if image_path else False}")
                else:
                    image_path = str(image_info)
                    filename = os.path.basename(image_path)
                    print(f"🔧 DEBUG: Test bed image {i}: path='{image_path}', filename='{filename}', exists={os.path.exists(image_path) if image_path else False}")
                
                # Check if this is a placeholder value
                if isinstance(image_path, str) and (image_path.startswith('_') and image_path.endswith('_placeholder') or 'placeholder' in image_path.lower()):
                    # This is a placeholder, show it as placeholder text
                    html_content += f'<div class="field"><span class="placeholder">[{image_path}]</span></div>'
                elif image_path and os.path.exists(image_path):
                    # Check if it's a script file
                    if image_path.lower().endswith(('.py', '.sh', '.bat', '.ps1', '.js', '.vbs')):
                        # This is a script file, show it as script content
                        html_content += f'<div class="field"><span class="value">📄 Script: {filename}</span></div>'
                    else:
                        # This is an image file
                        # Try to encode the image as base64 for display
                        data_url = self._encode_image_as_base64(image_path)
                        if data_url:
                            print(f"🔧 DEBUG: Generated base64 data URL for: {filename}")
                            html_content += f'''
                            <div class="image-container" style="text-align: center; margin: 20px 0;">
                                <div class="screenshot-info">Figure 8.2.{i}: Test Bed Diagram</div>
                                <img src="{data_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                            </div>
                            '''
                        else:
                            # Fallback to file:// URL
                            file_url = f"file:///{image_path.replace(os.sep, '/')}"
                            print(f"🔧 DEBUG: Generated file URL: {file_url}")
                            html_content += f'''
                            <div class="image-container" style="text-align: center; margin: 20px 0;">
                                <div class="screenshot-info">Figure 8.2.{i}: Test Bed Diagram</div>
                                <img src="{file_url}" alt="{filename}" style="max-width: 100%; max-height: 400px; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" />
                            </div>
                            '''
                else:
                    print(f"🔧 DEBUG: Image not found or invalid path: '{image_path}'")
                    html_content += f'<div class="field"><span class="placeholder">[File placeholder 8.2.{i} - Path: {image_path}]</span></div>'
        else:
            # No test bed content available - section will be empty
            pass
        
        tools_required_heading = data.get('tools_required_heading', 'Tools Required')
        # For HTML preview, we'll show the heading as is since it's just for preview
        html_content += f"<h3>{tools_required_heading}</h3>"
        tools = data.get('tools_required', [])
        for tool in tools:
            # Extract tool name from dictionary or use as string
            tool_name = tool.get('name', tool) if isinstance(tool, dict) else tool
            
            # If tool name is empty, check for script descriptions as placeholders
            if not tool_name and isinstance(tool, dict):
                scripts = tool.get('scripts', [])
                if scripts:
                    # Use the first script's description as placeholder
                    first_script = scripts[0]
                    if isinstance(first_script, dict):
                        tool_name = first_script.get('description', '')
            
            if tool_name:
                html_content += f'<div class="field" style="margin: 0; padding: 0;"><span class="value">• {tool_name}</span></div>'
        
        test_execution_steps_heading = data.get('test_execution_steps_heading', 'Test Execution Steps')
        # For HTML preview, we'll show the heading as is since it's just for preview
        html_content += f"<h3>{test_execution_steps_heading}</h3>"
        execution_steps = data.get('execution_steps', [])
        manual_steps = data.get('manual_execution_steps', [])
        
        if execution_steps:
            for scenario_steps in execution_steps:
                steps = scenario_steps.get('steps', [])
                custom_inputs = scenario_steps.get('Custom_input', [])
                note_section = scenario_steps.get('note_section', [])
                
                # Add custom inputs first if they exist
                if custom_inputs:
                    for k, custom_input in enumerate(custom_inputs):
                        # Process custom input for bold formatting in HTML
                        parts = custom_input.split('*')
                        formatted_custom_input = ''
                        for i, part in enumerate(parts):
                            if i % 2 == 1:  # Odd index = between asterisks = bold
                                formatted_custom_input += f'<strong>{part}</strong>'
                            else:
                                formatted_custom_input += part
                        html_content += f'<div class="field"><span class="value">{formatted_custom_input}</span></div>'
                
                # Add scenario key and description if they exist
                scenario_key = scenario_steps.get('scenario_key', '')
                scenario_description = scenario_steps.get('scenario_description', '')
                if scenario_key or scenario_description:
                    if scenario_key:
                        html_content += f'<div class="field"><span class="value" style="font-weight: bold;">{scenario_key}:</span></div>'
                    if scenario_description:
                        # Process description for bold formatting in HTML
                        parts = scenario_description.split('*')
                        formatted_description = ''
                        for i, part in enumerate(parts):
                            if i % 2 == 1:  # Odd index = between asterisks = bold
                                formatted_description += f'<strong>{part}</strong>'
                            else:
                                formatted_description += part
                        html_content += f'<div class="field"><span class="value">{formatted_description}</span></div>'
                    
                    # Add notes after scenario description
                    if note_section:
                        for note in note_section:
                            if note.strip():
                                html_content += f'<div class="field"><span class="value" style="font-style: italic;">Note: {note}</span></div>'
                    
                    html_content += '<br>'
                
                # Only show the actual execution steps, not the scenario description
                for j, step in enumerate(steps):
                    html_content += f'<div class="field"><span class="value">{step}</span></div>'
                
                # Add spacing between scenarios
                html_content += '<br>'
        
        # Add manual execution steps if any
        if manual_steps:
            html_content += '<div class="field"><span class="value">Manual Steps (Not linked to scenarios):</span></div>'
            for j, step in enumerate(manual_steps):
                html_content += f'<div class="field"><span class="value">{step}</span></div>'
        
        if not execution_steps and not manual_steps:
            html_content += '<div class="field"><span class="placeholder">[Test execution steps to be defined]</span></div>'
        html_content += '</div>'
        
        # Section 9: Expected Result
        expected_results_heading = data.get('expected_results_heading', '9. Expected Result')
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{expected_results_heading}</h2>"
        expected_results = data.get('expected_results', '')
        if expected_results:
            html_content += f'<div class="field"><span class="value">{expected_results}</span></div>'
        else:
            if scenarios:
                for i, scenario in enumerate(scenarios):
                    # Remove the "Test Scenario 1.9.4.1" prefix and just show the expected result
                    result_text = "All vulnerabilities are identified with accurate scan results and minimal false positives. No critical issues remain undetected."
                    html_content += f'<div class="field"><span class="value">{result_text}</span></div>'
            else:
                html_content += '<div class="field"><span class="placeholder">[Expected results content to be added]</span></div>'
        html_content += '</div>'
        
        # Section 10: Expected Format of Evidence
        expected_format_evidence_heading = data.get('expected_format_evidence_heading', '10. Expected Format of Evidence')
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += f"<h2>{expected_format_evidence_heading}</h2>"
        
        # Get evidence format from new data format
        evidence_format_list = data.get('evidence_format', [])
        print(f"🔧 DEBUG: HTML Section 10 Evidence - evidence_format_list: {evidence_format_list}")
        print(f"🔧 DEBUG: HTML Section 10 Evidence - evidence_format_list type: {type(evidence_format_list)}")
        print(f"🔧 DEBUG: HTML Section 10 Evidence - evidence_format_list length: {len(evidence_format_list)}")
        
        if evidence_format_list:
            # Use the new evidence format list - remove bullet points
            for evidence_item in evidence_format_list:
                if evidence_item:
                    # Extract evidence text from dictionary or use as string
                    evidence_text = evidence_item.get('evidence_text', evidence_item) if isinstance(evidence_item, dict) else evidence_item
                    if evidence_text:
                        html_content += f'<div class="field"><span class="value">{evidence_text}</span></div>'
        else:
            # Fallback to old format
            evidence_format = data.get('evidence_format', '')
            if evidence_format:
                html_content += f'<div class="field"><span class="value">{evidence_format}</span></div>'
            # No hardcoded fallback - only show content if user has provided it
        html_content += '</div>'
        
        # Section 11: Test Execution (Page Break)
        html_content += '<div class="major-page-break"></div>'
        html_content += '<div class="section">'
        html_content += "<h2>11. Test Execution</h2>"
        execution_blocks = data.get('execution_blocks', [])
        test_execution_cases = data.get('test_execution_cases', [])
        print(f"🔧 DEBUG: HTML Section 11 - execution_blocks: {execution_blocks}")
        print(f"🔧 DEBUG: HTML Section 11 - execution_blocks type: {type(execution_blocks)}")
        print(f"🔧 DEBUG: HTML Section 11 - execution_blocks length: {len(execution_blocks)}")
        print(f"🔧 DEBUG: HTML Section 11 - test_execution_cases: {test_execution_cases}")
        print(f"🔧 DEBUG: HTML Section 11 - test_execution_cases type: {type(test_execution_cases)}")
        print(f"🔧 DEBUG: HTML Section 11 - test_execution_cases length: {len(test_execution_cases)}")
        if execution_blocks:
            for i, block in enumerate(execution_blocks):
                html_content += f'<h3>11.1.{i+1} Test Case Number:</h3>'
                html_content += f'<div class="field"><span class="label">a. Test Case Name:</span></div>'
                html_content += f'<div class="field"><span class="value">{block.get("test_case_name", "[Test Case Name]")}</span></div>'
                html_content += f'<div class="field"><span class="label">b. Test Case Description:</span></div>'
                html_content += f'<div class="field"><span class="value">{block.get("description", "[Test Case Description]")}</span></div>'
                html_content += f'<div class="field"><span class="label">c. Execution Steps:</span></div>'
                steps = block.get('steps', [])
                if isinstance(steps, list):
                    for step in steps:
                        html_content += f'<div class="field"><span class="value">• {step}</span></div>'
                else:
                    html_content += f'<div class="field"><span class="value">{steps}</span></div>'
                html_content += f'<div class="field"><span class="label">d. Test Observations:</span></div>'
                html_content += f'<div class="field"><span class="value">{block.get("observations", "[Test Observations]")}</span></div>'
                html_content += f'<div class="field"><span class="label">e. Evidence Provided:</span></div>'
                html_content += f'<div class="field"><span class="value">{block.get("evidence", "[Evidence Provided]")}</span></div>'
                html_content += '<br>'
        elif test_execution_cases:
            for i, case in enumerate(test_execution_cases):
                print(f"🔧 DEBUG: Processing test_execution_case {i}: {case}")
                html_content += f'<h3>11.1.{i+1} Test Case Number:</h3>'
                html_content += f'<div class="field"><span class="label">a. Test Case Name:</span></div>'
                html_content += f'<div class="field"><span class="value">{case.get("test_case_name", "[Test Case Name]")}</span></div>'
                html_content += f'<div class="field"><span class="label">b. Test Case Description:</span></div>'
                test_case_desc = case.get("test_case_description", "[Test Case Description]")
                print(f"🔧 DEBUG: test_case_description: '{test_case_desc}'")
                # Check if we have imported data from UI fields
                if hasattr(self, 'test_execution_cases_ui_data') and len(self.test_execution_cases_ui_data) > i:
                    ui_case = self.test_execution_cases_ui_data[i]
                    if 'case_desc_edit' in ui_case and ui_case['case_desc_edit']:
                        ui_description = ui_case['case_desc_edit'].toPlainText()
                        if ui_description.strip():
                            test_case_desc = ui_description
                
                if test_case_desc and test_case_desc != "[Test Case Description]":
                    html_content += f'<div class="field"><span class="value">{test_case_desc}</span></div>'
                else:
                    html_content += '<div class="field"><span class="placeholder">[Test Case Description]</span></div>'
                html_content += f'<div class="field"><span class="label">c. Execution Steps:</span></div>'
                steps_data = case.get('steps_data', [])
                if steps_data:
                    for step in steps_data:
                        step_text = step.get('step_text', step) if isinstance(step, dict) else step
                        html_content += f'<div class="field"><span class="value">• {step_text}</span></div>'
                else:
                    html_content += f'<div class="field"><span class="placeholder">[Execution Steps]</span></div>'
                html_content += f'<div class="field"><span class="label">d. Test Observations:</span></div>'
                test_observations = case.get("test_observations", "[Test Observations]")
                print(f"🔧 DEBUG: test_observations: '{test_observations}'")
                # Check if we have imported data from UI fields
                if hasattr(self, 'test_execution_cases_ui_data') and len(self.test_execution_cases_ui_data) > i:
                    ui_case = self.test_execution_cases_ui_data[i]
                    if 'observations_edit' in ui_case and ui_case['observations_edit']:
                        ui_observations = ui_case['observations_edit'].toPlainText()
                        if ui_observations.strip():
                            test_observations = ui_observations
                
                if test_observations and test_observations != "[Test Observations]":
                    html_content += f'<div class="field"><span class="value">{test_observations}</span></div>'
                else:
                    html_content += '<div class="field"><span class="placeholder">[Test Observations]</span></div>'
                html_content += f'<div class="field"><span class="label">e. Evidence Provided:</span></div>'
                evidence_provided = case.get("evidence_provided", "[Evidence Provided]")
                print(f"🔧 DEBUG: evidence_provided: '{evidence_provided}'")
                # Check if we have imported data from UI fields
                if hasattr(self, 'test_execution_cases_ui_data') and len(self.test_execution_cases_ui_data) > i:
                    ui_case = self.test_execution_cases_ui_data[i]
                    if 'evidence_edit' in ui_case and ui_case['evidence_edit']:
                        ui_evidence = ui_case['evidence_edit'].toPlainText()
                        if ui_evidence.strip():
                            evidence_provided = ui_evidence
                
                if evidence_provided and evidence_provided != "[Evidence Provided]":
                    html_content += f'<div class="field"><span class="value">{evidence_provided}</span></div>'
                else:
                    html_content += '<div class="field"><span class="placeholder">[Evidence Provided]</span></div>'
                html_content += '<br>'
        elif scenarios:
            for i, scenario in enumerate(scenarios):
                html_content += f'<h3>11.1.{i+1} Test Case Number:</h3>'
                html_content += f'<div class="field"><span class="label">a. Test Case Name:</span></div>'
                html_content += f'<div class="field"><span class="value">{scenario.get("test_case_name", "[Test Case Name]")}</span></div>'
                html_content += f'<div class="field"><span class="label">b. Test Case Description:</span></div>'
                html_content += f'<div class="field"><span class="value">{scenario.get("description", "[Test Case Description]")}</span></div>'
                html_content += f'<div class="field"><span class="label">c. Execution Steps:</span></div>'
                steps = scenario.get('steps', [])
                for step in steps:
                    html_content += f'<div class="field"><span class="value">• {step}</span></div>'
                html_content += f'<div class="field"><span class="label">d. Test Observations:</span></div>'
                html_content += f'<div class="field"><span class="placeholder">[Test Observations to be filled during testing]</span></div>'
                html_content += f'<div class="field"><span class="label">e. Evidence Provided:</span></div>'
                html_content += f'<div class="field"><span class="placeholder">[Screenshots and evidence to be inserted]</span></div>'
                html_content += '<br>'
        else:
            html_content += '<h3>11.1.1 Test Case Number:</h3>'
            html_content += f'<div class="field"><span class="label">a. Test Case Name:</span></div>'
            html_content += f'<div class="field"><span class="placeholder">[Test Case Name]</span></div>'
            html_content += f'<div class="field"><span class="label">b. Test Case Description:</span></div>'
            html_content += f'<div class="field"><span class="placeholder">[Test Case Description]</span></div>'
            html_content += f'<div class="field"><span class="label">c. Execution Steps:</span></div>'
            html_content += f'<div class="field"><span class="placeholder">[Execution Steps]</span></div>'
            html_content += f'<div class="field"><span class="label">d. Test Observations:</span></div>'
            html_content += f'<div class="field"><span class="placeholder">[Test Observations]</span></div>'
            html_content += f'<div class="field"><span class="label">e. Evidence Provided:</span></div>'
            html_content += f'<div class="field"><span class="placeholder">[Evidence Provided]</span></div>'
        html_content += '</div>'
        
        # Section 12: Test Case Result
        html_content += '<div class="section-break"></div>'
        html_content += '<div class="section">'
        html_content += "<h2>12. Test Case Result</h2>"
        html_content += '<table>'
        html_content += '<tr><th>S. No</th><th>TEST CASE No.</th><th>PASS FAIL</th><th>Remarks</th></tr>'
        
        # Get test case results from new data format
        test_case_results = data.get('test_case_results', [])
        
        if test_case_results:
            # Use the new test case results format
            for result_data in test_case_results:
                serial_no = result_data.get('serial_no', '')
                test_case_no = result_data.get('test_case_no', '')
                pass_fail = result_data.get('pass_fail', 'NA')
                remarks = result_data.get('remarks', '')
                
                html_content += f'<tr><td>{serial_no}</td><td>{test_case_no}</td><td>{pass_fail}</td><td>{remarks}</td></tr>'
        else:
            # Fallback to old format
            if execution_blocks:
                for i, block in enumerate(execution_blocks):
                    html_content += f'<tr><td>{i+1}</td><td>1.9.4.{i+1}</td><td>[PASS/FAIL]</td><td>[Remarks]</td></tr>'
            elif scenarios:
                for i, scenario in enumerate(scenarios):
                    html_content += f'<tr><td>{i+1}</td><td>1.9.4.{i+1}</td><td>[PASS/FAIL]</td><td>[Remarks]</td></tr>'
            else:
                html_content += '<tr><td>1</td><td>1.9.4.1</td><td>[PASS/FAIL]</td><td>[Remarks]</td></tr>'
        
        html_content += '</table>'
        html_content += '</div>'
        
        # Footer
        html_content += f'<div style="text-align: center; margin-top: 30px; color: #666666; font-size: 10px;">Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} | DELTAPHI LABS | Confidential</div>'
        
        html_content += "</body></html>"
        return html_content
    
    def add_placeholder_content_to_html(self, html_content, placeholder_content_management):
        """Add placeholder content to HTML with center alignment"""
        if not placeholder_content_management:
            return html_content
            
        for placeholder_item in placeholder_content_management:
            placeholder_text = placeholder_item.get('placeholder', '')
            if placeholder_text:
                html_content += f'<div class="placeholder-content">{placeholder_text}</div>'
        
        return html_content
        
    def format_field_html(self, label, value):
        """Format field for HTML preview"""
        display_value = value if value else "[To be filled]"
        return f'<div class="field"><span class="label">{label}:</span> <span class="value">{display_value}</span></div>'

    @staticmethod
    def resolve_steps_data_filenames(test_execution_cases, config_data):
        """Resolve steps_data filenames to full objects using global arrays"""
        print(f"🔍 DEBUG 11: resolve_steps_data_filenames called with {len(test_execution_cases) if test_execution_cases else 0} cases")
        if not test_execution_cases:
            print(f"🔍 DEBUG 11: No test execution cases to resolve")
            return test_execution_cases

        # Extract global arrays from configuration
        configuration = config_data.get('configuration', {})
        global_images = configuration.get('images', [])
        global_scripts = configuration.get('scripts', [])

        # Debug logging
        print(f"DEBUG: resolve_steps_data_filenames - global_images count: {len(global_images)}")
        print(f"DEBUG: resolve_steps_data_filenames - global_scripts count: {len(global_scripts)}")
        
        # Debug: Check for non-string filenames
        for i, img in enumerate(global_images):
            if isinstance(img, dict) and "filename" in img:
                if not isinstance(img["filename"], str):
                    print(f"DEBUG: resolve_steps_data_filenames - Found non-string filename in image {i}: {type(img['filename'])} = {img['filename']}")
        
        for i, script in enumerate(global_scripts):
            if isinstance(script, dict) and "filename" in script:
                if not isinstance(script["filename"], str):
                    print(f"DEBUG: resolve_steps_data_filenames - Found non-string filename in script {i}: {type(script['filename'])} = {script['filename']}")

        # Create lookup dictionaries for cross-referencing
        try:
            image_lookup = {img["filename"]: img for img in global_images if isinstance(img, dict) and isinstance(img.get("filename"), str)}
            script_lookup = {script["filename"]: script for script in global_scripts if isinstance(script, dict) and isinstance(script.get("filename"), str)}
        except TypeError as e:
            print(f"DEBUG: resolve_steps_data_filenames - Error creating lookup dictionaries: {e}")
            print(f"DEBUG: resolve_steps_data_filenames - global_images types: {[type(img.get('filename')) if isinstance(img, dict) else type(img) for img in global_images]}")
            print(f"DEBUG: resolve_steps_data_filenames - global_scripts types: {[type(script.get('filename')) if isinstance(script, dict) else type(script) for script in global_scripts]}")
            # Fallback to empty dictionaries
            image_lookup = {}
            script_lookup = {}

        # Process each test execution case
        resolved_cases = []
        for case in test_execution_cases:
            resolved_case = case.copy()

            # Process steps_data for this case
            if 'steps_data' in case:
                resolved_steps_data = []
                for step_data in case['steps_data']:
                    resolved_step = step_data.copy()

                    # Resolve image filenames to full image objects
                    resolved_images = []
                    for filename in step_data.get('images', []):
                        if isinstance(filename, str) and filename in image_lookup:
                            resolved_images.append(image_lookup[filename])
                        elif isinstance(filename, dict):
                            # If filename is a dict, it might already be a full object
                            resolved_images.append(filename)
                    resolved_step['resolved_images'] = resolved_images

                    # Resolve script filenames to full script objects
                    resolved_scripts = []
                    for filename in step_data.get('scripts', []):
                        if isinstance(filename, str) and filename in script_lookup:
                            resolved_scripts.append(script_lookup[filename])
                        elif isinstance(filename, dict):
                            # If filename is a dict, it might already be a full object
                            resolved_scripts.append(filename)
                    
                    # IMPORTANT: Also check for scripts stored in upload_scripts_list format
                    # This is where the actual script information with placeholders is stored
                    if 'upload_scripts_list' in step_data:
                        print(f"🔍 DEBUG 11: Found upload_scripts_list with {len(step_data['upload_scripts_list'])} scripts for step {step_data.get('step_index', 'unknown')}")
                        for script_info in step_data['upload_scripts_list']:
                            if isinstance(script_info, dict):
                                # Ensure the script has the description (placeholder) information
                                if 'description' in script_info and script_info['description']:
                                    resolved_scripts.append(script_info)
                                    print(f"🔍 DEBUG 11: Added script with placeholder '{script_info['description']}' to resolved_scripts")
                                else:
                                    # If no description, still add the script but log a warning
                                    resolved_scripts.append(script_info)
                                    print(f"🔍 DEBUG 11: Added script without placeholder to resolved_scripts: {script_info.get('filename', 'unknown')}")
                    
                    resolved_step['resolved_scripts'] = resolved_scripts

                    # Preserve placeholders from step_data
                    if 'placeholders' in step_data:
                        resolved_step['placeholders'] = step_data['placeholders']
                        print(f"🔍 DEBUG 11: Preserved {len(step_data['placeholders'])} placeholders for step {step_data.get('step_index', 'unknown')}")

                    resolved_steps_data.append(resolved_step)

                resolved_case['steps_data'] = resolved_steps_data

            resolved_cases.append(resolved_case)

        return resolved_cases

    @staticmethod
    def resolve_test_bed_files(test_bed_files, config_data):
        """Resolve test bed files (images/scripts) filenames to full objects using global arrays"""
        if not test_bed_files:
            return test_bed_files

        # Extract global arrays from configuration
        configuration = config_data.get('configuration', {})
        global_images = configuration.get('images', [])
        global_scripts = configuration.get('scripts', [])

        # Debug logging
        print(f"DEBUG: resolve_test_bed_files - global_images count: {len(global_images)}")
        print(f"DEBUG: resolve_test_bed_files - global_scripts count: {len(global_scripts)}")
        
        # Debug: Check for non-string filenames
        for i, img in enumerate(global_images):
            if isinstance(img, dict) and "filename" in img:
                if not isinstance(img["filename"], str):
                    print(f"DEBUG: resolve_test_bed_files - Found non-string filename in image {i}: {type(img['filename'])} = {img['filename']}")
        
        for i, script in enumerate(global_scripts):
            if isinstance(script, dict) and "filename" in script:
                if not isinstance(script["filename"], str):
                    print(f"DEBUG: resolve_test_bed_files - Found non-string filename in script {i}: {type(script['filename'])} = {script['filename']}")

        # Create lookup dictionaries for cross-referencing
        try:
            image_lookup = {img["filename"]: img for img in global_images if isinstance(img, dict) and isinstance(img.get("filename"), str)}
            script_lookup = {script["filename"]: script for script in global_scripts if isinstance(script, dict) and isinstance(script.get("filename"), str)}
        except TypeError as e:
            print(f"DEBUG: resolve_test_bed_files - Error creating lookup dictionaries: {e}")
            print(f"DEBUG: resolve_test_bed_files - global_images types: {[type(img.get('filename')) if isinstance(img, dict) else type(img) for img in global_images]}")
            print(f"DEBUG: resolve_test_bed_files - global_scripts types: {[type(script.get('filename')) if isinstance(script, dict) else type(script) for script in global_scripts]}")
            # Fallback to empty dictionaries
            image_lookup = {}
            script_lookup = {}

        # Process each test bed file
        resolved_files = []
        for file_item in test_bed_files:
            if isinstance(file_item, dict):
                # This is already a script object with path/filename
                if 'filename' in file_item and file_item['filename'] in script_lookup:
                    # It's a script, get the full object
                    resolved_script = script_lookup[file_item['filename']].copy()
                    # Keep any additional properties like content
                    if 'content' in file_item:
                        resolved_script['content'] = file_item['content']
                    resolved_files.append(resolved_script)
                else:
                    resolved_files.append(file_item)
            elif isinstance(file_item, str):
                # This is a filename string
                if file_item in image_lookup:
                    resolved_files.append(image_lookup[file_item])
                elif file_item in script_lookup:
                    resolved_files.append(script_lookup[file_item])
                else:
                    # Keep as string if not found (placeholder or old format)
                    resolved_files.append(file_item)
            elif isinstance(file_item, dict):
                # This might be a full object already
                resolved_files.append(file_item)
            else:
                resolved_files.append(file_item)

        return resolved_files


class EnhancedDocumentManager:
    """Manager class for enhanced document generation"""
    
    @staticmethod
    def generate_enhanced_report(data, output_path):
        """Generate enhanced security assessment report"""
        try:
            # Check if template path is provided
            template_path = data.get('template_path', '')
            if template_path and os.path.exists(template_path):
                # Use template-based generation
                return EnhancedDocumentManager.generate_enhanced_report_with_template(data, output_path, template_path)
            else:
                # Use standard generation
                enhanced_data = EnhancedDocumentManager.prepare_enhanced_data(data)
                generator = EnhancedDocumentGenerator()
                generator.generate_docx(enhanced_data, output_path)
                return True, f"Enhanced report generated successfully: {output_path}"
        except Exception as e:
            import traceback
            print(f"DEBUG: Enhanced report generation error: {str(e)}")
            print(f"DEBUG: Error traceback:")
            traceback.print_exc()
            return False, f"Error generating enhanced report: {str(e)}"
    
    @staticmethod
    def generate_enhanced_report_with_template(data, output_path, template_path):
        """Generate enhanced security assessment report using a template file"""
        try:
            # Load the template document
            template_doc = Document(template_path)
            
            # Create a new document based on the template
            doc = Document(template_path)
            
            # Prepare the data
            enhanced_data = EnhancedDocumentManager.prepare_enhanced_data(data)
            
            # The template document contains headers, footers, and watermarks
            # We want to preserve these but start our content from page 1
            
            # Preserve watermarks from template
            EnhancedDocumentManager.preserve_template_watermarks(template_doc, doc)
            
            # Apply configurable header if specified
            header_text = data.get('header_text', '')
            header_start_page = data.get('header_start_page', 0)
            if header_text:
                EnhancedDocumentManager.apply_configurable_header(doc, header_text, header_start_page)
            
            # Clear all existing content but keep the template structure (headers, footers, watermarks)
            # We'll remove all paragraphs but keep the document structure
            if len(doc.paragraphs) > 0:
                # Remove all existing paragraphs
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            
            # Instead of creating a temporary document, add content directly to the main document
            # This ensures images are properly embedded in the final document
            generator = EnhancedDocumentGenerator()
            generator.doc = doc  # Use the main document directly
            
            # Reset section counter for dynamic numbering
            generator.reset_section_counter()
            
            # Add content sections directly to the main document
            generator.add_header(enhanced_data)  # Add header with title
            generator.add_section_1_itsar(enhanced_data)
            generator.add_section_2_security_requirement(enhanced_data)
            generator.add_section_3_requirement_description(enhanced_data)
            generator.add_section_4_dut_confirmation(enhanced_data)
            generator.add_section_5_dut_configuration(enhanced_data)
            generator.add_section_6_preconditions(enhanced_data)
            generator.add_section_7_test_objective(enhanced_data)
            generator.add_section_8_test_plan(enhanced_data)
            generator.add_section_9_expected_result(enhanced_data)
            generator.add_section_10_evidence_format(enhanced_data)
            generator.add_section_11_test_execution(enhanced_data)
            generator.add_section_12_test_case_result(enhanced_data)
            generator.add_footer()  # Add footer
            
            # Save the document
            doc.save(output_path)
            return True, f"Enhanced report generated with template successfully: {output_path}"
            
        except Exception as e:
            import traceback
            print(f"DEBUG: Enhanced report generation with template error: {str(e)}")
            print(f"DEBUG: Error traceback:")
            traceback.print_exc()
            return False, f"Error generating enhanced report with template: {str(e)}"
    
    @staticmethod
    def preserve_template_watermarks(template_doc, target_doc):
        """Preserve watermarks from template document"""
        try:
            # Copy watermark settings from template to target document
            if len(template_doc.sections) > 0 and len(target_doc.sections) > 0:
                template_section = template_doc.sections[0]
                target_section = target_doc.sections[0]
                
                # Copy background settings that might contain watermarks
                if hasattr(template_section, '_sectPr') and hasattr(target_section, '_sectPr'):
                    # Copy background element if it exists
                    background = template_section._sectPr.find('.//w:background', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if background is not None:
                        target_background = target_section._sectPr.find('.//w:background', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if target_background is not None:
                            target_background.getparent().remove(target_background)
                        target_section._sectPr.append(background)
                
                # Copy any drawing elements that might be watermarks
                for paragraph in template_doc.paragraphs:
                    for run in paragraph.runs:
                        if run._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            # This might be a watermark, preserve it
                            pass
                            
        except Exception as e:
            # If watermark preservation fails, continue without it
            print(f"Warning: Could not preserve watermarks: {e}")
            pass
    
    @staticmethod
    def generate_enhanced_preview(data):
        """Generate enhanced HTML preview"""
        generator = EnhancedDocumentGenerator()
        return generator.generate_preview(data)

    @staticmethod
    def prepare_enhanced_data(config_data):
        """Prepare data for enhanced document generation"""
        # Debug logging
        print(f"DEBUG: prepare_enhanced_data - config_data keys: {list(config_data.keys())}")
        if 'configuration' in config_data:
            config = config_data['configuration']
            print(f"DEBUG: prepare_enhanced_data - configuration keys: {list(config.keys())}")
            if 'images' in config:
                print(f"DEBUG: prepare_enhanced_data - images count: {len(config['images'])}")
                for i, img in enumerate(config['images']):
                    if isinstance(img, dict) and 'filename' in img:
                        print(f"DEBUG: prepare_enhanced_data - image {i} filename type: {type(img['filename'])} = {img['filename']}")
            if 'scripts' in config:
                print(f"DEBUG: prepare_enhanced_data - scripts count: {len(config['scripts'])}")
                for i, script in enumerate(config['scripts']):
                    if isinstance(script, dict) and 'filename' in script:
                        print(f"DEBUG: prepare_enhanced_data - script {i} filename type: {type(script['filename'])} = {script['filename']}")
        
        # Get the user input title and format it properly
        user_title = config_data.get('report_title', '').strip()
        if user_title:
            # If user provided a title, use it with proper formatting
            test_report = f"TEST REPORT FOR: {user_title}"
        else:
            # Only use default if no user input
            test_report = 'TEST REPORT FOR: [Test Case Title]'
        
        # Extract configuration data if it exists
        config = config_data.get('configuration', {})
        print(f"🔍 DEBUG 11: config_data keys: {list(config_data.keys())}")
        print(f"🔍 DEBUG 11: config keys: {list(config.keys())}")
        print(f"🔍 DEBUG 11: test_execution_cases in config: {'test_execution_cases' in config}")
        if 'test_execution_cases' in config:
            print(f"🔍 DEBUG 11: test_execution_cases count: {len(config['test_execution_cases'])}")
            for i, case in enumerate(config['test_execution_cases']):
                print(f"🔍 DEBUG 11: Case {i}: {case.get('case_number', 'unknown')} with {len(case.get('steps_data', []))} steps_data")
        
        # Determine number of scenarios for dynamic result/remarks generation
        test_execution_cases = config.get('test_execution_cases', []) or config_data.get('test_execution_cases', [])
        execution_blocks = config_data.get('execution_blocks', [])
        test_scenarios = config_data.get('test_scenarios', []) or config.get('test_scenarios', [])
        
        num_scenarios = 0
        if test_execution_cases:
            num_scenarios = len(test_execution_cases)
        elif execution_blocks:
            num_scenarios = len(execution_blocks)
        elif test_scenarios:
            num_scenarios = len(test_scenarios)
        else:
            num_scenarios = 1  # Default to 1 if no scenarios found
        
        # Generate dynamic result and remarks fields
        dynamic_results = {}
        for i in range(1, num_scenarios + 1):
            result_key = f"result_{i}"
            remarks_key = f"remarks_{i}"
            
            # Get from config_data first, then from config, then default to "none"
            dynamic_results[result_key] = config_data.get(result_key, config.get(result_key, "none"))
            dynamic_results[remarks_key] = config_data.get(remarks_key, config.get(remarks_key, "none"))
        
        enhanced_data = {
            'test_report': test_report,
            'itsar_section': config_data.get('itsar_section', 'Section 1.9 Vulnerability Testing Requirements'),
            'security_req': config_data.get('security_req', '1.9.4: Vulnerability Scanning'),
            'req_description': config_data.get('req_description', 'The system shall be tested for known vulnerabilities using automated scanning tools.'),
            'header_text': config_data.get('header_text', ''),
            'header_start_page': config_data.get('header_start_page', 0),
            'dut_configuration': config_data.get('dut_configuration', ''),
            'preconditions': config_data.get('preconditions', ''),
            'test_objective': config_data.get('test_objective', 'Identify and document all security vulnerabilities in the system.'),
            'expected_results': config_data.get('expected_results', 'All vulnerabilities identified and documented with remediation plans.'),
            'evidence_format': config_data.get('evidence_format', 'Screenshots of scanning tools and results.'),
            'screenshots': config_data.get('screenshots', []),
            'interfaces': config_data.get('interfaces', []),
            'dut_fields': config.get('dut_fields', []) or config_data.get('dut_fields', []),  # Extract from configuration or config_data directly
            'hash_sections': config.get('hash_sections', []) or config_data.get('hash_sections', []),  # Extract from configuration or config_data directly
            'itsar_fields': config.get('itsar_fields', []) or config_data.get('itsar_fields', []), # Extract from configuration or config_data directly
            'machine_ip': config.get('machine_ip', ''),
            'target_ip': config.get('target_ip', ''),
            'ssh_username': config.get('ssh_username', 'admin'),
            'ssh_password': config.get('ssh_password', ''),
            'tools_required': config_data.get('tools_required', []) or config.get('tools_required', []),  # Extract from config_data directly or configuration as fallback
            'test_plan_overview': config_data.get('test_plan_overview', '') or config.get('test_plan_overview', ''),
            'test_scenarios': config_data.get('test_scenarios', []) or config.get('test_scenarios', []),  # Extract from config_data directly or configuration as fallback
            'execution_steps': config_data.get('execution_steps', []) or config.get('execution_steps', []),  # Extract from config_data directly or configuration as fallback
            'manual_execution_steps': config_data.get('manual_execution_steps', []) or config.get('manual_execution_steps', []),  # Extract from config_data directly or configuration as fallback
            'test_bed_images': config_data.get('test_bed_images', []) or config.get('test_bed_images', []),  # Extract from config_data directly or configuration as fallback
            'test_bed_scripts': config_data.get('test_bed_scripts', []) or config.get('test_bed_scripts', []),  # Extract from config_data directly or configuration as fallback
            'test_execution_cases': config.get('test_execution_cases', []) or config_data.get('test_execution_cases', []), # Extract from configuration or root
            'step_images': config.get('step_images', []), # Extract from configuration

            # Add test_bed_diagram data from configuration
            'test_bed_diagram': config_data.get('test_bed_diagram', {}) or config.get('test_bed_diagram', {}),  # Extract from config_data directly or configuration as fallback

            # Resolve steps_data filenames to full objects for document generation
            'test_execution_cases_resolved': EnhancedDocumentGenerator.resolve_steps_data_filenames(config.get('test_execution_cases', []) or config_data.get('test_execution_cases', []), config_data),
            'resolved_test_bed_images': EnhancedDocumentGenerator.resolve_test_bed_files(config.get('test_bed_images', []), config_data),
            'resolved_test_bed_scripts': EnhancedDocumentGenerator.resolve_test_bed_files(config.get('test_bed_scripts', []), config_data),
            'sections_1_7': config_data.get('sections_1_7', {}) or config.get('sections_1_7', {}), # Extract from config_data directly or configuration as fallback
            'expected_results': config_data.get('expected_results', []) or config.get('expected_results', []), # Extract from config_data directly or configuration as fallback
            'evidence_format': config_data.get('evidence_format', []) or config.get('evidence_format', []), # Extract from config_data directly or configuration as fallback
            'test_case_results': config_data.get('test_case_results', []) or config.get('test_case_results', []), # Extract from config_data first, then configuration as fallback
            'test_scenarios_heading': config.get('test_scenarios_heading', 'Number of Test Scenarios'),
            'test_bed_diagram_heading': config.get('test_bed_diagram_heading', 'Test Bed Diagram'),
            'test_bed_diagram_notes': config.get('test_bed_diagram_notes', ''),
            'tools_required_heading': config.get('tools_required_heading', 'Tools Required'),
            'test_execution_steps_heading': config.get('test_execution_steps_heading', 'Test Execution Steps'),
            'expected_results_heading': config.get('expected_results_heading', '9. Expected Result'),
            'expected_format_evidence_heading': config.get('expected_format_evidence_heading', '10. Expected Format of Evidence'),
            'scenario_placeholders': config_data.get('scenario_placeholders', {}),  # Add scenario placeholders for DUT placeholders
            
            # Add dynamic result and remarks fields
            **dynamic_results
        }
        
        print(f"🔍 DEBUG 11: enhanced_data test_execution_cases count: {len(enhanced_data.get('test_execution_cases', []))}")
        print(f"🔍 DEBUG 11: enhanced_data test_execution_cases_resolved count: {len(enhanced_data.get('test_execution_cases_resolved', []))}")
        
        return enhanced_data
    
    @staticmethod
    def apply_configurable_header(doc, header_text, header_start_page):
        """Apply configurable header text to document starting from specified page number"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            # Get the first section of the document
            if len(doc.sections) == 0:
                print("Warning: No sections found in document")
                return
            
            section = doc.sections[0]
            
            # DON'T clear existing header content - preserve images and watermarks
            # Instead, add the configurable header text to the left side
            
            # Apply header logic based on start page
            # Note: Page 1 is the first page (without visible page number)
            # Page 0 = All pages, Page 1+ = Start from that page number
            
            if header_start_page == 0 or header_start_page == 1:
                # Apply header to all pages (0 = all pages, 1 = from first page which means all pages)
                
                # Apply to regular header (for pages 2+)
                header = section.header
                EnhancedDocumentManager.add_header_text_to_existing_header(header, header_text)
                
                # Also apply to first page header (for page 1)
                first_page_header = section.first_page_header
                if first_page_header is not None:
                    EnhancedDocumentManager.add_header_text_to_existing_header(first_page_header, header_text)
                
                print(f"Applied configurable header: '{header_text}' to all pages including first page (preserving existing header content)")
            else:
                # For pages starting from a specific page number (2+), we need to create sections
                # This is more complex and requires page breaks and section management
                # For now, we'll apply the header to all pages and add a note
                # In a full implementation, you would need to:
                # 1. Create page breaks before the start page
                # 2. Create a new section starting from the specified page
                # 3. Apply the header only to that section
                
                # Apply to regular header (for pages 2+)
                header = section.header
                EnhancedDocumentManager.add_header_text_to_existing_header(header, header_text)
                
                # Also apply to first page header (for page 1) if start page is 1
                if header_start_page == 1:
                    first_page_header = section.first_page_header
                    if first_page_header is not None:
                        EnhancedDocumentManager.add_header_text_to_existing_header(first_page_header, header_text)
                
                print(f"Applied configurable header: '{header_text}' starting from page {header_start_page} (Note: Currently applied to all pages. Page-specific logic requires advanced section management)")
            
        except Exception as e:
            print(f"Error applying configurable header: {e}")
            import traceback
            traceback.print_exc()
    
    
    @staticmethod
    def add_header_text_to_existing_header(header, header_text):
        """Add header text to existing header while preserving existing content"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            # Check if header already has content
            if len(header.paragraphs) > 0:
                # Find the first paragraph that doesn't have content or has minimal content
                target_paragraph = None
                for paragraph in header.paragraphs:
                    if not paragraph.text.strip():
                        target_paragraph = paragraph
                        break
                
                # If no empty paragraph found, create a new one
                if target_paragraph is None:
                    target_paragraph = header.add_paragraph()
                
                # Add the header text to the left side
                target_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add the header text with specified formatting
                run = target_paragraph.add_run(header_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.font.bold = True
                # Set the font for East Asian characters as well
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                
            else:
                # No existing paragraphs, create a new one
                header_paragraph = header.add_paragraph()
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add the header text with specified formatting
                run = header_paragraph.add_run(header_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.font.bold = True
                # Set the font for East Asian characters as well
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                
        except Exception as e:
            print(f"Error adding header text to existing header: {e}")
            import traceback
            traceback.print_exc()