#!/usr/bin/env python3
"""
Security Assessment Report Generator - CLI Version
A command-line interface for generating security assessment reports when GUI is not available
"""

import os
import sys
import json
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import screenshot functionality for CLI
try:
    from screenshot_annotation_tool import ScreenshotCaptureTool
    SCREENSHOT_AVAILABLE = True
except ImportError:
    SCREENSHOT_AVAILABLE = False
    print("Warning: Screenshot annotation tool not available. Basic screenshot capture will be used.")

# Import test scenarios functionality for CLI
try:
    from test_scenarios_manager import TestScenariosManager
    TEST_SCENARIOS_AVAILABLE = True
except ImportError:
    TEST_SCENARIOS_AVAILABLE = False
    print("Warning: Test scenarios manager not available. Basic test scenarios will be used.")


# Import enhanced document generator for CLI
try:
    from enhanced_document_generator import EnhancedDocumentManager
    ENHANCED_DOCUMENT_AVAILABLE = True
except ImportError:
    ENHANCED_DOCUMENT_AVAILABLE = False
    print("Warning: Enhanced document generator not available. Basic document generation will be used.")

# Configuration file path
CONFIG_FILE_PATH = "security_report_config.json"

class SecurityReportGeneratorCLI:
    """Command-line version of the Security Report Generator"""
    
    def __init__(self):
        self.config = {}
        self.screenshots = []
        self.load_config()
    
    def load_config(self):
        """Load configuration from file"""
        if os.path.exists(CONFIG_FILE_PATH):
            try:
                with open(CONFIG_FILE_PATH, 'r') as f:
                    self.config = json.load(f)
                print(f"Configuration loaded from {CONFIG_FILE_PATH}")
            except Exception as e:
                print(f"Warning: Could not load configuration: {e}")
                self.config = {}
        else:
            self.config = {}
    
    def save_config(self):
        """Save configuration to file"""
        try:
            with open(CONFIG_FILE_PATH, 'w') as f:
                json.dump(self.config, f, indent=2)
            print(f"Configuration saved to {CONFIG_FILE_PATH}")
        except Exception as e:
            print(f"Error saving configuration: {e}")
    
    def prompt_input(self, prompt, default="", required=False):
        """Prompt user for input with optional default"""
        if default:
            user_input = input(f"{prompt} [{default}]: ").strip()
            return user_input if user_input else default
        else:
            while True:
                user_input = input(f"{prompt}: ").strip()
                if user_input or not required:
                    return user_input
                print("This field is required.")
    
    def collect_report_info(self):
        """Collect basic report information from user"""
        print("\n=== Report Information ===")
        
        self.config['report_title'] = self.prompt_input(
            "Test Case Title", 
            self.config.get('report_title', ''),
            required=True
        )
        
        self.config['report_number'] = self.prompt_input(
            "Report Number (e.g., DLPL-JC24-1128-01-1.4.1-V1)", 
            self.config.get('report_number', ''),
            required=True
        )
        
        self.config['itsar_section'] = self.prompt_input(
            "ITSAR Section (e.g., Section 1.4: System Secure Execution Environment)", 
            self.config.get('itsar_section', '')
        )
        
        self.config['requirement_number'] = self.prompt_input(
            "Security Requirement (e.g., 1.4.1: No unused functions)", 
            self.config.get('requirement_number', '')
        )
        
        self.config['test_environment'] = self.prompt_input(
            "Test Environment", 
            self.config.get('test_environment', 'Kali Linux 2024.4')
        )
        
        self.config['test_date'] = self.prompt_input(
            "Test Date", 
            self.config.get('test_date', datetime.now().strftime('%Y-%m-%d'))
        )
    
    def collect_dut_info(self):
        """Collect Device Under Test information"""
        print("\n=== DUT Configuration ===")
        
        # Initialize DUT fields list
        self.config['dut_fields'] = []
        
        print("Enter DUT fields (label and value pairs). Enter empty label to finish:")
        
        while True:
            label = input("Field label (or empty to finish): ").strip()
            if not label:
                break
                
            value = input(f"Value for '{label}': ").strip()
            
            self.config['dut_fields'].append({
                'label': label,
                'value': value
            })
            
            print(f"Added: {label} = {value}")
        
        print(f"Total DUT fields added: {len(self.config['dut_fields'])}")
        
        # Keep IP address for backward compatibility
        self.config['dut_ip'] = self.prompt_input(
            "DUT IP Address", 
            self.config.get('dut_ip', '192.168.1.2')
        )
    
    def collect_hash_info(self):
        """Collect dynamic hash information"""
        print("\n=== Hash Sections ===")
        
        # Initialize hash fields list
        self.config['hash_fields'] = []
        
        # Ask for hash section heading
        hash_heading = self.prompt_input(
            "Hash Section Heading (e.g., 'Digest Hashes', 'Configuration Hashes', 'OS Hashes')", 
            self.config.get('hash_section_heading', 'Digest Hashes')
        )
        self.config['hash_section_heading'] = hash_heading
        
        # Ask for number of hash types
        while True:
            try:
                hash_count = int(input("How many types of hashes do you want to add? (0-20): ").strip())
                if 0 <= hash_count <= 20:
                    break
                else:
                    print("Please enter a number between 0 and 20.")
            except ValueError:
                print("Please enter a valid number.")
        
        if hash_count == 0:
            print("No hash fields will be added.")
            return
        
        print(f"\nYou will add {hash_count} hash type(s) under '{hash_heading}'.")
        
        # Collect hash information
        for i in range(hash_count):
            print(f"\n--- Hash Type {i+1} ---")
            
            hash_name = input(f"Hash {i+1} name (e.g., 'FortiAP OS Hash'): ").strip()
            if not hash_name:
                hash_name = f"Hash {i+1}"
            
            hash_value = input(f"Hash {i+1} value: ").strip()
            
            self.config['hash_fields'].append({
                'hash_name': hash_name,
                'hash_value': hash_value
            })
            
            print(f"Added: {hash_name} = {hash_value}")
        
        print(f"Total hash fields added: {len(self.config['hash_fields'])}")
    
    def collect_screenshots(self):
        """Collect screenshot file paths with optional capture"""
        print("\n=== Screenshots & Evidence ===")
        
        # Ask if user wants to capture new screenshots
        capture_new = input("Do you want to capture new screenshots? (y/n): ").lower().startswith('y')
        
        if capture_new and SCREENSHOT_AVAILABLE:
            print("Capturing screenshots with annotation capabilities...")
            try:
                # Try to capture a screenshot
                screenshot_path = ScreenshotCaptureTool.capture_screenshot_with_annotation()
                if screenshot_path:
                    self.screenshots.append(screenshot_path)
                    print(f"Captured and annotated screenshot: {screenshot_path}")
                else:
                    print("Screenshot capture cancelled or failed.")
            except Exception as e:
                print(f"Error capturing screenshot: {e}")
                print("Falling back to manual file input...")
        
        print("Enter screenshot file paths (one per line, empty line to finish):")
        
        while True:
            screenshot_path = input("Screenshot file path: ").strip()
            if not screenshot_path:
                break
            
            if os.path.exists(screenshot_path):
                self.screenshots.append(screenshot_path)
                print(f"Added: {screenshot_path}")
                
                # Offer annotation for existing screenshots
                if SCREENSHOT_AVAILABLE:
                    annotate = input("Annotate this screenshot? (y/n): ").lower().startswith('y')
                    if annotate:
                        try:
                            annotated_path = ScreenshotCaptureTool.annotate_existing_screenshot(screenshot_path)
                            if annotated_path and annotated_path != screenshot_path:
                                # Replace with annotated version
                                self.screenshots[-1] = annotated_path
                                print(f"Annotated screenshot saved: {annotated_path}")
                        except Exception as e:
                            print(f"Error annotating screenshot: {e}")
            else:
                print(f"Warning: File not found: {screenshot_path}")
                continue_anyway = input("Add anyway? (y/n): ").lower().startswith('y')
                if continue_anyway:
                    self.screenshots.append(screenshot_path)
        
        print(f"Total screenshots added: {len(self.screenshots)}")
    
    def collect_test_scenarios(self):
        """Collect test scenarios and execution data"""
        print("\n=== Test Scenarios & Execution ===")
        
        # Test Plan
        test_plan = input("Enter test plan overview (optional): ").strip()
        self.config['test_plan'] = test_plan
        
        # Test Scenarios
        self.config['test_scenarios'] = []
        print("\nEnter test scenarios (empty title to finish):")
        
        while True:
            title = input("Scenario title (or empty to finish): ").strip()
            if not title:
                break
                
            description = input(f"Description for '{title}': ").strip()
            
            # Collect steps
            steps = []
            print(f"Enter steps for '{title}' (empty step to finish):")
            while True:
                step = input(f"Step {len(steps) + 1}: ").strip()
                if not step:
                    break
                steps.append(step)
            
            scenario = {
                'title': title,
                'description': description,
                'steps': steps
            }
            
            self.config['test_scenarios'].append(scenario)
            print(f"Added scenario: {title} with {len(steps)} steps")
        
        print(f"Total scenarios added: {len(self.config['test_scenarios'])}")
        
        # Expected Results
        expected_results = input("Enter expected results for pass: ").strip()
        self.config['expected_results'] = expected_results
        
        # Generate execution blocks
        if TEST_SCENARIOS_AVAILABLE and self.config['test_scenarios']:
            generate_execution = input("Generate execution blocks from scenarios? (y/n): ").lower().startswith('y')
            if generate_execution:
                self.config['execution_blocks'] = TestScenariosManager.generate_execution_from_scenarios(self.config['test_scenarios'])
                print(f"Generated {len(self.config['execution_blocks'])} execution blocks")
    
    def run_system_commands(self):
        """Run system commands and capture output"""
        print("\n=== System Commands ===")
        
        commands = [
            ("ifconfig", "Network interface configuration"),
            ("netstat -tulpn", "Network connections and listening ports"),
            ("ps aux", "Running processes"),
            ("uname -a", "System information")
        ]
        
        self.config['command_outputs'] = {}
        
        for cmd, description in commands:
            print(f"\nRunning: {description} ({cmd})")
            run_cmd = input(f"Run '{cmd}'? (y/n): ").lower().startswith('y')
            
            if run_cmd:
                try:
                    result = subprocess.run(
                        cmd, 
                        shell=True, 
                        capture_output=True, 
                        text=True, 
                        timeout=30
                    )
                    
                    output = result.stdout
                    if result.stderr:
                        output += f"\n--- STDERR ---\n{result.stderr}"
                    
                    self.config['command_outputs'][cmd] = {
                        'description': description,
                        'output': output,
                        'return_code': result.returncode,
                        'timestamp': datetime.now().isoformat()
                    }
                    
                    print(f"Command completed with return code: {result.returncode}")
                    
                except subprocess.TimeoutExpired:
                    print("Command timed out after 30 seconds")
                    self.config['command_outputs'][cmd] = {
                        'description': description,
                        'output': 'Command timed out',
                        'return_code': -1,
                        'timestamp': datetime.now().isoformat()
                    }
                except Exception as e:
                    print(f"Error running command: {e}")
                    self.config['command_outputs'][cmd] = {
                        'description': description,
                        'output': f'Error: {str(e)}',
                        'return_code': -1,
                        'timestamp': datetime.now().isoformat()
                    }
    
    def generate_report(self):
        """Generate the security assessment report using enhanced document generator"""
        print("\n=== Generating Enhanced Report ===")
        
        # Get output file path
        default_filename = f"security_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.prompt_input(
            "Output file path", 
            default_filename
        )
        
        try:
            # Prepare configuration data for enhanced document generation
            config_data = self.prepare_enhanced_config_data()
            
            # Use enhanced document generator if available
            if ENHANCED_DOCUMENT_AVAILABLE:
                print("Using enhanced document generator...")
                success, message = EnhancedDocumentManager.generate_enhanced_report(config_data, output_path)
                if success:
                    print(f"\n✅ Enhanced report successfully generated: {output_path}")
                    print(f"📄 File size: {os.path.getsize(output_path)} bytes")
                    return output_path
                else:
                    print(f"❌ Error generating enhanced report: {message}")
                    print("Falling back to basic document generation...")
                    return self.generate_basic_report(output_path)
            else:
                print("Enhanced document generator not available, using basic generation...")
                return self.generate_basic_report(output_path)
            
        except Exception as e:
            print(f"❌ Error generating report: {e}")
            return None
            
    def prepare_enhanced_config_data(self):
        """Prepare configuration data for enhanced document generation"""
        config_data = {
            'report_title': self.config.get('report_title', 'Security Assessment Report'),
            'report_number': self.config.get('report_number', ''),
            'screenshots': self.screenshots,
            'interfaces': self.config.get('interfaces', []),
            'dut_fields': self.config.get('dut_fields', []),
            'hash_fields': self.config.get('hash_fields', []),
            'test_scenarios': self.config.get('test_scenarios', []),
            'execution_blocks': self.config.get('execution_blocks', []),
            'test_plan': self.config.get('test_plan', ''),
            'expected_results': self.config.get('expected_results', ''),
            'test_objective': self.config.get('test_objective', ''),
            'req_description': self.config.get('req_description', ''),
            'dut_configuration': self.config.get('dut_configuration', ''),
            'preconditions': self.config.get('preconditions', ''),
            'evidence_format': self.config.get('evidence_format', ''),
            'itsar_section': self.config.get('itsar_section', ''),
            'security_req': self.config.get('security_req', ''),
            'machine_ip': self.config.get('machine_ip', ''),
            'target_ip': self.config.get('target_ip', ''),
            'ssh_username': self.config.get('ssh_username', 'admin'),
            'ssh_password': self.config.get('ssh_password', ''),
            'tools_required': self.config.get('tools_required', [
                "Firefox (Latest Version)",
                "Burp Suite Professional", 
                "Nmap",
                "Nessus Scanner"
            ])
        }
        
        return config_data
        
    def generate_basic_report(self, output_path):
        """Generate basic DOCX report (fallback method)"""
        try:
            print("Generating basic report...")
            
            # Create document
            doc = Document()
            
            # Set up document properties
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            # Title
            title = doc.add_heading('Security Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Report Information
            doc.add_heading('Report Information', level=1)
            
            report_info = [
                ('Test Case Title', self.config.get('report_title', 'N/A')),
                ('Report Number', self.config.get('report_number', 'N/A')),
                ('ITSAR Section', self.config.get('itsar_section', 'N/A')),
                ('Security Requirement', self.config.get('requirement_number', 'N/A')),
                ('Test Environment', self.config.get('test_environment', 'N/A')),
                ('Test Date', self.config.get('test_date', 'N/A'))
            ]
            
            table = doc.add_table(rows=len(report_info), cols=2)
            
            for i, (label, value) in enumerate(report_info):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = value
            
            # DUT Information
            if any(self.config.get(key) for key in ['dut_model', 'dut_version', 'dut_serial', 'dut_ip']) or self.config.get('dut_fields'):
                doc.add_heading('Device Under Test (DUT) Configuration', level=1)
                
                dut_info = []
                # Add dynamic DUT fields
                for field in self.config.get('dut_fields', []):
                    dut_info.append((field['label'], field['value']))
                
                # Add IP address
                dut_info.append(('IP Address', self.config.get('dut_ip', 'N/A')))
                
                dut_table = doc.add_table(rows=len(dut_info), cols=2)
                
                for i, (label, value) in enumerate(dut_info):
                    dut_table.cell(i, 0).text = label
                    dut_table.cell(i, 1).text = value
            
            # Hash Information
            if self.config.get('hash_fields'):
                doc.add_heading(self.config.get('hash_section_heading', 'Digest Hashes'), level=1)
                
                hash_info = []
                for hash_field in self.config.get('hash_fields', []):
                    hash_info.append((hash_field['hash_name'], hash_field['hash_value']))
                
                hash_table = doc.add_table(rows=len(hash_info), cols=2)
                
                for i, (hash_name, hash_value) in enumerate(hash_info):
                    hash_table.cell(i, 0).text = hash_name
                    hash_table.cell(i, 1).text = hash_value
            
            # Command Outputs
            if self.config.get('command_outputs'):
                doc.add_heading('System Command Outputs', level=1)
                
                for cmd, cmd_data in self.config['command_outputs'].items():
                    doc.add_heading(f"{cmd_data['description']}", level=2)
                    doc.add_paragraph(f"Command: {cmd}")
                    doc.add_paragraph(f"Timestamp: {cmd_data['timestamp']}")
                    doc.add_paragraph(f"Return Code: {cmd_data['return_code']}")
                    
                    # Add command output in a code block style
                    output_para = doc.add_paragraph()
                    output_run = output_para.add_run(cmd_data['output'])
                    output_run.font.name = 'Courier New'
                    output_run.font.size = Pt(9)
            
            # Screenshots
            if self.screenshots:
                doc.add_heading('Screenshots & Evidence', level=1)
                
                for i, screenshot_path in enumerate(self.screenshots, 1):
                    doc.add_heading(f'Screenshot {i}', level=2)
                    doc.add_paragraph(f'File: {os.path.basename(screenshot_path)}')
                    
                    if os.path.exists(screenshot_path):
                        try:
                            doc.add_picture(screenshot_path, width=Inches(6))
                        except Exception as e:
                            doc.add_paragraph(f'Error inserting image: {str(e)}')
                    else:
                        doc.add_paragraph(f'Image file not found: {screenshot_path}')
            
            # Test Scenarios and Execution
            if self.config.get('test_plan') or self.config.get('test_scenarios') or self.config.get('expected_results'):
                doc.add_heading('Test Scenarios & Execution', level=1)
                
                if self.config.get('test_plan'):
                    doc.add_heading('Test Plan Overview', level=2)
                    doc.add_paragraph(self.config['test_plan'])
                
                if self.config.get('test_scenarios'):
                    doc.add_heading('Test Scenarios', level=2)
                    for i, scenario in enumerate(self.config['test_scenarios'], 1):
                        doc.add_heading(f'{i}. {scenario.get("title", "Untitled Scenario")}', level=3)
                        if scenario.get('description'):
                            doc.add_paragraph(f'Description: {scenario["description"]}')
                        if scenario.get('steps'):
                            doc.add_paragraph('Steps:')
                            for j, step in enumerate(scenario['steps'], 1):
                                doc.add_paragraph(f'{j}. {step}')
                
                if self.config.get('expected_results'):
                    doc.add_heading('Expected Results for Pass', level=2)
                    doc.add_paragraph(self.config['expected_results'])
            
            # Save document
            doc.save(output_path)
            print(f"\n✅ Basic report successfully generated: {output_path}")
            print(f"📄 File size: {os.path.getsize(output_path)} bytes")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating basic report: {e}")
            return None
    
    def interactive_mode(self):
        """Run in interactive mode"""
        print("Security Assessment Report Generator - CLI Version")
        print("=" * 50)
        
        while True:
            print("\nOptions:")
            print("1. Collect Report Information")
            print("2. Collect DUT Configuration")
            print("3. Collect Hash Information")
            print("4. Add Screenshots")
            print("5. Collect Test Scenarios")
            print("6. Run System Commands")
            print("7. Generate Report")
            print("8. Save Configuration")
            print("9. Load Configuration")
            print("10. Show Current Configuration")
            print("11. Exit")
            
            choice = input("\nSelect option (1-11): ").strip()
            
            if choice == '1':
                self.collect_report_info()
            elif choice == '2':
                self.collect_dut_info()
            elif choice == '3':
                self.collect_hash_info()
            elif choice == '4':
                self.collect_screenshots()
            elif choice == '5':
                self.collect_test_scenarios()
            elif choice == '6':
                self.run_system_commands()
            elif choice == '7':
                result = self.generate_report()
                if result:
                    print(f"Report generated successfully: {result}")
            elif choice == '8':
                self.save_config()
            elif choice == '9':
                self.load_config()
            elif choice == '10':
                print("\nCurrent Configuration:")
                print(json.dumps(self.config, indent=2, default=str))
            elif choice == '11':
                print("Exiting...")
                break
            else:
                print("Invalid option. Please try again.")

def main():
    """Main entry point"""
    generator = SecurityReportGeneratorCLI()
    
    if len(sys.argv) > 1:
        if sys.argv[1] == '--auto':
            # Auto mode - collect all info and generate report
            generator.collect_report_info()
            generator.collect_dut_info()
            generator.collect_hash_info()
            generator.collect_screenshots()
            generator.collect_test_scenarios()
            generator.run_system_commands()
            result = generator.generate_report()
            if result:
                print(f"Report generated: {result}")
        else:
            print("Usage: python security_report_cli.py [--auto]")
    else:
        # Interactive mode
        generator.interactive_mode()

if __name__ == '__main__':
    main()